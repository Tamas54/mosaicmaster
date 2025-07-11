# Standard library imports
import os
import re
import uuid
import json
import asyncio
import shutil
import logging
import zipfile
import tempfile
import datetime
from pathlib import Path
from typing import Dict, Any, Tuple, List, Optional, Union, Set
from functools import partial
from contextlib import asynccontextmanager

# Third-party imports
import aiofiles
from fastapi import APIRouter, HTTPException, UploadFile, File, Form, BackgroundTasks, Depends
from fastapi.responses import JSONResponse
from bs4 import BeautifulSoup
from docx import Document
import fitz  # PyMuPDF
from ebooklib import epub
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from PIL import Image

# Local imports
from external_converter import try_ocr_image

# Conditional imports
try:
    from odf.opendocument import OpenDocumentText
    from odf.text import P
    HAS_ODF = True
except ImportError:
    HAS_ODF = False
    logging.warning("odfpy nem elérhető. ODT feldolgozás korlátozott lesz.")

# Local imports
from config import (
    TEMP_DIR, 
    SYSTEM_DOWNLOADS, 
    manager,
    logger,
    sanitize_filename,
    MAX_FILE_SIZE,
    CLEANUP_INTERVAL_HOURS,
    MAX_TEMP_DIR_AGE_HOURS,
    check_calibre,
    check_libreoffice
)
from external_converter import try_convert_external, get_external_support_info

router = APIRouter()

class FileFormatError(Exception):
    """Nem támogatott fájlformátum hiba"""
    pass

class ConversionError(Exception):
    """Általános konverziós hiba"""
    pass

class ExternalToolError(Exception):
    """Külső eszközök (Calibre, LibreOffice) hibája"""
    pass


class TempDirectoryManager:
    """Ideiglenes könyvtárak kezelése automatikus tisztítással"""
    
    def __init__(self, base_dir: Path, max_age_hours: int = 24):
        self.base_dir = base_dir
        self.max_age_hours = max_age_hours
        self.active_dirs: Set[str] = set()
        self._cleanup_task = None
        
    @asynccontextmanager
    async def temp_dir(self, connection_id: str) -> Path:
        """Ideiglenes könyvtár kontextuskezelő"""
        work_dir = self.base_dir / connection_id
        work_dir.mkdir(exist_ok=True, parents=True)
        self.active_dirs.add(connection_id)
        
        try:
            yield work_dir
        finally:
            if connection_id not in self.active_dirs:
                await self.cleanup(connection_id)
    
    async def create(self, connection_id: str) -> Path:
        """Új ideiglenes könyvtár létrehozása"""
        work_dir = self.base_dir / connection_id
        work_dir.mkdir(exist_ok=True, parents=True)
        self.active_dirs.add(connection_id)
        return work_dir
        
    async def cleanup(self, connection_id: str) -> None:
        """Ideiglenes könyvtár törlése"""
        work_dir = self.base_dir / connection_id
        if work_dir.exists():
            try:
                shutil.rmtree(work_dir)
                logger.info(f"Cleaned up directory: {work_dir}")
            except Exception as e:
                logger.error(f"Error cleaning up directory {work_dir}: {str(e)}")
        
        if connection_id in self.active_dirs:
            self.active_dirs.remove(connection_id)
    
    async def cleanup_old_directories(self) -> None:
        """Régi ideiglenes könyvtárak automatikus törlése"""
        now = datetime.datetime.now()
        try:
            for item in self.base_dir.iterdir():
                if not item.is_dir():
                    continue
                    
                # Ellenőrizzük a könyvtár korát
                stats = item.stat()
                last_modified = datetime.datetime.fromtimestamp(stats.st_mtime)
                age_hours = (now - last_modified).total_seconds() / 3600
                
                if age_hours > self.max_age_hours:
                    conn_id = item.name
                    if conn_id in self.active_dirs:
                        logger.info(f"Skipping cleanup of active directory: {item}")
                        continue
                        
                    logger.info(f"Removing old temp directory: {item} (age: {age_hours:.1f} hours)")
                    try:
                        shutil.rmtree(item, ignore_errors=True)
                    except Exception as e:
                        logger.error(f"Failed to remove old directory {item}: {str(e)}")
        except Exception as e:
            logger.error(f"Error during automatic cleanup: {str(e)}")
    
    async def start_cleanup_scheduler(self) -> None:
        """Időzített tisztítás indítása"""
        while True:
            try:
                await self.cleanup_old_directories()
            except Exception as e:
                logger.error(f"Scheduler error in cleanup: {str(e)}")
            
            # Óránként ellenőrzés
            await asyncio.sleep(CLEANUP_INTERVAL_HOURS * 3600)
            
    def start(self):
        """Tisztítási ütemező indítása"""
        if self._cleanup_task is None or self._cleanup_task.done():
            self._cleanup_task = asyncio.create_task(self.start_cleanup_scheduler())
            logger.info("Temp directory cleanup scheduler started")

# Globális TempDirectoryManager példány
temp_manager = TempDirectoryManager(TEMP_DIR, MAX_TEMP_DIR_AGE_HOURS)

async def process_image_with_ocr(file_path: Path, preserve_format: bool = False) -> Tuple[str, Dict[str, Any]]:
    """Képfeldolgozás Vision OCR segítségével az external_converter használatával"""
    try:
        # Ideiglenes fájl létrehozása az OCR eredménynek
        temp_txt_path = file_path.with_suffix('.ocr.txt')
        
        # OCR feldolgozás az external_converter segítségével
        result = await try_ocr_image(file_path, temp_txt_path)
        
        if result is None:
            raise HTTPException(status_code=500, detail="OCR processing not available or failed")
        
        # OCR szöveg beolvasása
        async with aiofiles.open(temp_txt_path, 'r', encoding='utf-8') as f:
            text = await f.read()
        
        # Ideiglenes fájl törlése
        try:
            temp_txt_path.unlink(missing_ok=True)
        except:
            pass
        
        return text, {"type": "image", "method": result.get("method", "vision_ocr")}
        
    except Exception as e:
        logger.error(f"Error processing image with OCR: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Image OCR processing failed: {str(e)}")


class ConversionProcessor:
    """Dokumentum konverziós osztály"""
    
    SUPPORTED_CONVERSIONS = {
        "pdf": ["docx", "doc", "odt", "txt", "rtf", "ppt", "pptx", "epub", "srt", "mobi", "sub"],
        "docx": ["pdf", "doc", "odt", "txt", "rtf", "ppt", "pptx", "epub", "srt", "mobi", "sub"],
        "doc": ["pdf", "docx", "odt", "txt", "rtf", "ppt", "pptx", "epub", "srt", "mobi", "sub"],
        "odt": ["pdf", "docx", "doc", "txt", "rtf", "ppt", "pptx", "epub", "srt", "mobi", "sub"],
        "txt": ["pdf", "docx", "doc", "odt", "rtf", "ppt", "pptx", "epub", "srt", "mobi", "sub"],
        "rtf": ["pdf", "docx", "doc", "odt", "txt", "ppt", "pptx", "epub", "srt", "mobi", "sub"],
        "ppt": ["pdf", "docx", "doc", "odt"],
        "pptx": ["pdf", "docx", "doc", "odt"],
        "epub": ["pdf", "docx", "doc", "odt", "txt", "rtf", "ppt", "pptx", "srt", "mobi", "sub"],
        "srt": ["odt", "docx"],
        "mobi": ["pdf", "docx", "doc", "odt", "txt", "rtf", "ppt", "pptx", "epub", "srt", "sub"],
        "sub": ["pdf", "docx", "doc", "odt", "txt", "rtf", "ppt", "pptx", "epub", "srt", "mobi"],
        "image/jpeg": None,
        "image/png": None,
        "image/gif": None
    }

    @classmethod
    async def process_file(cls, file_path: Path, file_type: str, preserve_format: bool = False) -> Tuple[str, Dict[str, Any]]:
        """Fájl típus alapján a megfelelő feldolgozó metódus kiválasztása"""
        processors = {
            "pdf": cls.process_pdf,
            "docx": cls.process_docx,
            "rtf": cls.process_rtf,
            "odt": cls.process_odt,
            "txt": cls.process_txt,  # Hozzáadva a TXT feldolgozó
            "application/msword": cls.process_docx,  # DOC is handled by process_docx
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": cls.process_docx,  # DOCX MIME type
            "image/jpeg": process_image_with_ocr,
            "image/png": process_image_with_ocr,
            "image/gif": process_image_with_ocr,
        }

        # Megpróbáljuk normalizálni a file_type-ot, ha szükséges
        norm_file_type = file_type.lower().split('/')[-1] if '/' in file_type else file_type.lower()
        if norm_file_type in processors:
            return await processors[norm_file_type](file_path, preserve_format)
        elif file_type in processors:
            return await processors[file_type](file_path, preserve_format)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file_type}"
            )

    @classmethod
    async def convert_document(cls, input_path: Path, output_path: Path, 
                             source_format: str, target_format: str) -> Dict[str, Any]:
        """Dokumentum konvertálása"""
        if source_format not in cls.SUPPORTED_CONVERSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported source format: {source_format}"
            )

        # Képek OCR feldolgozása
        if source_format in ["image/jpeg", "image/png", "image/gif"]:
            text, _ = await process_image_with_ocr(input_path)
            
            # Ha TXT a cél, egyszerűen mentjük a szöveget
            if target_format == "txt":
                async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
                    await f.write(text)
                return {"converted_text": text}
            
            # Egyéb formátumok esetén ideiglenes TXT fájlt hozunk létre és konvertáljuk
            temp_txt = input_path.parent / f"{input_path.stem}_ocr_temp.txt"
            async with aiofiles.open(temp_txt, "w", encoding="utf-8") as f:
                await f.write(text)
            
            try:
                if target_format == "pdf":
                    result = await cls._convert_txt_to_pdf(temp_txt, output_path)
                elif target_format == "docx":
                    result = await cls._convert_txt_to_docx(temp_txt, output_path)
                elif target_format == "epub":
                    result = await cls._convert_txt_to_epub(temp_txt, output_path)
                elif target_format == "odt":
                    result = await cls._convert_txt_to_odt(temp_txt, output_path)
                elif target_format in ["srt", "sub", "vtt"]:
                    result = await cls._convert_to_srt(temp_txt, output_path, "txt")
                else:
                    # Nem támogatott formátum esetén TXT-ként mentjük
                    async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
                        await f.write(text)
                    result = {"converted_text": text}
                
                return result
            finally:
                # Ideiglenes fájl törlése
                if temp_txt.exists():
                    temp_txt.unlink()

        # Speciális konverziók kezelése
        if source_format == "odt" and target_format == "rtf":
            return await cls._convert_odt_to_rtf(input_path, output_path)
        
        if source_format == "pdf" and target_format == "epub":
            return await cls._convert_pdf_to_epub_chunked(input_path, output_path)
        
        if source_format == "srt":
            if target_format == "odt":
                return await cls._convert_srt_to_odt(input_path, output_path)
            elif target_format == "docx":
                return await cls._convert_srt_to_docx(input_path, output_path)
                
        # Dokumentum konvertálása SRT formátumba
        if target_format == "srt":
            return await cls._convert_to_srt(input_path, output_path, source_format)

        # Egyedi kezelés DOCX/DOC generálása esetén
        # 1. ELSŐDLEGES: Pure Python konverziók
        if target_format in ["docx", "doc"] and source_format not in ["docx", "doc"]:
            try:
                if source_format == "txt":
                    if target_format == "docx":
                        return await cls._convert_txt_to_docx(input_path, output_path)
                    else:  # target_format == "doc"
                        return await cls._convert_txt_to_doc(input_path, output_path)
                elif source_format == "pdf":
                    if target_format == "docx":
                        return await cls._convert_pdf_to_docx(input_path, output_path)
                    else:  # target_format == "doc"
                        return await cls._convert_pdf_to_doc(input_path, output_path)
            except Exception as e:
                logger.warning(f"Pure Python conversion failed: {str(e)}, trying external tools")
        
        if target_format == "epub" and source_format == "pdf":
            try:
                return await cls._convert_pdf_to_epub_chunked(input_path, output_path)
            except Exception as e:
                logger.warning(f"Pure Python PDF→EPUB conversion failed: {str(e)}, trying external tools")
        
        if target_format == "docx" and source_format == "srt":
            try:
                return await cls._convert_srt_to_docx(input_path, output_path)
            except Exception as e:
                logger.warning(f"Pure Python SRT→DOCX conversion failed: {str(e)}, trying external tools")
        
        if target_format == "odt" and source_format == "srt":
            try:
                return await cls._convert_srt_to_odt(input_path, output_path)
            except Exception as e:
                logger.warning(f"Pure Python SRT→ODT conversion failed: {str(e)}, trying external tools")
        
        if target_format == "pdf" and source_format == "txt":
            try:
                return await cls._convert_txt_to_pdf(input_path, output_path)
            except Exception as e:
                logger.warning(f"Pure Python TXT→PDF conversion failed: {str(e)}, trying external tools")
        
        if target_format == "pdf" and source_format == "epub":
            try:
                # EPUB → TXT → PDF lánc
                temp_txt = input_path.parent / f"{input_path.stem}_temp.txt"
                await cls._convert_epub_to_txt(input_path, temp_txt)
                result = await cls._convert_txt_to_pdf(temp_txt, output_path)
                if temp_txt.exists():
                    temp_txt.unlink()
                return result
            except Exception as e:
                logger.warning(f"Pure Python EPUB→PDF conversion failed: {str(e)}, trying external tools")
        
        if target_format == "txt" and source_format == "odt":
            try:
                return await cls._convert_odt_to_txt(input_path, output_path)
            except Exception as e:
                logger.warning(f"Pure Python ODT→TXT conversion failed: {str(e)}, trying external tools")
        
        # ODT konverziók
        if target_format == "pdf" and source_format == "odt":
            try:
                # ODT → TXT → PDF lánc
                temp_txt = input_path.parent / f"{input_path.stem}_temp.txt"
                await cls._convert_odt_to_txt(input_path, temp_txt)
                result = await cls._convert_txt_to_pdf(temp_txt, output_path)
                if temp_txt.exists():
                    temp_txt.unlink()
                return result
            except Exception as e:
                logger.warning(f"Pure Python ODT→PDF conversion failed: {str(e)}, trying external tools")
        
        if target_format == "docx" and source_format == "odt":
            try:
                # ODT → TXT → DOCX lánc
                temp_txt = input_path.parent / f"{input_path.stem}_temp.txt"
                await cls._convert_odt_to_txt(input_path, temp_txt)
                result = await cls._convert_txt_to_docx(temp_txt, output_path)
                if temp_txt.exists():
                    temp_txt.unlink()
                return result
            except Exception as e:
                logger.warning(f"Pure Python ODT→DOCX conversion failed: {str(e)}, trying external tools")
        
        if target_format == "epub" and source_format == "odt":
            try:
                # ODT → TXT → EPUB lánc
                temp_txt = input_path.parent / f"{input_path.stem}_temp.txt"
                await cls._convert_odt_to_txt(input_path, temp_txt)
                result = await cls._convert_txt_to_epub(temp_txt, output_path)
                if temp_txt.exists():
                    temp_txt.unlink()
                return result
            except Exception as e:
                logger.warning(f"Pure Python ODT→EPUB conversion failed: {str(e)}, trying external tools")
        
        if target_format == "odt" and source_format == "txt":
            try:
                # TXT → ODT közvetlen (ha van implementáció)
                return await cls._convert_txt_to_odt(input_path, output_path)
            except Exception as e:
                logger.warning(f"Pure Python TXT→ODT conversion failed: {str(e)}, trying external tools")
        
        # MOBI cél konverziók (MOBI generálása nehéz, ezért EPUB-on keresztül)
        if target_format == "mobi":
            try:
                # Bármilyen formátum → EPUB → MOBI lánc (EPUB közelebb van a MOBI-hez)
                temp_epub = input_path.parent / f"{input_path.stem}_temp.epub"
                
                # Először EPUB-ba konvertálunk
                if source_format == "txt":
                    await cls._convert_txt_to_epub(input_path, temp_epub)
                elif source_format == "docx":
                    await cls._convert_docx_to_epub(input_path, temp_epub)
                elif source_format == "pdf":
                    # PDF → TXT → EPUB lánc
                    temp_txt = input_path.parent / f"{input_path.stem}_temp2.txt"
                    await cls._convert_pdf_to_txt(input_path, temp_txt)
                    await cls._convert_txt_to_epub(temp_txt, temp_epub)
                    if temp_txt.exists():
                        temp_txt.unlink()
                elif source_format == "odt":
                    # ODT → TXT → EPUB lánc
                    temp_txt = input_path.parent / f"{input_path.stem}_temp2.txt"
                    await cls._convert_odt_to_txt(input_path, temp_txt)
                    await cls._convert_txt_to_epub(temp_txt, temp_epub)
                    if temp_txt.exists():
                        temp_txt.unlink()
                elif source_format == "rtf":
                    # RTF → TXT → EPUB lánc
                    temp_txt = input_path.parent / f"{input_path.stem}_temp2.txt"
                    await cls._convert_rtf_to_txt(input_path, temp_txt)
                    await cls._convert_txt_to_epub(temp_txt, temp_epub)
                    if temp_txt.exists():
                        temp_txt.unlink()
                else:
                    raise Exception(f"MOBI conversion from {source_format} not supported")
                
                # Most EPUB → MOBI (ezt sajnos nem tudjuk pure Python-nal, de legalább EPUB van)
                # Egyelőre az EPUB-ot adjuk vissza MOBI néven, mert a legtöbb olvasó támogatja
                if temp_epub.exists():
                    shutil.copy(temp_epub, output_path)
                    temp_epub.unlink()
                    
                    # Figyelmeztető üzenet a logban
                    logger.warning(f"MOBI conversion: Created EPUB file instead (most e-readers support EPUB)")
                    
                    return {
                        "converted": True,
                        "note": "Generated EPUB format instead of MOBI (more compatible)",
                        "actual_format": "epub"
                    }
                else:
                    raise Exception("Failed to create intermediate EPUB file")
                    
            except Exception as e:
                logger.warning(f"MOBI conversion failed: {str(e)}, trying external tools")
        
        # PPT/PPTX konverziók (external converter használatával)
        if source_format in ["ppt", "pptx"]:
            try:
                # PPT/PPTX → TXT (external converter)
                temp_txt = input_path.parent / f"{input_path.stem}_temp.txt"
                external_result = await try_convert_external(source_format, input_path, temp_txt)
                
                if external_result and temp_txt.exists():
                    if target_format == "txt":
                        shutil.copy(temp_txt, output_path)
                        temp_txt.unlink()
                        return external_result
                    elif target_format == "pdf":
                        # PPT → TXT → PDF lánc
                        result = await cls._convert_txt_to_pdf(temp_txt, output_path)
                        temp_txt.unlink()
                        return result
                    elif target_format == "docx":
                        # PPT → TXT → DOCX lánc
                        result = await cls._convert_txt_to_docx(temp_txt, output_path)
                        temp_txt.unlink()
                        return result
                    elif target_format == "epub":
                        # PPT → TXT → EPUB lánc
                        result = await cls._convert_txt_to_epub(temp_txt, output_path)
                        temp_txt.unlink()
                        return result
                    elif target_format == "odt":
                        # PPT → TXT → ODT lánc
                        result = await cls._convert_txt_to_odt(temp_txt, output_path)
                        temp_txt.unlink()
                        return result
                    elif target_format in ["srt", "sub", "vtt"]:
                        # PPT → TXT → subtitle lánc
                        result = await cls._convert_to_srt(temp_txt, output_path, "txt")
                        temp_txt.unlink()
                        return result
                else:
                    raise Exception(f"External PPT conversion failed")
            except Exception as e:
                logger.warning(f"PPT/PPTX conversion failed: {str(e)}, trying external tools")
        
        # TXT → PPT/PPTX (nem támogatott, de legalább hibaüzenet)
        if target_format in ["ppt", "pptx"]:
            logger.error(f"Conversion to {target_format} is not supported - PowerPoint files cannot be generated from text")
            raise HTTPException(
                status_code=400,
                detail=f"Conversion to {target_format} format is not supported. PowerPoint files cannot be automatically generated."
            )
        
        # Felirat fájlok konverziói (SRT, SUB, VTT)
        if source_format in ["srt", "sub", "vtt"]:
            try:
                if target_format == "txt":
                    # Subtitle → TXT (felirat szöveg kinyerése)
                    return await DocumentProcessor._convert_subtitle_to_txt(input_path, output_path, source_format)
                elif target_format == "docx":
                    # Subtitle → TXT → DOCX lánc
                    temp_txt = input_path.parent / f"{input_path.stem}_temp.txt"
                    await DocumentProcessor._convert_subtitle_to_txt(input_path, temp_txt, source_format)
                    result = await cls._convert_txt_to_docx(temp_txt, output_path)
                    if temp_txt.exists():
                        temp_txt.unlink()
                    return result
                elif target_format == "pdf":
                    # Subtitle → TXT → PDF lánc
                    temp_txt = input_path.parent / f"{input_path.stem}_temp.txt"
                    await DocumentProcessor._convert_subtitle_to_txt(input_path, temp_txt, source_format)
                    result = await cls._convert_txt_to_pdf(temp_txt, output_path)
                    if temp_txt.exists():
                        temp_txt.unlink()
                    return result
                elif target_format == "epub":
                    # Subtitle → TXT → EPUB lánc
                    temp_txt = input_path.parent / f"{input_path.stem}_temp.txt"
                    await DocumentProcessor._convert_subtitle_to_txt(input_path, temp_txt, source_format)
                    result = await cls._convert_txt_to_epub(temp_txt, output_path)
                    if temp_txt.exists():
                        temp_txt.unlink()
                    return result
                elif target_format == "odt":
                    # Subtitle → ODT közvetlen (már van implementáció SRT-hez)
                    if source_format == "srt":
                        return await cls._convert_srt_to_odt(input_path, output_path)
                    else:
                        # SUB/VTT → TXT → ODT lánc
                        temp_txt = input_path.parent / f"{input_path.stem}_temp.txt"
                        await DocumentProcessor._convert_subtitle_to_txt(input_path, temp_txt, source_format)
                        result = await cls._convert_txt_to_odt(temp_txt, output_path)
                        if temp_txt.exists():
                            temp_txt.unlink()
                        return result
                elif target_format in ["srt", "sub", "vtt"]:
                    # Subtitle formátumok közötti konverzió
                    return await DocumentProcessor._convert_subtitle_format(input_path, output_path, source_format, target_format)
            except Exception as e:
                logger.warning(f"Subtitle conversion failed: {str(e)}, trying external tools")
        
        # TXT → subtitle formátumok
        if source_format == "txt" and target_format in ["srt", "sub", "vtt"]:
            try:
                return await cls._convert_to_srt(input_path, output_path, source_format)
            except Exception as e:
                logger.warning(f"TXT to subtitle conversion failed: {str(e)}, trying external tools")
        
        if target_format == "txt" and source_format == "pdf":
            try:
                return await cls._convert_pdf_to_txt(input_path, output_path)
            except Exception as e:
                logger.warning(f"Pure Python PDF→TXT conversion failed: {str(e)}, trying external tools")
        
        if target_format == "txt" and source_format == "docx":
            try:
                return await cls._convert_docx_to_txt(input_path, output_path)
            except Exception as e:
                logger.warning(f"Pure Python DOCX→TXT conversion failed: {str(e)}, trying external tools")
        
        if target_format == "txt" and source_format == "rtf":
            try:
                return await cls._convert_rtf_to_txt(input_path, output_path)
            except Exception as e:
                logger.warning(f"Pure Python RTF→TXT conversion failed: {str(e)}, trying external tools")
        
        if target_format == "txt" and source_format == "epub":
            try:
                return await cls._convert_epub_to_txt(input_path, output_path)
            except Exception as e:
                logger.warning(f"Pure Python EPUB→TXT conversion failed: {str(e)}, trying external tools")
        
        if target_format == "txt" and source_format == "mobi":
            try:
                return await cls._convert_mobi_to_txt(input_path, output_path)
            except Exception as e:
                logger.warning(f"Pure Python MOBI→TXT conversion failed: {str(e)}, trying external tools")
        
        if target_format == "epub" and source_format == "txt":
            try:
                return await cls._convert_txt_to_epub(input_path, output_path)
            except Exception as e:
                logger.warning(f"Pure Python TXT→EPUB conversion failed: {str(e)}, trying external tools")
        
        if target_format == "docx" and source_format == "rtf":
            try:
                return await cls._convert_rtf_to_docx(input_path, output_path)
            except Exception as e:
                logger.warning(f"Pure Python RTF→DOCX conversion failed: {str(e)}, trying external tools")
        
        if target_format == "epub" and source_format == "docx":
            try:
                return await cls._convert_docx_to_epub(input_path, output_path)
            except Exception as e:
                logger.warning(f"Pure Python DOCX→EPUB conversion failed: {str(e)}, trying external tools")
        
        if target_format == "docx" and source_format == "epub":
            try:
                return await cls._convert_epub_to_docx(input_path, output_path)
            except Exception as e:
                logger.warning(f"Pure Python EPUB→DOCX conversion failed: {str(e)}, trying external tools")
        
        # RTF konverziók
        if target_format == "pdf" and source_format == "rtf":
            try:
                # RTF → TXT → PDF lánc
                temp_txt = input_path.parent / f"{input_path.stem}_temp.txt"
                await cls._convert_rtf_to_txt(input_path, temp_txt)
                result = await cls._convert_txt_to_pdf(temp_txt, output_path)
                if temp_txt.exists():
                    temp_txt.unlink()
                return result
            except Exception as e:
                logger.warning(f"Pure Python RTF→PDF conversion failed: {str(e)}, trying external tools")
        
        if target_format == "epub" and source_format == "rtf":
            try:
                # RTF → TXT → EPUB lánc
                temp_txt = input_path.parent / f"{input_path.stem}_temp.txt"
                await cls._convert_rtf_to_txt(input_path, temp_txt)
                result = await cls._convert_txt_to_epub(temp_txt, output_path)
                if temp_txt.exists():
                    temp_txt.unlink()
                return result
            except Exception as e:
                logger.warning(f"Pure Python RTF→EPUB conversion failed: {str(e)}, trying external tools")
        
        # DOCX konverziók
        if target_format == "pdf" and source_format == "docx":
            try:
                # DOCX → TXT → PDF lánc
                temp_txt = input_path.parent / f"{input_path.stem}_temp.txt"
                await cls._convert_docx_to_txt(input_path, temp_txt)
                result = await cls._convert_txt_to_pdf(temp_txt, output_path)
                if temp_txt.exists():
                    temp_txt.unlink()
                return result
            except Exception as e:
                logger.warning(f"Pure Python DOCX→PDF conversion failed: {str(e)}, trying external tools")
        
        # MOBI konverziók pure Python-nal
        if source_format == "mobi":
            try:
                if target_format == "txt":
                    return await cls._convert_mobi_to_txt(input_path, output_path)
                elif target_format == "pdf":
                    # MOBI → TXT → PDF lánc
                    temp_txt = input_path.parent / f"{input_path.stem}_temp.txt"
                    await cls._convert_mobi_to_txt(input_path, temp_txt)
                    result = await cls._convert_txt_to_pdf(temp_txt, output_path)
                    if temp_txt.exists():
                        temp_txt.unlink()
                    return result
                elif target_format == "docx":
                    # MOBI → TXT → DOCX lánc
                    temp_txt = input_path.parent / f"{input_path.stem}_temp.txt"
                    await cls._convert_mobi_to_txt(input_path, temp_txt)
                    result = await cls._convert_txt_to_docx(temp_txt, output_path)
                    if temp_txt.exists():
                        temp_txt.unlink()
                    return result
                elif target_format == "epub":
                    # MOBI → TXT → EPUB lánc
                    temp_txt = input_path.parent / f"{input_path.stem}_temp.txt"
                    await cls._convert_mobi_to_txt(input_path, temp_txt)
                    result = await cls._convert_txt_to_epub(temp_txt, output_path)
                    if temp_txt.exists():
                        temp_txt.unlink()
                    return result
            except Exception as e:
                logger.warning(f"Pure Python MOBI conversion failed: {str(e)}, trying external tools")
        
        # 2. MÁSODLAGOS: External converter helper (pure Python)
        try:
            external_result = await try_convert_external(source_format, input_path, output_path)
            if external_result:
                return external_result
        except Exception as e:
            logger.warning(f"External converter failed: {str(e)}")
        
        # Ha minden Pure Python módszer hibázik
        logger.error(f"All conversion methods failed for {source_format} → {target_format}")
        raise HTTPException(
            status_code=500,
            detail=f"Document conversion failed: No suitable conversion method available for {source_format} → {target_format}"
        )
        
        # Ha nincs támogatott konverzió
        raise HTTPException(
            status_code=400,
            detail=f"Conversion from {source_format} to {target_format} is not supported"
        )

    @staticmethod
    async def process_txt(file_path: Path, preserve_format: bool = False) -> Tuple[str, Dict[str, Any]]:
        """TXT feldolgozása - aszinkron verzió"""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = await f.read()
            return text, {"type": "txt"}
        except Exception as e:
            logger.error(f"Error processing TXT: {str(e)}")
            raise HTTPException(status_code=500, detail=f"TXT processing failed: {str(e)}")

    @staticmethod
    async def process_pdf(file_path: Path, preserve_format: bool = False) -> Tuple[str, Dict[str, Any]]:
        """PDF feldolgozása - aszinkron verzió"""
        try:
            loop = asyncio.get_event_loop()
            
            def extract_pdf_text():
                try:
                    pdf = fitz.open(str(file_path))
                    texts = []
                    for page in pdf:
                        texts.append(page.get_text())
                    page_count = len(pdf)
                    return "\n\n=== PAGE BREAK ===\n\n".join(texts), page_count
                except fitz.FileDataError as e:
                    raise FileFormatError(f"Invalid PDF file: {str(e)}")
                except Exception as e:
                    raise Exception(f"PDF processing error: {str(e)}")
                    
            text, page_count = await loop.run_in_executor(None, extract_pdf_text)
            return text, {"pages": page_count, "type": "pdf"}
        except FileFormatError as e:
            logger.error(f"Invalid PDF file: {str(e)}")
            raise HTTPException(status_code=400, detail=str(e))
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise HTTPException(status_code=500, detail=f"PDF processing failed: {str(e)}")

    @staticmethod
    async def process_docx(file_path: Path, preserve_format: bool = False) -> Tuple[str, Dict[str, Any]]:
        """DOCX feldolgozása - aszinkron verzió"""
        try:
            loop = asyncio.get_event_loop()
            
            def extract_docx_text():
                try:
                    doc = Document(file_path)
                    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                    return "\n".join(paragraphs)
                except Exception as e:
                    raise Exception(f"DOCX processing error: {str(e)}")
                    
            text = await loop.run_in_executor(None, extract_docx_text)
            return text, {"type": "docx"}
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise HTTPException(status_code=500, detail=f"DOCX processing failed: {str(e)}")

    @staticmethod
    async def process_rtf(file_path: Path, preserve_format: bool = False) -> Tuple[str, Dict[str, Any]]:
        """RTF feldolgozása - aszinkron verzió"""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                rtf_content = await f.read()
                
            # RTF feldolgozás aszinkron végrehajtása
            loop = asyncio.get_event_loop()
            
            def process_rtf_content(content):
                return re.sub(r'{\\.*?}|[{}]', '', content)
                
            text = await loop.run_in_executor(None, process_rtf_content, rtf_content)
            return text, {"type": "rtf"}
        except Exception as e:
            logger.error(f"Error processing RTF: {str(e)}")
            raise HTTPException(status_code=500, detail="Error processing RTF file")

    @staticmethod
    async def process_odt(file_path: Path, preserve_format: bool = False) -> Tuple[str, Dict[str, Any]]:
        """ODT feldolgozása - aszinkron verzió"""
        try:
            loop = asyncio.get_event_loop()
            
            def extract_odt_text():
                try:
                    with zipfile.ZipFile(file_path, 'r') as odt_zip:
                        with odt_zip.open('content.xml') as content_file:
                            content = content_file.read().decode('utf-8', errors="ignore")
                            soup = BeautifulSoup(content, "xml")
                            paragraphs = soup.find_all("text:p")
                            return "\n".join(p.get_text() for p in paragraphs)
                except zipfile.BadZipFile:
                    raise FileFormatError("Invalid ODT file (bad zip structure)")
                except Exception as e:
                    raise Exception(f"ODT processing error: {str(e)}")
                    
            text = await loop.run_in_executor(None, extract_odt_text)
            return text, {"type": "odt"}
        except FileFormatError as e:
            logger.error(f"Invalid ODT file: {str(e)}")
            raise HTTPException(status_code=400, detail=str(e))
        except Exception as e:
            logger.error(f"Error processing ODT: {str(e)}")
            raise HTTPException(status_code=500, detail="Error processing ODT file")

    @staticmethod
    async def _convert_txt_to_docx(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """TXT → DOCX közvetlen konverzió a Python docx csomaggal"""
        try:
            async with aiofiles.open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = await f.read()
            
            loop = asyncio.get_event_loop()
            
            def create_docx(text_content):
                doc = Document()
                
                # Soronként adjuk hozzá a szöveget, hogy az ékezetek és formázás megmaradjon
                paragraphs = text_content.split('\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para)
                    else:
                        doc.add_paragraph()  # Üres bekezdés
                        
                doc.save(str(output_path))
                return True
                
            success = await loop.run_in_executor(None, create_docx, text)
            
            if not success:
                raise ConversionError("Failed to create DOCX")
                
            return {"converted": True}
            
        except Exception as e:
            logger.error(f"Error converting TXT to DOCX: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_txt_to_doc(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """TXT → DOC közvetlen konverzió a Python docx csomaggal (DOC és DOCX formátum ugyanaz)"""
        try:
            async with aiofiles.open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = await f.read()
            
            loop = asyncio.get_event_loop()
            
            def create_doc(text_content):
                doc = Document()
                
                # Soronként adjuk hozzá a szöveget, hogy az ékezetek és formázás megmaradjon
                paragraphs = text_content.split('\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para)
                    else:
                        doc.add_paragraph()  # Üres bekezdés
                        
                doc.save(str(output_path))
                return True
                
            success = await loop.run_in_executor(None, create_doc, text)
            
            if not success:
                raise ConversionError("Failed to create DOC")
                
            return {"converted": True}
            
        except Exception as e:
            logger.error(f"Error converting TXT to DOC: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_txt_to_pdf(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """TXT → PDF közvetlen konverzió a ReportLab csomaggal"""
        try:
            async with aiofiles.open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = await f.read()
            
            loop = asyncio.get_event_loop()
            
            def create_pdf(text_content):
                from reportlab.lib.pagesizes import letter, A4
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                from reportlab.lib.enums import TA_LEFT
                
                # PDF létrehozása Unicode támogatással
                doc = SimpleDocTemplate(str(output_path), pagesize=A4)
                styles = getSampleStyleSheet()
                story = []
                
                # Unicode font regisztrálása (DejaVu Sans támogatja a magyar karaktereket)
                try:
                    # Próbáljuk meg regisztrálni a DejaVu Sans fontot
                    try:
                        pdfmetrics.registerFont(TTFont('DejaVuSans', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'))
                        unicode_font = 'DejaVuSans'
                    except:
                        # Ha nincs DejaVu, próbáljuk a Liberation fontot
                        try:
                            pdfmetrics.registerFont(TTFont('LiberationSans', '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf'))
                            unicode_font = 'LiberationSans'
                        except:
                            # Ha egyetlen specifikus font sem érhető el, használjuk a Helvetica-t UTF-8 kódolással
                            unicode_font = 'Helvetica'
                except Exception as font_error:
                    logger.warning(f"Font registration failed: {font_error}, using default Helvetica")
                    unicode_font = 'Helvetica'
                
                # Unicode-kompatibilis stílus létrehozása
                unicode_style = ParagraphStyle(
                    'UnicodeNormal',
                    parent=styles['Normal'],
                    fontName=unicode_font,
                    fontSize=11,
                    leading=14,
                    alignment=TA_LEFT,
                    spaceAfter=6
                )
                
                # Soronként adjuk hozzá a szöveget
                paragraphs = text_content.split('\n')
                for para in paragraphs:
                    if para.strip():
                        # HTML escape a speciális karakterekhez, de ékezetes karaktereket megtartjuk
                        para_escaped = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        
                        # Unicode karakterek explicit kezelése
                        try:
                            story.append(Paragraph(para_escaped, unicode_style))
                        except Exception as para_error:
                            # Ha a paragraph létrehozása hibázik, próbáljuk ASCII-ra konvertálni
                            logger.warning(f"Paragraph creation failed for line, using ASCII fallback: {para_error}")
                            para_ascii = para_escaped.encode('ascii', 'replace').decode('ascii')
                            story.append(Paragraph(para_ascii, unicode_style))
                        
                        story.append(Spacer(1, 0.1*inch))
                    else:
                        story.append(Spacer(1, 0.2*inch))  # Üres sor
                        
                doc.build(story)
                return True
                
            success = await loop.run_in_executor(None, create_pdf, text)
            
            if not success:
                raise ConversionError("Failed to create PDF")
                
            return {"converted": True}
            
        except Exception as e:
            logger.error(f"Error converting TXT to PDF: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_odt_to_txt(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """ODT → TXT közvetlen konverzió a zipfile és BeautifulSoup csomagokkal"""
        try:
            import zipfile
            from bs4 import BeautifulSoup
            
            loop = asyncio.get_event_loop()
            
            def extract_odt_text():
                try:
                    with zipfile.ZipFile(input_path, 'r') as odt_zip:
                        with odt_zip.open('content.xml') as content_file:
                            content = content_file.read().decode('utf-8', errors="ignore")
                            soup = BeautifulSoup(content, "xml")
                            paragraphs = soup.find_all("text:p")
                            text = "\n".join(p.get_text() for p in paragraphs)
                            return text
                except zipfile.BadZipFile:
                    raise Exception("Invalid ODT file (bad zip structure)")
                except Exception as e:
                    raise Exception(f"ODT text extraction error: {str(e)}")
                    
            text = await loop.run_in_executor(None, extract_odt_text)
            
            # TXT fájl mentése
            async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                await f.write(text)
                
            return {"converted": True}
            
        except Exception as e:
            logger.error(f"Error converting ODT to TXT: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_pdf_to_txt(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """PDF → TXT közvetlen konverzió a PyMuPDF csomaggal"""
        try:
            import fitz  # PyMuPDF
            
            loop = asyncio.get_event_loop()
            
            def extract_pdf_text():
                try:
                    pdf = fitz.open(str(input_path))
                    text = ""
                    for page_num in range(len(pdf)):
                        page = pdf[page_num]
                        text += f"\n--- Oldal {page_num + 1} ---\n"
                        text += page.get_text()
                        text += "\n"
                    pdf.close()
                    return text
                except Exception as e:
                    raise Exception(f"PDF text extraction error: {str(e)}")
                    
            text = await loop.run_in_executor(None, extract_pdf_text)
            
            # TXT fájl mentése
            async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                await f.write(text)
                
            return {"converted": True}
            
        except Exception as e:
            logger.error(f"Error converting PDF to TXT: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_docx_to_txt(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """DOCX → TXT közvetlen konverzió a python-docx csomaggal"""
        try:
            from docx import Document
            
            loop = asyncio.get_event_loop()
            
            def extract_docx_text():
                try:
                    doc = Document(str(input_path))
                    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                    return "\n".join(paragraphs)
                except Exception as e:
                    raise Exception(f"DOCX text extraction error: {str(e)}")
                    
            text = await loop.run_in_executor(None, extract_docx_text)
            
            # TXT fájl mentése
            async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                await f.write(text)
                
            return {"converted": True}
            
        except Exception as e:
            logger.error(f"Error converting DOCX to TXT: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_rtf_to_txt(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """RTF → TXT közvetlen konverzió regex-alapú parsing-gal"""
        try:
            import re
            
            async with aiofiles.open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = await f.read()
            
            # Egyszerű RTF parsing
            # RTF vezérlő kódok eltávolítása
            text = re.sub(r'\\[a-z]+\d*\s?', '', rtf_content)  # RTF commands
            text = re.sub(r'[{}]', '', text)  # Braces
            text = re.sub(r'\\\*.*?;', '', text)  # \* commands
            text = re.sub(r'\\[\'"][0-9a-fA-F]{2}', '', text)  # Hex codes
            text = re.sub(r'\r\n|\r|\n', '\n', text)  # Normalize line breaks
            text = re.sub(r'\n+', '\n', text)  # Multiple newlines
            text = text.strip()
            
            # TXT fájl mentése
            async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                await f.write(text)
                
            return {"converted": True}
            
        except Exception as e:
            logger.error(f"Error converting RTF to TXT: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_epub_to_txt(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """EPUB → TXT közvetlen konverzió a ebooklib csomaggal"""
        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup
            
            loop = asyncio.get_event_loop()
            
            def extract_epub_text():
                try:
                    book = epub.read_epub(str(input_path))
                    text_content = []
                    
                    for item in book.get_items():
                        if item.get_type() == ebooklib.ITEM_DOCUMENT:
                            soup = BeautifulSoup(item.get_content(), 'html.parser')
                            text_content.append(soup.get_text())
                    
                    return "\n\n".join(text_content)
                except Exception as e:
                    raise Exception(f"EPUB text extraction error: {str(e)}")
                    
            text = await loop.run_in_executor(None, extract_epub_text)
            
            # TXT fájl mentése
            async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                await f.write(text)
                
            return {"converted": True}
            
        except Exception as e:
            logger.error(f"Error converting EPUB to TXT: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_mobi_to_txt(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """MOBI → TXT konverzió külső segédeszközzel"""
        try:
            # Használjuk az external_converter helper-t
            from external_converter import try_convert_external
            
            result = await try_convert_external("mobi", input_path, output_path)
            if result:
                return result
            else:
                raise HTTPException(
                    status_code=500,
                    detail="MOBI conversion failed - no suitable converter available"
                )
            
        except Exception as e:
            logger.error(f"Error converting MOBI to TXT: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"MOBI conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_txt_to_epub(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """TXT → EPUB közvetlen konverzió a ebooklib csomaggal"""
        try:
            import ebooklib
            from ebooklib import epub
            
            async with aiofiles.open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = await f.read()
            
            loop = asyncio.get_event_loop()
            
            def create_epub(text_content):
                try:
                    book = epub.EpubBook()
                    book.set_identifier('txt_conversion')
                    book.set_title('Converted from TXT')
                    book.set_language('en')
                    book.add_author('MosaicMaster Converter')
                    
                    # Szöveg felosztása fejezetekre (üres sorok alapján)
                    chapters = text_content.split('\n\n')
                    
                    for i, chapter_text in enumerate(chapters):
                        if chapter_text.strip():
                            c = epub.EpubHtml(title=f'Chapter {i+1}', 
                                            file_name=f'chapter_{i+1}.xhtml', 
                                            lang='en')
                            c.content = f'<html><body><p>{chapter_text.replace(chr(10), "</p><p>")}</p></body></html>'
                            book.add_item(c)
                    
                    # Tartalomjegyzék
                    book.toc = tuple(book.get_items_of_type(ebooklib.ITEM_DOCUMENT))
                    book.add_item(epub.EpubNcx())
                    book.add_item(epub.EpubNav())
                    
                    # Spine
                    book.spine = ['nav'] + list(book.get_items_of_type(ebooklib.ITEM_DOCUMENT))
                    
                    epub.write_epub(str(output_path), book)
                    return True
                except Exception as e:
                    raise Exception(f"EPUB creation error: {str(e)}")
                    
            success = await loop.run_in_executor(None, create_epub, text)
            
            if not success:
                raise ConversionError("Failed to create EPUB")
                
            return {"converted": True}
            
        except Exception as e:
            logger.error(f"Error converting TXT to EPUB: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_rtf_to_docx(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """RTF → DOCX közvetlen konverzió"""
        try:
            # Először RTF → TXT
            temp_txt = input_path.parent / f"temp_{input_path.stem}.txt"
            await ConversionProcessor._convert_rtf_to_txt(input_path, temp_txt)
            
            # Aztán TXT → DOCX
            result = await ConversionProcessor._convert_txt_to_docx(temp_txt, output_path)
            
            # Temp fájl törlése
            try:
                temp_txt.unlink()
            except:
                pass
                
            return result
            
        except Exception as e:
            logger.error(f"Error converting RTF to DOCX: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_docx_to_epub(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """DOCX → EPUB közvetlen konverzió"""
        try:
            # Először DOCX → TXT
            temp_txt = input_path.parent / f"temp_{input_path.stem}.txt"
            await ConversionProcessor._convert_docx_to_txt(input_path, temp_txt)
            
            # Aztán TXT → EPUB
            result = await ConversionProcessor._convert_txt_to_epub(temp_txt, output_path)
            
            # Temp fájl törlése
            try:
                temp_txt.unlink()
            except:
                pass
                
            return result
            
        except Exception as e:
            logger.error(f"Error converting DOCX to EPUB: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_epub_to_docx(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """EPUB → DOCX közvetlen konverzió"""
        try:
            # Először EPUB → TXT
            temp_txt = input_path.parent / f"temp_{input_path.stem}.txt"
            await ConversionProcessor._convert_epub_to_txt(input_path, temp_txt)
            
            # Aztán TXT → DOCX
            result = await ConversionProcessor._convert_txt_to_docx(temp_txt, output_path)
            
            # Temp fájl törlése
            try:
                temp_txt.unlink()
            except:
                pass
                
            return result
            
        except Exception as e:
            logger.error(f"Error converting EPUB to DOCX: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_pdf_to_docx(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """PDF → DOCX közvetlen konverzió a PyMuPDF és Python docx csomagokkal"""
        try:
            loop = asyncio.get_event_loop()
            
            def extract_and_create_docx():
                try:
                    pdf = fitz.open(str(input_path))
                    doc = Document()
                    
                    for page_num in range(len(pdf)):
                        page = pdf[page_num]
                        text = page.get_text()
                        
                        # Oldalcím hozzáadása
                        doc.add_heading(f"Oldal {page_num + 1}", level=1)
                        
                        # Bekezdések hozzáadása
                        paragraphs = text.split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                doc.add_paragraph(para)
                    
                    doc.save(str(output_path))
                    return True
                except Exception as e:
                    logger.error(f"PDF to DOCX conversion error: {str(e)}")
                    return False
            
            success = await loop.run_in_executor(None, extract_and_create_docx)
            
            if not success:
                raise ConversionError("Failed to convert PDF to DOCX")
                
            return {"converted": True}
            
        except Exception as e:
            logger.error(f"Error converting PDF to DOCX: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_pdf_to_doc(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """PDF → DOC közvetlen konverzió a PyMuPDF és Python docx csomagokkal"""
        try:
            loop = asyncio.get_event_loop()
            
            def extract_and_create_doc():
                try:
                    pdf = fitz.open(str(input_path))
                    doc = Document()
                    
                    for page_num in range(len(pdf)):
                        page = pdf[page_num]
                        text = page.get_text()
                        
                        # Oldalcím hozzáadása
                        doc.add_heading(f"Oldal {page_num + 1}", level=1)
                        
                        # Bekezdések hozzáadása
                        paragraphs = text.split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                doc.add_paragraph(para)
                    
                    doc.save(str(output_path))
                    return True
                except Exception as e:
                    logger.error(f"PDF to DOC conversion error: {str(e)}")
                    return False
            
            success = await loop.run_in_executor(None, extract_and_create_doc)
            
            if not success:
                raise ConversionError("Failed to convert PDF to DOC")
                
            return {"converted": True}
            
        except Exception as e:
            logger.error(f"Error converting PDF to DOC: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_odt_to_rtf(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """ODT → RTF konverzió LibreOffice használatával"""
        try:
            output_dir = output_path.parent
            process = await asyncio.create_subprocess_exec(
                "soffice",
                "--headless",
                "--convert-to", "rtf",
                str(input_path),
                "--outdir", str(output_dir),
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )
            stdout, stderr = await process.communicate()
            
            if process.returncode != 0:
                error_msg = stderr.decode() if stderr else "Unknown error"
                logger.error(f"LibreOffice conversion failed: {error_msg}")
                raise ExternalToolError(f"LibreOffice error: {error_msg}")

            converted_file = output_dir / (input_path.stem + ".rtf")
            if not converted_file.exists():
                raise ConversionError("Converted file not found")

            if converted_file != output_path:
                shutil.move(str(converted_file), str(output_path))

            return {"converted": True}
            
        except ExternalToolError as e:
            logger.error(f"External tool error: {str(e)}")
            raise HTTPException(status_code=500, detail=str(e))
        except Exception as e:
            logger.error(f"Error converting ODT to RTF: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_txt_to_odt(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """TXT → ODT konverzió"""
        if not HAS_ODF:
            raise HTTPException(
                status_code=500,
                detail="odfpy is not installed. Please install it via 'pip install odfpy'"
            )
            
        try:
            async with aiofiles.open(input_path, "r", encoding="utf-8") as f:
                content = await f.read()

            loop = asyncio.get_event_loop()
            
            def convert_txt_to_odt_sync():
                from odf.opendocument import OpenDocumentText
                from odf.text import P
                
                doc = OpenDocumentText()
                
                # Szöveg bekezdésekre bontása
                paragraphs = content.split('\n')
                
                for para_text in paragraphs:
                    if para_text.strip():
                        # Nem üres bekezdés hozzáadása
                        para = P(text=para_text.strip())
                        doc.text.addElement(para)
                    else:
                        # Üres sor hozzáadása
                        para = P(text="")
                        doc.text.addElement(para)

                doc.save(str(output_path))
                return True
                
            await loop.run_in_executor(None, convert_txt_to_odt_sync)
            return {"converted": True}

        except Exception as e:
            logger.error(f"Error converting TXT to ODT: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"TXT to ODT conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_pdf_to_epub_chunked(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """PDF → EPUB konverzió nagy fájlokhoz optimalizálva, darabolt feldolgozással - JAVÍTOTT"""
        try:
            loop = asyncio.get_event_loop()
            
            # A PDF megnyitását külön végrehajtjuk
            def open_pdf():
                try:
                    return fitz.open(str(input_path))
                except Exception as e:
                    raise FileFormatError(f"Failed to open PDF: {str(e)}")
            
            pdf = await loop.run_in_executor(None, open_pdf)
            total_pages = len(pdf)
            
            # EPUB létrehozása
            book = epub.EpubBook()
            book.set_title(input_path.stem)
            book.set_language("en")

            # CSS stílus a szebb megjelenítéshez
            default_css = epub.EpubItem(
                uid="style_default",
                file_name="style/default.css",
                media_type="text/css",
                content="body { font-family: serif; line-height: 1.5; margin: 1em; } "
                        "p { margin: 0.5em 0; text-indent: 1em; } "
                        "h1 { text-align: center; margin: 1em 0; } "
                        "img { max-width: 100%; height: auto; }"
            )
            book.add_item(default_css)

            # Laponként dolgozzuk fel, nem az egész PDF-et egyszerre
            chapters = []
            
            # Darabolt feldolgozási méret - egyszerre ennyi lapot dolgozunk fel
            CHUNK_SIZE = 10
            
            for chunk_start in range(0, total_pages, CHUNK_SIZE):
                chunk_end = min(chunk_start + CHUNK_SIZE, total_pages)
                logger.debug(f"Processing PDF chunk {chunk_start}-{chunk_end} of {total_pages}")
                
                # Lapok feldolgozása ebben a darabban
                for page_num in range(chunk_start, chunk_end):
                    # Oldal feldolgozása - JAVÍTOTT
                    def process_page(page_idx):
                        page = pdf[page_idx]
                        
                        # Szöveget kinyerjük strukturált formában
                        # Ezt a "text" módot használjuk a legjobban strukturált szöveghez
                        raw_text = page.get_text("text")
                        
                        # Szöveg darabolása bekezdésekre - JAVÍTOTT 
                        # Több sorköz = új bekezdés
                        paragraphs = []
                        current_para = []
                        
                        for line in raw_text.split('\n'):
                            if line.strip():
                                current_para.append(line.strip())
                            elif current_para:  # Üres sor és van tartalom
                                paragraphs.append(' '.join(current_para))
                                current_para = []
                        
                        # Az utolsó bekezdés hozzáadása, ha nem üres
                        if current_para:
                            paragraphs.append(' '.join(current_para))
                        
                        # HTML tartalom generálása bekezdésekből
                        html_content = f"<html><head><title>Page {page_idx + 1}</title></head><body>"
                        html_content += f"<h1>Page {page_idx + 1}</h1>"
                        
                        for para in paragraphs:
                            html_content += f"<p>{para}</p>"
                        
                        html_content += "</body></html>"
                        
                        return html_content
                    
                    html_content = await loop.run_in_executor(None, process_page, page_num)
                    
                    # Fejezet létrehozása
                    chapter = epub.EpubHtml(
                        title=f"Page {page_num + 1}",
                        file_name=f"page_{page_num + 1}.xhtml",
                        content=html_content
                    )
                    chapter.add_item(default_css)
                    book.add_item(chapter)
                    chapters.append(chapter)
                
                # Időnként yield-elünk az eseményhuroknak
                await asyncio.sleep(0)

            # EPUB struktúra összeállítása
            book.toc = [(epub.Section("Pages"), chapters)]
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            book.spine = ['nav'] + chapters

            # EPUB fájl mentése
            def write_epub():
                epub.write_epub(str(output_path), book, {})
                
            await loop.run_in_executor(None, write_epub)
            return {"pages": total_pages}

        except FileFormatError as e:
            logger.error(f"PDF file error: {str(e)}")
            raise HTTPException(status_code=400, detail=str(e))
        except Exception as e:
            logger.error(f"Error converting PDF to EPUB: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_srt_to_odt(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """SRT → ODT konverzió"""
        if not HAS_ODF:
            raise HTTPException(
                status_code=500,
                detail="odfpy is not installed. Please install it via 'pip install odfpy'"
            )
            
        try:
            async with aiofiles.open(input_path, "r", encoding="utf-8") as f:
                content = await f.read()

            loop = asyncio.get_event_loop()
            
            def convert_srt_to_odt_sync():
                from odf.opendocument import OpenDocumentText
                from odf.text import P
                
                blocks = re.split(r'\n\s*\n', content.strip())
                doc = OpenDocumentText()
                
                for block in blocks:
                    lines = block.splitlines()
                    if len(lines) >= 3:
                        subtitle_text = " ".join(lines[2:]).strip()
                    else:
                        subtitle_text = " ".join(lines).strip()
                    
                    if subtitle_text:
                        para = P(text=subtitle_text)
                        doc.text.addElement(para)

                doc.save(str(output_path))
                return True
                
            await loop.run_in_executor(None, convert_srt_to_odt_sync)
            return {"converted": True}

        except Exception as e:
            logger.error(f"Error converting SRT to ODT: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_srt_to_docx(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """SRT → DOCX konverzió"""
        try:
            loop = asyncio.get_event_loop()
            
            # Fájlolvasás aszinkron módon
            async with aiofiles.open(input_path, "r", encoding="utf-8", errors="ignore") as f:
                content = await f.read()
            
            def convert_srt_to_docx_sync(srt_content):
                doc = Document()
                blocks = re.split(r'\n\s*\n', srt_content.strip())
                for block in blocks:
                    lines = block.splitlines()
                    if len(lines) >= 3:
                        text = " ".join(lines[2:]).strip()
                    else:
                        text = " ".join(lines).strip()
                    
                    if text:
                        doc.add_paragraph(text)
                
                doc.save(str(output_path))
                return True

            await loop.run_in_executor(None, convert_srt_to_docx_sync, content)
            return {"converted": True}

        except Exception as e:
            logger.error(f"Error converting SRT to DOCX: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )
            
    @staticmethod
    async def _convert_to_srt(input_path: Path, output_path: Path, source_format: str) -> Dict[str, Any]:
        """Bármely támogatott dokumentum konvertálása SRT feliratozási formátumba"""
        try:
            loop = asyncio.get_event_loop()
            text = ""
            
            # Először kivonjuk a szöveget a forrásdokumentumból
            if source_format == "pdf":
                # PDF-ből kivonjuk a szöveget
                def extract_pdf_text():
                    try:
                        pdf = fitz.open(str(input_path))
                        texts = []
                        for page in pdf:
                            texts.append(page.get_text())
                        return "\n\n".join(texts)
                    except Exception as e:
                        logger.error(f"PDF feldolgozási hiba: {str(e)}")
                        raise
                
                text = await loop.run_in_executor(None, extract_pdf_text)
                
            elif source_format in ["docx", "doc"]:
                # DOCX/DOC fájlból kivonjuk a szöveget
                def extract_docx_text():
                    try:
                        doc = Document(input_path)
                        return "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())
                    except Exception as e:
                        logger.error(f"DOCX feldolgozási hiba: {str(e)}")
                        raise
                
                text = await loop.run_in_executor(None, extract_docx_text)
                
            elif source_format == "txt":
                # TXT fájlból egyszerűen beolvassuk a szöveget
                async with aiofiles.open(input_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = await f.read()
            
            elif source_format == "odt":
                # ODT fájlból kivonjuk a szöveget
                def extract_odt_text():
                    try:
                        with zipfile.ZipFile(input_path, 'r') as odt_zip:
                            with odt_zip.open('content.xml') as content_file:
                                content = content_file.read().decode('utf-8', errors="ignore")
                                soup = BeautifulSoup(content, "xml")
                                paragraphs = soup.find_all("text:p")
                                return "\n\n".join(p.get_text() for p in paragraphs)
                    except Exception as e:
                        logger.error(f"ODT feldolgozási hiba: {str(e)}")
                        raise
                
                text = await loop.run_in_executor(None, extract_odt_text)
                
            else:
                # Egyéb formátumoknál megpróbálunk közbeeső TXT konverziót használni
                temp_txt_path = input_path.with_suffix('.converted.txt')
                try:
                    # Közbenső TXT konverzió Calibre-vel
                    process = await asyncio.create_subprocess_exec(
                        "ebook-convert",
                        str(input_path),
                        str(temp_txt_path),
                        stdout=asyncio.subprocess.PIPE,
                        stderr=asyncio.subprocess.PIPE
                    )
                    stdout, stderr = await process.communicate()
                    
                    if process.returncode != 0:
                        raise Exception(f"Calibre konverziós hiba: {stderr.decode() if stderr else 'Ismeretlen hiba'}")
                    
                    # TXT olvasása
                    async with aiofiles.open(temp_txt_path, "r", encoding="utf-8", errors="ignore") as f:
                        text = await f.read()
                        
                except Exception as e:
                    logger.error(f"Közbülső konverziós hiba: {str(e)}")
                    raise HTTPException(
                        status_code=500,
                        detail=f"Nem sikerült a {source_format} formátumú dokumentumot SRT formátumra konvertálni: {str(e)}"
                    )
                finally:
                    # Ideiglenes fájl törlése
                    if temp_txt_path.exists():
                        try:
                            os.remove(temp_txt_path)
                        except:
                            pass
            
            # Most a szöveget SRT formátumra alakítjuk
            # Bekezdéseket felismerjük és feliratokká alakítjuk
            def create_srt_from_text(input_text):
                # Szöveget bekezdésekre osztjuk
                paragraphs = re.split(r'\n\s*\n', input_text.strip())
                
                # Rövid szakaszokra osztjuk a hosszú bekezdéseket (max ~40-50 karakter/sor)
                segments = []
                for para in paragraphs:
                    if not para.strip():
                        continue
                        
                    words = para.split()
                    current_line = []
                    current_length = 0
                    
                    for word in words:
                        # Ha az új szóval együtt a sor túl hosszú lenne, új sort kezdünk
                        if current_length + len(word) + 1 > 45:  # +1 a szóköz miatt
                            segments.append(" ".join(current_line))
                            current_line = [word]
                            current_length = len(word)
                        else:
                            current_line.append(word)
                            current_length += len(word) + 1
                    
                    # Az utolsó sort is hozzáadjuk
                    if current_line:
                        segments.append(" ".join(current_line))
                
                # SRT formátum összeállítása
                srt_content = []
                for i, segment in enumerate(segments, 1):
                    # SRT sorszáma
                    srt_content.append(str(i))
                    
                    # Időkód (egyszerű, 3 másodperces szakaszokkal)
                    start_time = (i - 1) * 3
                    end_time = i * 3
                    
                    start_formatted = f"{start_time//3600:02d}:{(start_time%3600)//60:02d}:{start_time%60:02d},000"
                    end_formatted = f"{end_time//3600:02d}:{(end_time%3600)//60:02d}:{end_time%60:02d},000"
                    
                    srt_content.append(f"{start_formatted} --> {end_formatted}")
                    
                    # A felirat szövege
                    srt_content.append(segment)
                    
                    # Üres sor a felirat blokkok között
                    srt_content.append("")
                
                return "\n".join(srt_content)
            
            # SRT tartalom létrehozása és mentése
            srt_content = await loop.run_in_executor(None, create_srt_from_text, text)
            
            async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
                await f.write(srt_content)
            
            return {"converted": True, "segments": len(re.findall(r'\d+\n\d{2}:\d{2}:\d{2},\d{3}', srt_content))}
            
        except Exception as e:
            logger.error(f"Error converting to SRT: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion to SRT failed: {str(e)}"
            )

    @staticmethod
    async def _convert_with_calibre(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """Konvertálás Calibre használatával"""
        try:
            # Folyamat változók inicializálása előre
            process = None
            
            try:
                process = await asyncio.create_subprocess_exec(
                    "ebook-convert",
                    str(input_path),
                    str(output_path),
                    stdout=asyncio.subprocess.PIPE,
                    stderr=asyncio.subprocess.PIPE
                )
                stdout, stderr = await process.communicate()
                
                if process.returncode != 0:
                    error_msg = stderr.decode() if stderr else "Unknown Calibre error"
                    logger.error(f"Calibre conversion failed: {error_msg}")
                    raise ExternalToolError(f"Calibre error: {error_msg}")
                
            except asyncio.CancelledError:
                # Hatékony folyamat megszakítás
                if process and process.returncode is None:
                    try:
                        process.terminate()
                        await asyncio.sleep(0.5)
                        if process.returncode is None:
                            process.kill()
                    except Exception as e:
                        logger.error(f"Error terminating Calibre process: {str(e)}")
                raise
                
            return {"converted": True}
            
        except ExternalToolError as e:
            raise HTTPException(status_code=500, detail=str(e))
        except Exception as e:
            logger.error(f"Error in Calibre conversion: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Calibre conversion failed: {str(e)}"
            )

    @staticmethod
    async def _convert_with_libreoffice(input_path: Path, output_path: Path) -> Dict[str, Any]:
        """Konvertálás LibreOffice használatával"""
        try:
            # LibreOffice elérhetőségének ellenőrzése
            check_libreoffice()
            
            output_dir = output_path.parent
            # Folyamat változók inicializálása előre
            process = None
            
            try:
                target_ext = output_path.suffix[1:]  # A kezdő pont nélkül
                
                process = await asyncio.create_subprocess_exec(
                    "soffice",
                    "--headless",
                    "--convert-to", target_ext,
                    str(input_path),
                    "--outdir", str(output_dir),
                    stdout=asyncio.subprocess.PIPE,
                    stderr=asyncio.subprocess.PIPE
                )
                stdout, stderr = await process.communicate()
                
                if process.returncode != 0:
                    error_msg = stderr.decode() if stderr else "Unknown LibreOffice error"
                    logger.error(f"LibreOffice conversion failed: {error_msg}")
                    raise ExternalToolError(f"LibreOffice error: {error_msg}")
            
            except asyncio.CancelledError:
                # Hatékony folyamat megszakítás
                if process and process.returncode is None:
                    try:
                        process.terminate()
                        await asyncio.sleep(0.5)
                        if process.returncode is None:
                            process.kill()
                    except Exception as e:
                        logger.error(f"Error terminating LibreOffice process: {str(e)}")
                raise

            converted_file = output_dir / (input_path.stem + output_path.suffix)
            if not converted_file.exists():
                raise ConversionError("Converted file not found after LibreOffice conversion")

            if converted_file != output_path:
                shutil.move(str(converted_file), str(output_path))

            return {"converted": True}

        except ExternalToolError as e:
            raise HTTPException(status_code=500, detail=str(e))
        except ConversionError as e:
            raise HTTPException(status_code=500, detail=str(e))
        except Exception as e:
            logger.error(f"Error converting with LibreOffice: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Conversion failed: {str(e)}"
            )

def get_temp_manager():
    """TempDirectoryManager függőségi injektáláshoz"""
    return temp_manager

async def check_file_size(file: UploadFile) -> None:
    """Fájlméret ellenőrzése"""
    # Fájl méret ellenőrzése a feldolgozás előtt
    file_size = 0
    chunk_size = 1024 * 1024  # 1MB
    chunk = await file.read(chunk_size)
    while chunk:
        file_size += len(chunk)
        if file_size > MAX_FILE_SIZE:
            await file.seek(0)  # Visszatekerés
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum allowed size is {MAX_FILE_SIZE/(1024*1024):.1f}MB"
            )
        chunk = await file.read(chunk_size)
    
    # Fájl mutatót visszaállítjuk az elejére
    await file.seek(0)

@router.post("/")
async def convert_document(
    file: UploadFile = File(...),
    target_format: str = Form(...),
    connection_id: str = Form(None),
    background_tasks: BackgroundTasks = None,
    temp_mgr: TempDirectoryManager = Depends(get_temp_manager)
):
    """Dokumentum konvertálása"""
    if not connection_id:
        connection_id = str(uuid.uuid4())
    
    # Fájlméret ellenőrzése
    try:
        await check_file_size(file)
    except HTTPException as e:
        logger.warning(f"File size check failed: {e.detail}")
        raise
        
    # Aszinkron kontextuskezelővel kezeljük az ideiglenes könyvtárat
    async with temp_mgr.temp_dir(connection_id) as work_dir:
        try:
            # Fájl mentése
            input_path = work_dir / sanitize_filename(file.filename)
            async with aiofiles.open(input_path, "wb") as f:
                await f.write(await file.read())

            # Progress update - fájl feldolgozás kezdete
            await manager.send_progress(connection_id, 10, "Processing document...")

            # Forrás formátum meghatározása
            ext = Path(file.filename).suffix.lower()[1:]
            
            # JAVÍTÁS: Kiterjesztés vs MIME-type kezelés
            if file.content_type and (
                file.content_type.startswith("application/") or 
                file.content_type.startswith("text/") or
                file.content_type.startswith("image/")
            ):
                source_format = file.content_type.split('/')[-1]
                
                # Normalize common MIME types
                mime_to_ext = {
                    "vnd.openxmlformats-officedocument.wordprocessingml.document": "docx", 
                    "msword": "doc",
                    "vnd.oasis.opendocument.text": "odt",
                    "plain": "txt",
                    "rtf": "rtf",
                    "vnd.ms-powerpoint": "ppt",
                    "vnd.openxmlformats-officedocument.presentationml.presentation": "pptx",
                    "x-mobipocket-ebook": "mobi",
                    "epub+zip": "epub"
                }
                
                source_format = mime_to_ext.get(source_format, source_format)
            else:
                # Ha nincs megfelelő MIME-type, használjuk a kiterjesztést
                source_format = ext
                
                # txt fájl speciális kezelése - ha nem ismerhető fel más módon
                if ext == "txt" or (file.filename and file.filename.lower().endswith(".txt")):
                    source_format = "txt"
            
            # Kép formátumok normalizálása
            if ext in ["jpg", "jpeg", "mpo"]:
                source_format = "image/jpeg"
            elif ext in ["png"]:
                source_format = "image/png"
            elif ext in ["gif"]:
                source_format = "image/gif"
            
            logger.info(f"Source format determined: {source_format} for file {file.filename}")

            output_filename = f"converted_{Path(file.filename).stem}.{target_format}"
            output_path = SYSTEM_DOWNLOADS / output_filename

            # Progress update - konverzió indítása
            await manager.send_progress(connection_id, 30, "Starting conversion...")

            # Dokumentum konvertálása
            try:
                result = await ConversionProcessor.convert_document(
                    input_path, output_path, source_format, target_format.lower()
                )
            except HTTPException as e:
                # Pontos hibainformáció átadása a kliensnek
                await manager.send_progress(connection_id, 100, f"Error: {e.detail}")
                raise
            
            # Progress update - konverzió befejezve
            await manager.send_progress(connection_id, 100, "Conversion complete")
            
            # Ha van háttérfolyamat-kezelő, hozzáadjuk a tisztítási feladatot
            if background_tasks:
                # Ezt még nem töröljük azonnal, hanem később a háttérben
                background_tasks.add_task(temp_mgr.cleanup, connection_id)
            
            return JSONResponse({
                "download_url": f"/download/{output_filename}",
                "metadata": result
            })
        except Exception as e:
            logger.error(f"Error in /api/convert: {e}")
            
            # Folyamat állapot frissítése
            try:
                await manager.send_progress(connection_id, 100, f"Error: {str(e)}")
            except Exception as ws_error:
                logger.error(f"Failed to send error via WebSocket: {ws_error}")
                
            # Ha specifikus HTTP kivétel, azt továbbítjuk
            if isinstance(e, HTTPException):
                raise
                
            # Egyéb kivételek esetén általános HTTP hibát adunk
            raise HTTPException(status_code=500, detail=str(e))

@router.post("/ocr")
async def ocr_image(
    file: UploadFile = File(...),
    background_tasks: BackgroundTasks = None,
    temp_mgr: TempDirectoryManager = Depends(get_temp_manager)
):
    """Kép OCR feldolgozása"""
    connection_id = str(uuid.uuid4())
    
    # Fájlméret ellenőrzése
    try:
        await check_file_size(file)
    except HTTPException as e:
        logger.warning(f"File size check failed in OCR: {e.detail}")
        raise
    
    # Aszinkron kontextuskezelővel kezeljük az ideiglenes könyvtárat
    async with temp_mgr.temp_dir(connection_id) as work_dir:
        try:
            input_path = work_dir / sanitize_filename(file.filename)
            async with aiofiles.open(input_path, "wb") as f:
                await f.write(await file.read())

            # Progress update
            await manager.send_progress(connection_id, 30, "Processing image with OCR...")
            
            # OCR feldolgozás
            try:
                text, _ = await process_image_with_ocr(input_path)
            except HTTPException as e:
                # Pontos hibainformáció átadása a kliensnek
                await manager.send_progress(connection_id, 100, f"OCR Error: {e.detail}")
                raise
            
            # Progress update
            await manager.send_progress(connection_id, 100, "OCR processing complete")
            
            # Ha van háttérfolyamat-kezelő, hozzáadjuk a tisztítási feladatot
            if background_tasks:
                background_tasks.add_task(temp_mgr.cleanup, connection_id)
            
            return JSONResponse({"ocr_text": text})
        except Exception as e:
            logger.error(f"OCR failed: {str(e)}")
            
            # Ha specifikus HTTP kivétel, azt továbbítjuk
            if isinstance(e, HTTPException):
                raise
                
            # Egyéb kivételek esetén általános HTTP hibát adunk
            raise HTTPException(status_code=500, detail=f"OCR failed: {str(e)}")

@router.post("/image_to_pdf")
async def convert_image_to_pdf(
    file: UploadFile = File(...),
    connection_id: Optional[str] = Form(None),
    background_tasks: BackgroundTasks = None,
    temp_mgr: TempDirectoryManager = Depends(get_temp_manager)
):
    """Kép konvertálása PDF formátumba - MPO támogatással"""
    if not connection_id:
        connection_id = str(uuid.uuid4())
    
    # Aszinkron kontextuskezelővel kezeljük az ideiglenes könyvtárat
    async with temp_mgr.temp_dir(connection_id) as work_dir:
        try:
            input_path = work_dir / sanitize_filename(file.filename)
            async with aiofiles.open(input_path, "wb") as f:
                await f.write(await file.read())
            
            await manager.send_progress(connection_id, 10, "Processing image file...")
            
            # Ellenőrizzük, hogy képfájl-e
            ext = Path(file.filename).suffix.lower()[1:]
            if ext not in ["jpg", "jpeg", "png", "gif", "bmp", "tiff", "mpo"]:
                raise HTTPException(status_code=400, detail="Unsupported file format. Must be an image file.")
            
            # Kimenet elkészítése
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"image_to_pdf_{timestamp}.pdf"
            output_path = SYSTEM_DOWNLOADS / output_filename
            
            # Kép PDF-fé konvertálása
            try:
                # PIL-t használjuk a konverzióhoz
                from PIL import Image
                
                loop = asyncio.get_event_loop()
                
                def convert_image_to_pdf():
                    try:
                        img = Image.open(input_path)
                        
                        # MPO és más speciális formátumok kezelése
                        if hasattr(img, 'format') and img.format and img.format.upper() not in ['JPEG', 'JPG', 'PNG', 'GIF', 'BMP', 'TIFF']:
                            logger.info(f"Converting image from {img.format} format to JPEG format for PDF generation")
                            
                            # RGB-re konvertálás, ha szükséges
                            if img.mode != 'RGB':
                                img = img.convert('RGB')
                            
                            # Ideiglenes fájl létrehozása a konvertált képnek
                            temp_jpg_path = input_path.with_suffix('.converted.jpg')
                            img.save(temp_jpg_path, format='JPEG', quality=95)
                            
                            # Újra megnyitjuk a konvertált képet
                            img = Image.open(temp_jpg_path)
                        elif img.mode == 'RGBA':
                            # RGB-re konvertálás, ha szükséges
                            img = img.convert('RGB')
                        
                        # A4 méret pixel-ben (72 DPI-vel)
                        a4_size_px = (int(8.27 * 72), int(11.69 * 72))
                        
                        # Kép méretezése, ha nagyobb, mint az A4
                        img_width, img_height = img.size
                        a4_width, a4_height = a4_size_px
                        
                        # Méretarány számítás
                        width_ratio = a4_width / img_width
                        height_ratio = a4_height / img_height
                        ratio = min(width_ratio, height_ratio)
                        
                        # Csak akkor méretezzük, ha a kép nagyobb mint A4
                        if ratio < 1:
                            new_width = int(img_width * ratio)
                            new_height = int(img_height * ratio)
                            img = img.resize((new_width, new_height), Image.LANCZOS)
                        
                        # PDF mentése
                        img.save(output_path, "PDF", resolution=72.0)
                        return True
                    except Exception as e:
                        logger.error(f"Image to PDF conversion error: {str(e)}")
                        return False
                
                success = await loop.run_in_executor(None, convert_image_to_pdf)
                
                if not success:
                    raise HTTPException(status_code=500, detail="Failed to convert image to PDF")
                
                await manager.send_progress(connection_id, 100, "Image converted to PDF successfully")
                
                return JSONResponse({
                    "download_url": f"/download/{output_filename}"
                })
                
            except ImportError:
                # Ha nincs PIL, próbáljunk meg más megoldást
                try:
                    # ImageMagick használata
                    cmd = [
                        "convert",
                        str(input_path),
                        str(output_path)
                    ]
                    
                    process = await asyncio.create_subprocess_exec(
                        *cmd,
                        stdout=asyncio.subprocess.PIPE,
                        stderr=asyncio.subprocess.PIPE
                    )
                    
                    stdout, stderr = await process.communicate()
                    
                    if process.returncode != 0:
                        logger.error(f"ImageMagick conversion failed: {stderr.decode()}")
                        raise HTTPException(status_code=500, detail="Image to PDF conversion failed")
                    
                    await manager.send_progress(connection_id, 100, "Image converted to PDF successfully")
                    
                    return JSONResponse({
                        "download_url": f"/download/{output_filename}"
                    })
                except:
                    # Ha az ImageMagick sem elérhető, próbáljuk meg reportlab-bal
                    from reportlab.lib.pagesizes import letter
                    from reportlab.lib.utils import ImageReader
                    from reportlab.pdfgen import canvas
                    
                    def convert_with_reportlab():
                        try:
                            img = ImageReader(str(input_path))
                            img_width, img_height = img.getSize()
                            aspect = img_height / float(img_width)
                            
                            # A4 méret számítása
                            a4_width, a4_height = letter
                            
                            # Megfelelő méret számítása
                            if img_width > a4_width or img_height > a4_height:
                                if aspect > 1:  # Magasabb, mint széles
                                    img_height = a4_height
                                    img_width = img_height / aspect
                                else:  # Szélesebb, mint magas
                                    img_width = a4_width
                                    img_height = img_width * aspect
                            
                            c = canvas.Canvas(str(output_path), pagesize=letter)
                            # Kép elhelyezése középre
                            x = (a4_width - img_width) / 2
                            y = (a4_height - img_height) / 2
                            c.drawImage(img, x, y, width=img_width, height=img_height)
                            c.save()
                            return True
                        except Exception as e:
                            logger.error(f"ReportLab conversion error: {str(e)}")
                            return False
                    
                    success = await loop.run_in_executor(None, convert_with_reportlab)
                    
                    if not success:
                        raise HTTPException(status_code=500, detail="Failed to convert image to PDF with ReportLab")
                    
                    await manager.send_progress(connection_id, 100, "Image converted to PDF successfully")
                    
                    return JSONResponse({
                        "download_url": f"/download/{output_filename}"
                    })
                
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Image to PDF conversion error: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Image to PDF conversion failed: {str(e)}")
        
        finally:
            # Ha van háttérfolyamat, adjuk hozzá a munkakönyvtár tisztítását
            if background_tasks:
                background_tasks.add_task(temp_mgr.cleanup, connection_id)
            else:
                # Egyébként töröljük most azonnal
                await temp_mgr.cleanup(connection_id)

# Batch feldolgozás implementálása nagy gyűjteményekhez
@router.post("/batch")
async def batch_convert_documents(
    files: List[UploadFile] = File(...),
    target_format: str = Form(...),
    connection_id: str = Form(None),
    background_tasks: BackgroundTasks = None,
    temp_mgr: TempDirectoryManager = Depends(get_temp_manager)
):
    """Több dokumentum konvertálása egyetlen műveletben"""
    if not connection_id:
        connection_id = str(uuid.uuid4())
    
    if not files:
        raise HTTPException(
            status_code=400,
            detail="No files provided for batch conversion"
        )
    
    # Aszinkron kontextuskezelővel kezeljük az ideiglenes könyvtárat
    async with temp_mgr.temp_dir(connection_id) as work_dir:
        try:
            result_files = []
            total_files = len(files)
            
            for idx, file in enumerate(files, 1):
                # Fájlméret ellenőrzése
                try:
                    await check_file_size(file)
                except HTTPException as e:
                    logger.warning(f"File size check failed for {file.filename}: {e.detail}")
                    result_files.append({
                        "filename": file.filename,
                        "status": "error",
                        "error": e.detail
                    })
                    continue
                
                # Fájl mentése
                input_path = work_dir / sanitize_filename(file.filename)
                async with aiofiles.open(input_path, "wb") as f:
                    await f.write(await file.read())

                # Progress update
                progress = int(30 + (idx / total_files) * 60)
                await manager.send_progress(
                    connection_id, 
                    progress, 
                    f"Processing file {idx}/{total_files}: {file.filename}"
                )

                # Forrás formátum meghatározása - JAVÍTVA
                ext = Path(file.filename).suffix.lower()[1:]
                
                # MIME-type és kiterjesztés kezelés
                if file.content_type and (
                    file.content_type.startswith("application/") or 
                    file.content_type.startswith("text/") or
                    file.content_type.startswith("image/")
                ):
                    source_format = file.content_type.split('/')[-1]
                    
                    # Normalize common MIME types
                    mime_to_ext = {
                        "vnd.openxmlformats-officedocument.wordprocessingml.document": "docx", 
                        "msword": "doc",
                        "vnd.oasis.opendocument.text": "odt",
                        "plain": "txt",
                        "rtf": "rtf",
                        "vnd.ms-powerpoint": "ppt",
                        "vnd.openxmlformats-officedocument.presentationml.presentation": "pptx",
                        "x-mobipocket-ebook": "mobi",
                        "epub+zip": "epub"
                    }
                    
                    source_format = mime_to_ext.get(source_format, source_format)
                else:
                    # Ha nincs megfelelő MIME-type, használjuk a kiterjesztést
                    source_format = ext
                    
                    # txt fájl speciális kezelése - ha nem ismerhető fel más módon
                    if ext == "txt" or (file.filename and file.filename.lower().endswith(".txt")):
                        source_format = "txt"
                
                # Kép formátumok normalizálása
                if ext in ["jpg", "jpeg", "mpo"]:
                    source_format = "image/jpeg"
                elif ext in ["png"]:
                    source_format = "image/png"
                elif ext in ["gif"]:
                    source_format = "image/gif"
                
                logger.info(f"Source format determined: {source_format} for file {file.filename}")

                output_filename = f"converted_{Path(file.filename).stem}.{target_format}"
                output_path = SYSTEM_DOWNLOADS / output_filename

                # Dokumentum konvertálása
                try:
                    conversion_result = await ConversionProcessor.convert_document(
                        input_path, output_path, source_format, target_format.lower()
                    )
                    
                    result_files.append({
                        "filename": file.filename,
                        "output_filename": output_filename,
                        "download_url": f"/download/{output_filename}",
                        "status": "success",
                        "metadata": conversion_result
                    })
                except Exception as e:
                    logger.error(f"Error converting {file.filename}: {str(e)}")
                    error_message = str(e)
                    if isinstance(e, HTTPException):
                        error_message = e.detail
                        
                    result_files.append({
                        "filename": file.filename,
                        "status": "error",
                        "error": error_message
                    })

            # Progress update - batch folyamat befejezve
            await manager.send_progress(connection_id, 100, "Batch conversion complete")
            
            # Ha van háttérfolyamat-kezelő, hozzáadjuk a tisztítási feladatot
            if background_tasks:
                background_tasks.add_task(temp_mgr.cleanup, connection_id)
            
            # Statisztika a folyamatról
            success_count = sum(1 for file in result_files if file["status"] == "success")
            
            return JSONResponse({
                "batch_results": result_files,
                "summary": {
                    "total": total_files,
                    "success": success_count,
                    "failed": total_files - success_count
                }
            })
        except Exception as e:
            logger.error(f"Error in batch conversion: {str(e)}")
            raise HTTPException(status_code=500, detail=str(e))


# Standalone functions for text_reader_service integration
def convert_document_to_text(file_path: str, file_type: str = None) -> str:
    """
    Synchronous wrapper for document text extraction
    Used by text_reader_service.py
    Converts all document types to TXT format to avoid async event loop issues
    """
    import zipfile
    from pathlib import Path
    
    file_path = Path(file_path)
    
    # Determine file type from extension if not provided
    if not file_type:
        file_type = file_path.suffix.lower()[1:]  # Remove the dot
    
    # Convert all document types to TXT using synchronous methods
    try:
        if file_type.lower() == 'txt':
            return _extract_txt_text_sync(file_path)
        elif file_type.lower() == 'docx':
            return _extract_docx_text_sync(file_path)
        elif file_type.lower() == 'pdf':
            return _extract_pdf_text_sync(file_path)
        elif file_type.lower() == 'rtf':
            return _extract_rtf_text_sync(file_path)
        elif file_type.lower() == 'odt':
            return _extract_odt_text_sync(file_path)
        elif file_type.lower() == 'doc':
            return _extract_doc_text_sync(file_path)
        elif file_type.lower() in ['jpg', 'jpeg', 'png', 'gif']:
            return _extract_image_text_sync(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    except Exception as e:
        raise ValueError(f"Document conversion failed for {file_type}: {str(e)}")


def process_ocr(file_path: str, language: str = 'hun+eng') -> str:
    """
    Synchronous wrapper for OCR processing
    Used by text_reader_service.py
    """
    import asyncio
    from pathlib import Path
    
    file_path = Path(file_path)
    
    # Run async function synchronously  
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    try:
        result = loop.run_until_complete(process_image_with_ocr(file_path))
        return result[0]  # Return just the text part
    finally:
        loop.close()


# Synchronous text extraction functions for text_reader_service
def _extract_txt_text_sync(file_path: Path) -> str:
    """Extract text from TXT file synchronously"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()


def _extract_docx_text_sync(file_path: Path) -> str:
    """Extract text from DOCX file synchronously"""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = []
        for paragraph in doc.paragraphs:
            paragraphs.append(paragraph.text)
        return '\n'.join(paragraphs)
    except Exception as e:
        raise Exception(f"DOCX text extraction failed: {str(e)}")


def _extract_pdf_text_sync(file_path: Path) -> str:
    """Extract text from PDF file synchronously"""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        raise Exception(f"PDF text extraction failed: {str(e)}")


def _extract_rtf_text_sync(file_path: Path) -> str:
    """Extract text from RTF file synchronously"""
    try:
        from striprtf.striprtf import rtf_to_text
        with open(file_path, 'r', encoding='utf-8') as f:
            rtf_content = f.read()
        return rtf_to_text(rtf_content)
    except ImportError:
        # Fallback: try to read as plain text (basic RTF handling)
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            # Basic RTF text extraction - remove RTF codes
            import re
            text = re.sub(r'\\[a-z]+\d*\s?', '', content)
            text = re.sub(r'[{}]', '', text)
            return text.strip()
        except Exception:
            raise Exception("RTF text extraction failed - striprtf package not available")
    except Exception as e:
        raise Exception(f"RTF text extraction failed: {str(e)}")


def _extract_odt_text_sync(file_path: Path) -> str:
    """Extract text from ODT file synchronously"""
    try:
        import zipfile
        from bs4 import BeautifulSoup
        
        with zipfile.ZipFile(file_path, 'r') as odt_zip:
            with odt_zip.open('content.xml') as content_file:
                content = content_file.read().decode('utf-8', errors="ignore")
                soup = BeautifulSoup(content, "xml")
                paragraphs = soup.find_all("text:p")
                return "\n".join(p.get_text() for p in paragraphs)
    except zipfile.BadZipFile:
        raise Exception("Invalid ODT file (bad zip structure)")
    except Exception as e:
        raise Exception(f"ODT text extraction failed: {str(e)}")


def _extract_doc_text_sync(file_path: Path) -> str:
    """Extract text from DOC file synchronously using LibreOffice"""
    try:
        import subprocess
        import tempfile
        import os
        
        # Create temporary directory for conversion
        with tempfile.TemporaryDirectory() as temp_dir:
            # Convert DOC to TXT using LibreOffice
            result = subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'txt',
                '--outdir', temp_dir, str(file_path)
            ], capture_output=True, text=True, timeout=30)
            
            if result.returncode != 0:
                raise Exception(f"LibreOffice conversion failed: {result.stderr}")
            
            # Read the converted TXT file
            txt_file = os.path.join(temp_dir, file_path.stem + '.txt')
            if os.path.exists(txt_file):
                with open(txt_file, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                raise Exception("Converted TXT file not found")
                
    except subprocess.TimeoutExpired:
        raise Exception("DOC conversion timed out")
    except FileNotFoundError:
        raise Exception("LibreOffice not found - please install LibreOffice")
    except Exception as e:
        raise Exception(f"DOC text extraction failed: {str(e)}")


def _extract_image_text_sync(file_path: Path) -> str:
    """Extract text from image file using OCR synchronously"""
    try:
        import pytesseract
        from PIL import Image
        image = Image.open(file_path)
        return pytesseract.image_to_string(image, lang='hun+eng')
    except Exception as e:
        raise Exception(f"OCR text extraction failed: {str(e)}")


class DocumentProcessor:
    """Additional helper methods for document processing"""
    
    @staticmethod
    async def _convert_subtitle_to_txt(input_path: Path, output_path: Path, source_format: str) -> Dict[str, Any]:
        """Felirat fájlok szövegének kinyerése (SRT, SUB, VTT)"""
        try:
            async with aiofiles.open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = await f.read()
            
            extracted_text = []
            
            if source_format == "srt":
                # SRT formátum: időkódok és számozás eltávolítása
                import re
                # SRT pattern: szám, időkód, szöveg, üres sor
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    # Kihagyjuk a számokat és időkódokat
                    if line and not line.isdigit() and '-->' not in line:
                        extracted_text.append(line)
            
            elif source_format == "vtt":
                # WebVTT formátum
                import re
                lines = content.split('\n')
                in_cue = False
                for line in lines:
                    line = line.strip()
                    if line.startswith('WEBVTT') or line.startswith('NOTE'):
                        continue
                    if '-->' in line:
                        in_cue = True
                        continue
                    if line == '':
                        in_cue = False
                        continue
                    if in_cue and line:
                        # HTML tagek eltávolítása
                        clean_line = re.sub(r'<[^>]+>', '', line)
                        extracted_text.append(clean_line)
            
            elif source_format == "sub":
                # SUB formátum (egyszerű)
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    if line and not line.startswith('{') and not line.startswith('['):
                        extracted_text.append(line)
            
            # Szöveg összeállítása
            final_text = '\n'.join(extracted_text)
            
            # Szöveg mentése
            async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                await f.write(final_text)
            
            return {
                "converted": True,
                "lines_extracted": len(extracted_text),
                "characters": len(final_text)
            }
            
        except Exception as e:
            logger.error(f"Error converting {source_format} to TXT: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Subtitle conversion failed: {str(e)}"
            )
    
    @staticmethod
    async def _convert_subtitle_format(input_path: Path, output_path: Path, source_format: str, target_format: str) -> Dict[str, Any]:
        """Felirat formátumok közötti konverzió"""
        try:
            async with aiofiles.open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = await f.read()
            
            # Először tiszta szöveget nyerünk ki
            temp_txt = input_path.parent / f"{input_path.stem}_temp.txt"
            await DocumentProcessor._convert_subtitle_to_txt(input_path, temp_txt, source_format)
            
            # Majd a célformátumba konvertáljuk
            result = await DocumentProcessorService._convert_to_srt(temp_txt, output_path, "txt")
            
            # Takarítás
            if temp_txt.exists():
                temp_txt.unlink()
            
            return result
            
        except Exception as e:
            logger.error(f"Error converting {source_format} to {target_format}: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Subtitle format conversion failed: {str(e)}"
            )
