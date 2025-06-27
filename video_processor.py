import asyncio
import uuid
import shutil
import json
import logging
import os
import re
import subprocess
import random
import time
from datetime import datetime, timedelta
from pathlib import Path
from functools import partial
from typing import Dict, Optional, List, Tuple, Any
from concurrent.futures import ThreadPoolExecutor
from urllib.parse import quote, urlparse, parse_qs

import aiofiles
import aiohttp
from fastapi import APIRouter, HTTPException, UploadFile, File, Form, WebSocket, WebSocketDisconnect, BackgroundTasks
from fastapi.responses import JSONResponse, FileResponse
from pydub import AudioSegment
from pydub.utils import make_chunks
from pydub.exceptions import CouldntDecodeError

from config import (
    TEMP_DIR,
    SYSTEM_DOWNLOADS,
    manager,
    logger,
    client,
    sanitize_filename,
    process_semaphore,
    chunk_text_by_tokens
)

# VideoDownloader importálása - pontos elérési út a fájlrendszerben
from videodownloader import VideoDownloader

from transcriber import _transcribe_audio  # Importáljuk a transzkripciós függvényt

logger = logging.getLogger(__name__)
router = APIRouter()

async def async_rmtree(path: Path):
    """Könyvtár aszinkron törlése"""
    if not path.exists():
        return
    loop = asyncio.get_event_loop()
    await loop.run_in_executor(None, partial(shutil.rmtree, path, ignore_errors=True))

def sanitize_youtube_filename(filename):
    """Speciális karakterek eltávolítása a YouTube fájlnevekből"""
    # A # karakterek és más problémás karakterek lecserélése
    sanitized = re.sub(r'[#?&]', '_', str(filename))
    # Szóközök cseréje aláhúzásra
    sanitized = sanitized.replace(" ", "_")
    # Maximális fájlnév hosszúság (255 karakter)
    if len(sanitized) > 255:
        base, ext = os.path.splitext(sanitized)
        sanitized = f"{base[:245]}{ext}"
    return sanitized

async def convert_to_wav(input_path: Path, output_path: Path) -> bool:
    """Audio fájl konvertálása WAV formátumba"""
    try:
        process = await asyncio.create_subprocess_exec(
            "ffmpeg",
            "-i", str(input_path),
            "-vn",
            "-acodec", "pcm_s16le",
            "-ar", "44100",
            "-ac", "2",
            str(output_path),
            "-y",
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await process.communicate()
        if process.returncode != 0:
            logger.error(f"FFmpeg conversion failed: {stderr.decode()}")
            return False
        return True
    except Exception as e:
        logger.error(f"Error converting to WAV with FFmpeg: {e}")
        return False

async def convert_to_mp3(input_path: Path, output_path: Path) -> bool:
    """Audio fájl konvertálása MP3 formátumba"""
    try:
        process = await asyncio.create_subprocess_exec(
            "ffmpeg",
            "-i", str(input_path),
            "-vn",
            "-acodec", "libmp3lame",
            "-q:a", "2",
            str(output_path),
            "-y",
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await process.communicate()
        if process.returncode != 0:
            logger.error(f"FFmpeg MP3 conversion failed: {stderr.decode()}")
            return False
        return True
    except Exception as e:
        logger.error(f"Error converting to MP3 with FFmpeg: {e}")
        return False

async def generate_srt(segments: List[dict]) -> str:
    """SRT felirat generálása a szegmensekből"""
    srt_output = ""
    for i, segment in enumerate(segments, 1):
        start_time = format_timestamp(segment['start'])
        end_time = format_timestamp(segment['end'])
        srt_output += f"{i}\n{start_time} --> {end_time}\n{segment['text']}\n\n"
    return srt_output

async def generate_vtt(segments: List[dict]) -> str:
    """VTT felirat generálása a szegmensekből"""
    vtt_output = "WEBVTT\n\n"
    for segment in segments:
        start_time = "0" + str(timedelta(seconds=segment['start']))
        end_time = "0" + str(timedelta(seconds=segment['end']))
        vtt_output += f"{start_time} --> {end_time}\n{segment['text']}\n\n"
    return vtt_output

def format_timestamp(seconds):
    """Convert seconds to SRT timestamp format (HH:MM:SS,mmm)"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = seconds % 60
    milliseconds = int((secs % 1) * 1000)
    secs = int(secs)

    return f"{hours:02d}:{minutes:02d}:{secs:02d},{milliseconds:03d}"

async def _convert_with_ffmpeg(ffmpeg_cmd: list, converted_path: Path, connection_id: str):
    """FFmpeg konvertálást végző aszinkron függvény."""
    process = await asyncio.create_subprocess_exec(*ffmpeg_cmd,
                                                    stdout=asyncio.subprocess.PIPE,
                                                    stderr=asyncio.subprocess.PIPE)
    stdout, stderr = await process.communicate()
    if process.returncode != 0:
        logger.error(stderr.decode())
        raise HTTPException(status_code=500, detail=f"Video conversion failed: {stderr.decode()}")

async def generate_transcript_docx(text: str, segments: List[dict], output_path: Path):
    """Transcript generálása DOCX formátumba"""
    from docx import Document
    
    doc = Document()
    if segments:
        for segment in segments:
            doc.add_paragraph(f"[{segment['start']:.2f}-{segment['end']:.2f}] {segment['text']}")
    else:
        doc.add_paragraph(text)
    doc.save(str(output_path))

async def generate_transcript_pdf(text: str, segments: List[dict], output_path: Path):
    """Transcript generálása PDF formátumba"""
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    
    doc = SimpleDocTemplate(str(output_path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    if segments:
        for segment in segments:
            story.append(Paragraph(f"[{segment['start']:.2f}-{segment['end']:.2f}] {segment['text']}", styles['Normal']))
            story.append(Spacer(1, 12))
    else:
        story.append(Paragraph(text, styles['Normal']))
    doc.build(story)

async def translate_chunks(chunks: List[str], target_lang: str, connection_id: str) -> List[str]:
    """Szöveg darabok fordítása"""
    tasks = []
    async def translate_with_rate_limit(chunk: str, index: int) -> Tuple[int, str]:
        async with process_semaphore:
            try:
                response = await asyncio.wait_for(
                    client.chat.completions.create(
                        model=os.getenv("MODEL_NAME", "gpt-3.5-turbo"),
                        messages=[
                            {
                                "role": "system",
                                "content": (
                                    "You are a professional translation system. "
                                    "Accurately translate the provided text while preserving its original meaning and formatting. "
                                    "Return only the translated text."
                                )
                            },
                            {
                                "role": "user",
                                "content": f"Please translate this text to {target_lang}:\n\n{chunk}"
                            }
                        ],
                        temperature=0.3
                    ),
                    timeout=45
                )
                result = response.choices[0].message.content.strip()
                progress = min(90, ((index + 1) / len(chunks)) * 90)
                await manager.send_progress(connection_id, progress, f"Translating: {index + 1}/{len(chunks)} chunks...")
                await asyncio.sleep(0.2)
                return index, result
            except asyncio.TimeoutError:
                logger.error(f"Timeout during translation of chunk {index}.")
                raise HTTPException(status_code=408, detail="Translation request timed out.")
            except Exception as e:
                logger.error(f"Error translating chunk {index}: {str(e)}")
                raise
    for i, chunk in enumerate(chunks):
        tasks.append(asyncio.create_task(translate_with_rate_limit(chunk, i)))
    results = await asyncio.gather(*tasks)
    sorted_results = sorted(results, key=lambda x: x[0])
    return [result[1] for result in sorted_results]

@router.post("/")
async def process_video_endpoint(
    url: str = Form(None),
    file: UploadFile = File(None),
    convert_mp3: bool = Form(False),
    generate_subtitles: bool = Form(False),
    subtitle_format: str = Form("srt"),
    platform: str = Form("auto"),
    target_video_format: str = Form(None),
    resolution: str = Form(None),
    bitrate: str = Form(None),
    live_translation: bool = Form(False),
    target_lang: str = Form("en"),
    identify_speakers: bool = Form(False),
    connection_id: str = Form(None)
):
    """Videó feldolgozása"""
    # Bemeneti paraméterek naplózása
    logger.info(f"Video processing request params: url={url}, convert_mp3={convert_mp3}, "
                f"generate_subtitles={generate_subtitles}, subtitle_format={subtitle_format}, "
                f"platform={platform}")
                
    if not connection_id:
        connection_id = str(uuid.uuid4())
    work_dir = TEMP_DIR / connection_id
    work_dir.mkdir(exist_ok=True)

    try:
        await manager.send_progress(connection_id, 5, "Starting video download/processing...")

        video_filepath = None
        system_video_path = None

        # VideóDownloader inicializálása
        downloader = VideoDownloader(connection_id, work_dir, manager)

        if url:
            # Videó letöltése a VideoDownloader segítségével
            video_filepath = await downloader.download_video(url, platform)
            
            if not video_filepath:
                raise HTTPException(status_code=500, detail=f"Failed to download video from {platform} URL")
            
            await manager.send_progress(connection_id, 30, "Video downloaded successfully.")
        elif file:  # Ha fájlt töltött fel a felhasználó
            video_filepath = work_dir / sanitize_filename(file.filename)
            async with aiofiles.open(video_filepath, "wb") as f:
                await f.write(await file.read())
            await manager.send_progress(connection_id, 10, "Processing video file...")
        else:
            raise HTTPException(status_code=400, detail="Either URL or file must be provided")

        download_urls = {}
        
        # A letöltött/feltöltött fájl áthelyezése a SYSTEM_DOWNLOADS-ba
        if video_filepath and video_filepath.exists():
            system_video_path = SYSTEM_DOWNLOADS / video_filepath.name
            if system_video_path.exists():
                system_video_path.unlink()  # Ha már létezik, töröljük
            
            # Másolás helyett egyszerű fájlmásolás, hogy megőrizzük az eredeti fájlt is
            shutil.copy2(str(video_filepath), str(system_video_path))
            download_urls["video_mp4"] = f"/download/{system_video_path.name}"
            logger.info(f"Video copied to: {system_video_path} (size: {system_video_path.stat().st_size} bytes)")
            
            # További feldolgozáshoz az eredeti fájlt használjuk a temp könyvtárban
            # A video_filepath marad az eredeti értéken, nem váltunk át a system_video_path-ra
        else:
            logger.warning(f"Video file does not exist at path: {video_filepath}")
            raise HTTPException(status_code=500, detail="Video file not found after download")

        # Opcionális MP3 konverzió
        if convert_mp3:
            await manager.send_progress(connection_id, 40, "Converting video to MP3...")
            mp3_filename = system_video_path.stem + ".mp3"
            system_mp3_path = SYSTEM_DOWNLOADS / mp3_filename
            # Az eredeti videófájl elérési útját használjuk a konverzióhoz (video_filepath), nem a system_video_path-t
            success = await convert_to_mp3(video_filepath, system_mp3_path)
            if not success:
                raise HTTPException(status_code=500, detail="MP3 conversion failed.")
            download_urls["audio_mp3"] = f"/download/{mp3_filename}"
            await manager.send_progress(connection_id, 60, "MP3 conversion complete.")

        # Felirat generálása és opcionális fordítás
        text = ""
        live_translation_text = ""
        segments = []
        
        if generate_subtitles:
            await manager.send_progress(connection_id, 70, "Generating subtitles...")
            
            try:
                # Használjuk a _transcribe_audio függvényt a transcriber modulból gyors feldolgozásra
                # Az eredeti videófájl elérési útját használjuk a feliratozáshoz (video_filepath), nem a system_video_path-t
                # fast_mode=True paraméter hozzáadva a gyors leiratoláshoz
                text, format_info = await _transcribe_audio(
                    video_filepath, 
                    identify_speakers=identify_speakers,
                    fast_mode=True
                )
                segments = format_info.get("segments", []) if format_info else []

                if live_translation:
                    chunks = chunk_text_by_tokens(text)
                    translated_chunks = await translate_chunks(chunks, target_lang, connection_id)
                    live_translation_text = "\n".join(translated_chunks)

                # Felirat formátum készítése
                if subtitle_format.lower() == "srt":
                    subtitles = await generate_srt(segments)
                elif subtitle_format.lower() == "vtt":
                    subtitles = await generate_vtt(segments)
                else:
                    raise HTTPException(status_code=400, detail="Invalid subtitle format.")

                subtitle_filename = system_video_path.stem + f".{subtitle_format}"
                system_subtitle_path = SYSTEM_DOWNLOADS / subtitle_filename

                async with aiofiles.open(system_subtitle_path, "w", encoding="utf-8") as f:
                    await f.write(subtitles)
                download_urls["subtitles"] = f"/download/{subtitle_filename}"
                
                # Ha fordítás is van, akkor a fordított feliratot is mentsük el
                if live_translation:
                    translated_subtitle_filename = system_video_path.stem + f"_translated.{subtitle_format}"
                    translated_subtitle_path = SYSTEM_DOWNLOADS / translated_subtitle_filename
                    
                    # Itt feltételezzük hogy a szegmensek ugyanarra a time-indexre vonatkoznak, csak a szöveg változik
                    if subtitle_format.lower() == "srt":
                        translated_segments = []
                        # A fordított szöveget elosztjuk a szegmensek között
                        if segments:
                            chunk_length = len(live_translation_text) / len(segments)
                            for i, seg in enumerate(segments):
                                start_idx = int(i * chunk_length)
                                end_idx = int((i+1) * chunk_length) if i < len(segments)-1 else len(live_translation_text)
                                segment_text = live_translation_text[start_idx:end_idx].strip()
                                translated_segments.append({
                                    'start': seg['start'],
                                    'end': seg['end'],
                                    'text': segment_text
                                })
                        else:
                            translated_segments = [{'start': 0, 'end': 5, 'text': live_translation_text}]
                            
                        translated_subtitles = await generate_srt(translated_segments)
                    else:
                        # VTT esetén is hasonlóan
                        translated_segments = []
                        if segments:
                            chunk_length = len(live_translation_text) / len(segments)
                            for i, seg in enumerate(segments):
                                start_idx = int(i * chunk_length)
                                end_idx = int((i+1) * chunk_length) if i < len(segments)-1 else len(live_translation_text)
                                segment_text = live_translation_text[start_idx:end_idx].strip()
                                translated_segments.append({
                                    'start': seg['start'],
                                    'end': seg['end'],
                                    'text': segment_text
                                })
                        else:
                            translated_segments = [{'start': 0, 'end': 5, 'text': live_translation_text}]
                        
                        translated_subtitles = await generate_vtt(translated_segments)
                    
                    async with aiofiles.open(translated_subtitle_path, "w", encoding="utf-8") as f:
                        await f.write(translated_subtitles)
                    download_urls["translated_subtitles"] = f"/download/{translated_subtitle_filename}"
                
                await manager.send_progress(connection_id, 90, "Subtitle generation complete.")
            except Exception as e:
                logger.error(f"Error generating subtitles: {str(e)}")
                await manager.send_progress(connection_id, 80, f"Error generating subtitles: {str(e)}")

        # Transcript adatok mentése
        transcript_data = {
            "text": text,
            "live_translation_text": live_translation_text,
            "segments": segments,
            "original_filename": str(system_video_path.name) if system_video_path else ""
        }
        
        transcript_data_path = work_dir / "transcript_data.json"
        async with aiofiles.open(transcript_data_path, "w") as f:
            await f.write(json.dumps(transcript_data))

        # Video formátum konverzió
        if target_video_format or resolution or bitrate:
            supported_video_formats = ["mp4", "webm", "mkv", "avi", "mov", "flv"]
            if target_video_format and target_video_format.lower() not in supported_video_formats:
                target_video_format = "mp4"

            new_ext = target_video_format if target_video_format else system_video_path.suffix.lstrip('.')
            converted_filename = f"{system_video_path.stem}_converted.{new_ext}"
            converted_path = SYSTEM_DOWNLOADS / converted_filename

            ffmpeg_cmd = ["ffmpeg", "-i", str(system_video_path)]
            if resolution:
                res_map = {"4k": "2160", "1080p": "1080", "720p": "720", "480p": "480", "360p": "360", "240p": "240"}
                res = res_map.get(resolution.lower())
                if res:
                    ffmpeg_cmd += ["-vf", f"scale=-2:{res}"]

            if bitrate:
                quality_map = {"low": "500k", "medium": "1000k", "high": "2500k", "ultra": "5000k"}
                quality = quality_map.get(bitrate.lower())
                if quality:
                    ffmpeg_cmd += ["-b:v", quality]

            # Codec beállítások a formátum alapján
            if target_video_format:
                if target_video_format.lower() == "mp4":
                    ffmpeg_cmd = ["ffmpeg", "-hwaccel", "auto"] + ffmpeg_cmd[1:] + ["-c:v", "libx264", "-c:a", "aac", "-preset", "veryfast", "-crf", "22", "-threads", "0"]
                elif target_video_format.lower() == "webm":
                    ffmpeg_cmd = ["ffmpeg", "-hwaccel", "auto"] + ffmpeg_cmd[1:] + ["-c:v", "libvpx-vp9", "-c:a", "libopus", "-b:a", "128k", "-threads", "0"]
                elif target_video_format.lower() == "mkv":
                    ffmpeg_cmd = ["ffmpeg", "-hwaccel", "auto"] + ffmpeg_cmd[1:] + ["-c:v", "libx264", "-c:a", "copy", "-preset", "veryfast", "-crf", "22", "-threads", "0"]
                elif target_video_format.lower() == "avi":
                    ffmpeg_cmd = ["ffmpeg", "-hwaccel", "auto"] + ffmpeg_cmd[1:] + ["-c:v", "mpeg4", "-c:a", "mp2", "-q:v", "6", "-threads", "0"]
                elif target_video_format.lower() == "mov":
                    ffmpeg_cmd = ["ffmpeg", "-hwaccel", "auto"] + ffmpeg_cmd[1:] + ["-c:v", "prores", "-c:a", "pcm_s16le", "-profile:v", "2", "-threads", "0"]
                elif target_video_format.lower() == "flv":
                    ffmpeg_cmd = ["ffmpeg", "-hwaccel", "auto"] + ffmpeg_cmd[1:] + ["-c:v", "flv", "-c:a", "mp3", "-ar", "44100", "-threads", "0"]
                else:
                    ffmpeg_cmd = ["ffmpeg", "-hwaccel", "auto"] + ffmpeg_cmd[1:] + ["-c:v", "libx264", "-c:a", "copy", "-preset", "veryfast", "-crf", "22", "-threads", "0"]
            else:
                ffmpeg_cmd = ["ffmpeg", "-hwaccel", "auto"] + ffmpeg_cmd[1:] + ["-c:v", "libx264", "-c:a", "copy", "-preset", "veryfast", "-crf", "22", "-threads", "0"]
                
            ffmpeg_cmd += [str(converted_path), "-y"]

            #FFmpeg konvertálás
            try:
                await _convert_with_ffmpeg(ffmpeg_cmd, converted_path, connection_id)
                download_urls["video_converted"] = f"/download/{converted_filename}"
            except Exception as e:
                logger.error(f"Error converting video format: {str(e)}")
                await manager.send_progress(connection_id, 95, f"Error converting video format: {str(e)}")

        # A letöltött fájlok elérhetőségének ellenőrzése
        for key, path_str in list(download_urls.items()):
            if not path_str.startswith("/download/"):
                continue
                
            filename = path_str.split('/')[-1]
            file_path = SYSTEM_DOWNLOADS / filename
            if not file_path.exists() or file_path.stat().st_size == 0:
                logger.warning(f"Download file is missing or empty: {file_path}")
                download_urls.pop(key, None)

        await manager.send_progress(connection_id, 100, "Video processing complete!")
        return JSONResponse({
            "download_urls": download_urls,
            "transcript_download_url": f"/download_transcript/{connection_id}/txt" if generate_subtitles else None,
            "transcription_text": text if generate_subtitles else None,
            "live_translation_text": live_translation_text if generate_subtitles and live_translation else None
        })

    except Exception as e:
        logger.error(f"Video processing error: {e}")
        raise HTTPException(status_code=500, detail=f"Video processing failed: {str(e)}")

@router.get("/download_transcript/{connection_id}/{format}")
async def download_transcript(connection_id: str, format: str):
    """Transzkripciós fájl letöltése"""
    work_dir = TEMP_DIR / connection_id
    if not work_dir.exists():
        raise HTTPException(status_code=404, detail="Transcript not found")
    transcript_data_path = work_dir / "transcript_data.json"
    if not transcript_data_path.exists():
        raise HTTPException(status_code=404, detail="Transcript data not found.")
    
    async with aiofiles.open(transcript_data_path, "r") as f:
        transcript_data = json.loads(await f.read())
    
    text = transcript_data.get("text", "")
    segments = transcript_data.get("segments", [])
    original_filename = transcript_data.get("original_filename", "transcript")
    name_without_ext = Path(original_filename).stem
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    output_filename = f"{name_without_ext}_{timestamp}.{format}"
    output_path = SYSTEM_DOWNLOADS / output_filename
    
    if format == "docx":
        await generate_transcript_docx(text, segments, output_path)
    elif format == "pdf":
        await generate_transcript_pdf(text, segments, output_path)
    elif format == "txt":
        async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
            await f.write(text)
    elif format == "srt":
        srt_content = await generate_srt(segments)
        async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
            await f.write(srt_content)
    elif format == "vtt":
        vtt_content = await generate_vtt(segments)
        async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
            await f.write(vtt_content)
    else:
        raise HTTPException(status_code=400, detail="Invalid transcript format")
    
    return FileResponse(
        path=output_path,
        filename=output_filename,
        media_type="application/octet-stream",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(output_filename)}"}
    )

@router.post("/convert")
async def convert_video(
    file: UploadFile = File(...),
    target_format: str = Form("mp4"),
    connection_id: Optional[str] = Form(None)
):
    """Convert video to a different format (specifically for MKV to MP4 conversion)"""
    if not connection_id:
        connection_id = str(uuid.uuid4())
    
    # Log file details for debugging
    logger.info(f"Received file for conversion: {file.filename}, size: {file.size}, content_type: {file.content_type}")
    logger.info(f"Target format: {target_format}, connection_id: {connection_id}")
    
    work_dir = TEMP_DIR / connection_id
    work_dir.mkdir(exist_ok=True)
    
    try:
        # Save uploaded file
        input_path = work_dir / sanitize_filename(file.filename)
        logger.info(f"Saving file to: {input_path}")
        
        async with aiofiles.open(input_path, "wb") as f:
            content = await file.read()
            await f.write(content)
            logger.info(f"Saved file of size: {len(content)} bytes")
        
        await manager.send_progress(connection_id, 20, "Videó formátum konvertálása...")
        
        # Determine file extension
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        # Prepare output filename (changing extension to target format)
        base_filename = os.path.splitext(file.filename)[0]
        output_filename = f"{base_filename}_converted.{target_format}"
        output_path = SYSTEM_DOWNLOADS / output_filename
        
        # MKV to MP4 conversion using FFmpeg with maximum compatibility parameters
        ffmpeg_cmd = [
            "ffmpeg",
            "-hwaccel", "auto",    # Enable hardware acceleration if available 
            "-i", str(input_path),
            "-c:v", "libx264",     # Use H.264, most compatible codec
            "-profile:v", "high",  # Use high profile for better quality
            "-level", "4.0",       # Compatibility level
            "-preset", "veryfast", # Faster encoding (was medium)
            "-crf", "22",          # Higher quality (lower value)
            "-c:a", "aac",         # AAC audio
            "-b:a", "192k",        # Higher audio bitrate
            "-ar", "44100",        # Standard audio sample rate
            "-pix_fmt", "yuv420p", # Most compatible pixel format
            "-movflags", "+faststart",  # Optimize for web playback
            "-threads", "0",       # Use all available CPU threads
            str(output_path),
            "-y"
        ]
        
        # Run the FFmpeg command
        await manager.send_progress(connection_id, 30, "Videó konténer konvertálása...")
        process = await asyncio.create_subprocess_exec(
            *ffmpeg_cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await process.communicate()
        
        if process.returncode != 0:
            logger.error(f"FFmpeg command failed: {stderr.decode()}")
            await manager.send_progress(connection_id, 100, f"Error: {stderr.decode()[:100]}...")
            raise HTTPException(status_code=500, detail="Videó konvertálás sikertelen")
        
        await manager.send_progress(connection_id, 100, "Videó konvertálás befejeződött")
        
        # Return result
        return JSONResponse({
            "download_url": f"/download/{output_filename}",
            "original_format": file_ext.lstrip('.'),
            "converted_format": target_format,
            "message": f"A videó sikeresen konvertálva {target_format} formátumba"
        })
    
    except Exception as e:
        logger.error(f"Video conversion error: {e}")
        raise HTTPException(status_code=500, detail=f"Videó konvertálás sikertelen: {str(e)}")
    
    finally:
        # Clean up work directory (we keep it for debugging purposes)
        # await async_rmtree(work_dir)
        pass

@router.post("/extract_audio")
async def extract_audio(
    file: UploadFile = File(...),
    target_format: str = Form("mp3"),
    connection_id: Optional[str] = Form(None)
):
    """Hang kinyerése videófájlból"""
    if not connection_id:
        connection_id = str(uuid.uuid4())
    
    work_dir = TEMP_DIR / connection_id
    work_dir.mkdir(exist_ok=True)
    
    try:
        input_path = work_dir / sanitize_filename(file.filename)
        async with aiofiles.open(input_path, "wb") as f:
            await f.write(await file.read())
        
        await manager.send_progress(connection_id, 10, "Processing video file...")
        
        # Támogatott kimeneti formátumok
        supported_formats = ["mp3", "wav", "ogg", "aac", "flac"]
        if target_format not in supported_formats:
            target_format = "mp3"
        
        output_filename = f"{Path(file.filename).stem}_audio.{target_format}"
        output_path = SYSTEM_DOWNLOADS / output_filename
        
        ffmpeg_cmd = ["ffmpeg", "-i", str(input_path), "-vn"]
        
        # Format-specifikus beállítások
        if target_format == "mp3":
            ffmpeg_cmd += ["-c:a", "libmp3lame", "-q:a", "2", "-threads", "0"]
        elif target_format == "wav":
            ffmpeg_cmd += ["-c:a", "pcm_s16le", "-ar", "44100", "-ac", "2", "-threads", "0"]
        elif target_format == "ogg":
            ffmpeg_cmd += ["-c:a", "libvorbis", "-q:a", "4", "-threads", "0"]
        elif target_format == "aac":
            ffmpeg_cmd += ["-c:a", "aac", "-b:a", "192k", "-threads", "0"]
        elif target_format == "flac":
            ffmpeg_cmd += ["-c:a", "flac", "-threads", "0"]
        
        ffmpeg_cmd += [str(output_path), "-y"]
        
        process = await asyncio.create_subprocess_exec(
            *ffmpeg_cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await process.communicate()
        
        if process.returncode != 0:
            logger.error(f"Audio extraction failed: {stderr.decode()}")
            raise HTTPException(status_code=500, detail="Audio extraction failed")
        
        await manager.send_progress(connection_id, 100, "Audio extraction complete")
        
        return JSONResponse({
            "download_url": f"/download/{output_filename}"
        })
    
    except Exception as e:
        logger.error(f"Audio extraction error: {e}")
        raise HTTPException(status_code=500, detail=f"Audio extraction failed: {str(e)}")
    
    finally:
        await async_rmtree(work_dir)

@router.post("/video_effects")
async def apply_video_effects(
    file: UploadFile = File(...),
    effect_type: str = Form("none"),  # none, grayscale, sepia, vignette, blur, sharpen
    rotate: int = Form(0),  # 0, 90, 180, 270
    flip_horizontal: bool = Form(False),
    flip_vertical: bool = Form(False),
    output_format: str = Form("mp4"),
    connection_id: Optional[str] = Form(None)
):
    """Effektek alkalmazása videón"""
    if not connection_id:
        connection_id = str(uuid.uuid4())
    
    work_dir = TEMP_DIR / connection_id
    work_dir.mkdir(exist_ok=True)
    
    try:
        input_path = work_dir / sanitize_filename(file.filename)
        async with aiofiles.open(input_path, "wb") as f:
            await f.write(await file.read())
        
        await manager.send_progress(connection_id, 10, "Processing video file...")
        
        # Kimenet elkészítése
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"effect_video_{timestamp}.{output_format}"
        output_path = SYSTEM_DOWNLOADS / output_filename
        
        # Alap ffmpeg parancs
        ffmpeg_cmd = ["ffmpeg", "-i", str(input_path)]
        
        # Effektek hozzáadása
        filter_complex = []
        
        # Forgatás
        if rotate == 90:
            filter_complex.append("transpose=1")
        elif rotate == 180:
            filter_complex.append("transpose=2,transpose=2")
        elif rotate == 270:
            filter_complex.append("transpose=2")
        
        # Tükrözés
        if flip_horizontal:
            filter_complex.append("hflip")
        if flip_vertical:
            filter_complex.append("vflip")
        
        # Vizuális effektek
        if effect_type == "grayscale":
            filter_complex.append("colorchannelmixer=.3:.4:.3:0:.3:.4:.3:0:.3:.4:.3")
        elif effect_type == "sepia":
            filter_complex.append("colorchannelmixer=.393:.769:.189:0:.349:.686:.168:0:.272:.534:.131")
        elif effect_type == "vignette":
            filter_complex.append("vignette=PI/4")
        elif effect_type == "blur":
            filter_complex.append("boxblur=5:1")
        elif effect_type == "sharpen":
            filter_complex.append("unsharp=5:5:1.0:5:5:0.0")
        
        # Filterek hozzáadása a parancshoz
        if filter_complex:
            ffmpeg_cmd += ["-vf", ",".join(filter_complex)]
        
        # Codec és output
        ffmpeg_cmd += [
            "-c:v", "libx264",
            "-preset", "veryfast",
            "-crf", "22",
            "-threads", "0",
            "-hwaccel", "auto",
            "-c:a", "copy",
            str(output_path),
            "-y"
        ]
        
        process = await asyncio.create_subprocess_exec(
            *ffmpeg_cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await process.communicate()
        
        if process.returncode != 0:
            logger.error(f"Video effect application failed: {stderr.decode()}")
            raise HTTPException(status_code=500, detail="Video effect application failed")
        
        await manager.send_progress(connection_id, 100, "Video effects applied")
        
        return JSONResponse({
            "download_url": f"/download/{output_filename}"
        })
    
    except Exception as e:
        logger.error(f"Video effects error: {e}")
        raise HTTPException(status_code=500, detail=f"Video effects failed: {str(e)}")
    
    finally:
        await async_rmtree(work_dir)

@router.post("/video_watermark")
async def apply_watermark(
    video_file: UploadFile = File(...),
    watermark_file: UploadFile = File(None),
    watermark_text: str = Form(None),
    position: str = Form("bottomright"),  # topleft, topright, bottomleft, bottomright, center
    opacity: float = Form(0.5),  # 0.0 - 1.0
    output_format: str = Form("mp4"),
    connection_id: Optional[str] = Form(None)
):
    """Vízjel hozzáadása videóhoz"""
    if not connection_id:
        connection_id = str(uuid.uuid4())
    
    if not watermark_file and not watermark_text:
        raise HTTPException(status_code=400, detail="Either watermark file or text must be provided")
    
    work_dir = TEMP_DIR / connection_id
    work_dir.mkdir(exist_ok=True)
    
    try:
        video_path = work_dir / sanitize_filename(video_file.filename)
        async with aiofiles.open(video_path, "wb") as f:
            await f.write(await video_file.read())
        
        await manager.send_progress(connection_id, 10, "Processing video file...")
        
        # Pozíció beállítása
        positions = {
            "topleft": "10:10",
            "topright": "main_w-overlay_w-10:10",
            "bottomleft": "10:main_h-overlay_h-10",
            "bottomright": "main_w-overlay_w-10:main_h-overlay_h-10",
            "center": "(main_w-overlay_w)/2:(main_h-overlay_h)/2"
        }
        pos = positions.get(position, positions["bottomright"])
        
        # Kimenet elkészítése
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"watermarked_video_{timestamp}.{output_format}"
        output_path = SYSTEM_DOWNLOADS / output_filename
        
        if watermark_file:
            # Kép vízjel
            watermark_path = work_dir / sanitize_filename(watermark_file.filename)
            async with aiofiles.open(watermark_path, "wb") as f:
                await f.write(await watermark_file.read())
            
            # FFmpeg parancs kép vízjelhez
            ffmpeg_cmd = [
                "ffmpeg",
                "-i", str(video_path),
                "-i", str(watermark_path),
                "-filter_complex", f"[1:v]format=rgba,colorchannelmixer=aa={opacity}[logo];[0:v][logo]overlay={pos}",
                "-c:v", "libx264",
                "-c:a", "copy",
                str(output_path),
                "-y"
            ]
        else:
            # Szöveg vízjel
            # FFmpeg parancs szöveg vízjelhez
            ffmpeg_cmd = [
                "ffmpeg",
                "-i", str(video_path),
                "-vf", f"drawtext=text='{watermark_text}':fontsize=24:fontcolor=white@{opacity}:x={pos.split(':')[0]}:y={pos.split(':')[1]}:shadowcolor=black:shadowx=2:shadowy=2",
                "-c:v", "libx264",
                "-c:a", "copy",
                str(output_path),
                "-y"
            ]
        
        process = await asyncio.create_subprocess_exec(
            *ffmpeg_cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await process.communicate()
        
        if process.returncode != 0:
            logger.error(f"Video watermarking failed: {stderr.decode()}")
            raise HTTPException(status_code=500, detail="Video watermarking failed")
        
        await manager.send_progress(connection_id, 100, "Video watermarking complete")
        
        return JSONResponse({
            "download_url": f"/download/{output_filename}"
        })
    
    except Exception as e:
        logger.error(f"Video watermarking error: {e}")
        raise HTTPException(status_code=500, detail=f"Video watermarking failed: {str(e)}")
    
    finally:
        await async_rmtree(work_dir)
        
@router.post("/video_compress")
async def compress_video(
    file: UploadFile = File(...),
    quality: str = Form("medium"),  # low, medium, high
    output_format: str = Form("mp4"),
    connection_id: Optional[str] = Form(None)
):
    """Videó tömörítése"""
    if not connection_id:
        connection_id = str(uuid.uuid4())
    
    work_dir = TEMP_DIR / connection_id
    work_dir.mkdir(exist_ok=True)
    
    try:
        input_path = work_dir / sanitize_filename(file.filename)
        async with aiofiles.open(input_path, "wb") as f:
            await f.write(await file.read())
        
        await manager.send_progress(connection_id, 10, "Processing video file...")
        
        # Minőség beállítások
        quality_settings = {
            "low": {"crf": "28", "preset": "veryfast", "scale": "640:-2"},
            "medium": {"crf": "23", "preset": "veryfast", "scale": "1280:-2"},
            "high": {"crf": "18", "preset": "medium", "scale": "1920:-2"}
        }
        settings = quality_settings.get(quality, quality_settings["medium"])
        
        # Kimenet elkészítése
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"compressed_video_{timestamp}.{output_format}"
        output_path = SYSTEM_DOWNLOADS / output_filename
        
        # FFmpeg parancs a tömörítéshez
        ffmpeg_cmd = [
            "ffmpeg",
            "-hwaccel", "auto",
            "-i", str(input_path),
            "-vf", f"scale={settings['scale']}",
            "-c:v", "libx264",
            "-crf", settings["crf"],
            "-preset", settings["preset"],
            "-threads", "0",
            "-c:a", "aac",
            "-b:a", "128k",
            str(output_path),
            "-y"
        ]
        
        process = await asyncio.create_subprocess_exec(
            *ffmpeg_cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await process.communicate()
        
        if process.returncode != 0:
            logger.error(f"Video compression failed: {stderr.decode()}")
            raise HTTPException(status_code=500, detail="Video compression failed")
        
        await manager.send_progress(connection_id, 100, "Video compression complete")
        
        return JSONResponse({
            "download_url": f"/download/{output_filename}"
        })
    
    except Exception as e:
        logger.error(f"Video compression error: {e}")
        raise HTTPException(status_code=500, detail=f"Video compression failed: {str(e)}")
    
    finally:
        await async_rmtree(work_dir)

@router.post("/merge")
async def merge_videos(
    files: List[UploadFile] = File(...),
    output_format: str = Form("mp4"),
    connection_id: Optional[str] = Form(None)
):
    """Videók összefűzése"""
    if not connection_id:
        connection_id = str(uuid.uuid4())
    
    if len(files) < 2:
        raise HTTPException(status_code=400, detail="At least two video files are required for merging")
    
    work_dir = TEMP_DIR / connection_id
    work_dir.mkdir(exist_ok=True)
    
    try:
        # Bemeneti fájlok mentése
        file_list_path = work_dir / "filelist.txt"
        input_paths = []
        
        for i, file in enumerate(files):
            input_path = work_dir / f"{i}_{sanitize_filename(file.filename)}"
            async with aiofiles.open(input_path, "wb") as f:
                await f.write(await file.read())
            input_paths.append(input_path)
        
        # Filelist létrehozása a concat demuxer számára
        async with aiofiles.open(file_list_path, "w") as f:
            for path in input_paths:
                await f.write(f"file '{path}'\n")
        
        await manager.send_progress(connection_id, 20, "Processing video files...")
        
        # Kimenet elkészítése
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"merged_video_{timestamp}.{output_format}"
        output_path = SYSTEM_DOWNLOADS / output_filename
        
        # FFmpeg parancs az összefűzéshez
        ffmpeg_cmd = [
            "ffmpeg",
            "-f", "concat",
            "-safe", "0",
            "-i", str(file_list_path),
            "-c", "copy",
            str(output_path),
            "-y"
        ]
        
        process = await asyncio.create_subprocess_exec(
            *ffmpeg_cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await process.communicate()
        
        if process.returncode != 0:
            logger.error(f"Video merging failed: {stderr.decode()}")
            raise HTTPException(status_code=500, detail="Video merging failed")
        
        await manager.send_progress(connection_id, 100, "Video merging complete")
        
        return JSONResponse({
            "download_url": f"/download/{output_filename}"
        })
    
    except Exception as e:
        logger.error(f"Video merging error: {e}")
        raise HTTPException(status_code=500, detail=f"Video merging failed: {str(e)}")
    
    finally:
        await async_rmtree(work_dir)

@router.post("/trim")
async def trim_video(
    file: UploadFile = File(...),
    start_time: str = Form(...),  # Format: "00:00:00"
    end_time: str = Form(...),    # Format: "00:00:00"
    output_format: str = Form("mp4"),
    connection_id: Optional[str] = Form(None)
):
    """Videó vágása"""
    if not connection_id:
        connection_id = str(uuid.uuid4())
    
    work_dir = TEMP_DIR / connection_id
    work_dir.mkdir(exist_ok=True)
    
    try:
        input_path = work_dir / sanitize_filename(file.filename)
        async with aiofiles.open(input_path, "wb") as f:
            await f.write(await file.read())
        
        await manager.send_progress(connection_id, 10, "Processing video file...")
        
        # Kimenet elkészítése
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"trimmed_video_{timestamp}.{output_format}"
        output_path = SYSTEM_DOWNLOADS / output_filename
        
        # FFmpeg parancs a vágáshoz
        ffmpeg_cmd = [
            "ffmpeg",
            "-hwaccel", "auto",
            "-i", str(input_path),
            "-ss", start_time,
            "-to", end_time,
            "-c:v", "libx264",
            "-preset", "veryfast",
            "-threads", "0",
            "-c:a", "aac",
            str(output_path),
            "-y"
        ]
        
        process = await asyncio.create_subprocess_exec(
            *ffmpeg_cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await process.communicate()
        
        if process.returncode != 0:
            logger.error(f"Video trimming failed: {stderr.decode()}")
            raise HTTPException(status_code=500, detail="Video trimming failed")
        
        await manager.send_progress(connection_id, 100, "Video trimming complete")
        
        return JSONResponse({
            "download_url": f"/download/{output_filename}"
        })
    
    except Exception as e:
        logger.error(f"Video trimming error: {e}")
        raise HTTPException(status_code=500, detail=f"Video trimming failed: {str(e)}")
    
    finally:
        await async_rmtree(work_dir)
