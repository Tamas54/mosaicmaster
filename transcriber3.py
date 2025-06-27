import uuid
import json
import os
import asyncio
import logging
import re
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional

import aiofiles
from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from fastapi.responses import JSONResponse, FileResponse
from pydub import AudioSegment
from pydub.exceptions import CouldntDecodeError
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from openai import AsyncOpenAI

from config import (
    TEMP_DIR,
    SYSTEM_DOWNLOADS,
    manager,
    logger,
    client,
    sanitize_filename,
    async_rmtree
)

router = APIRouter()

# --- Segédfüggvény az időbélyeg formázásához ---
def format_timestamp(seconds: float) -> str:
    """Másodpercek átalakítása SRT időbélyeg formátumra (HH:MM:SS,mmm)."""
    if seconds < 0: seconds = 0
    total_milliseconds = int(seconds * 1000)
    hours, remainder = divmod(total_milliseconds, 3600000)
    minutes, remainder = divmod(remainder, 60000)
    secs, milliseconds = divmod(remainder, 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{milliseconds:03d}"

# --- Felirat Generátorok ---

# SRT generátor
def create_srt_content(segments: List[Dict[str, Any]]) -> str:
    """SRT formátumú tartalom generálása szegmensekből."""
    srt_parts = []
    for i, segment in enumerate(segments, 1):
        start_time = format_timestamp(segment['start'])
        end_time = format_timestamp(segment['end'])
        text = segment.get('text', '').strip()
        if not text: continue
        srt_parts.append(f"{i}\n{start_time} --> {end_time}\n{text}\n")
    return "\n".join(srt_parts)

# VTT generátor
async def generate_vtt(segments: List[Dict[str, Any]]) -> str:
    """VTT (WebVTT) formátumú tartalom generálása szegmensekből."""
    vtt_output = "WEBVTT\n\n"
    def format_vtt_timestamp(seconds: float) -> str:
        if seconds < 0: seconds = 0
        total_milliseconds = int(seconds * 1000)
        hours, remainder = divmod(total_milliseconds, 3600000)
        minutes, remainder = divmod(remainder, 60000)
        secs, milliseconds = divmod(remainder, 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{milliseconds:03d}"
    for segment in segments:
        start_time = format_vtt_timestamp(segment['start'])
        end_time = format_vtt_timestamp(segment['end'])
        text = segment.get('text', '').strip()
        if not text: continue
        vtt_output += f"{start_time} --> {end_time}\n{text}\n\n"
    return vtt_output

# TXT konvertáló
def convert_txt_to_srt(txt_content: str) -> str:
    """Speciális TXT formátum ([start-end] text) konvertálása SRT formátumra."""
    lines = txt_content.strip().split('\n')
    srt_parts = []
    counter = 1
    for line in lines:
        line = line.strip()
        match = re.match(r'^\[(\d+\.?\d*)-(\d+\.?\d*)\]\s*(.*)', line)
        if not match: continue
        try:
            start_sec_str, end_sec_str, text = match.groups()
            start_sec = float(start_sec_str)
            end_sec = float(end_sec_str)
            if start_sec >= end_sec:
                logger.warning(f"Skipping invalid time: {line}")
                continue
            start_time = format_timestamp(start_sec)
            end_time = format_timestamp(end_sec)
            text_content = text.strip()
            if not text_content: continue
            srt_parts.append(f"{counter}\n{start_time} --> {end_time}\n{text_content}\n")
            counter += 1
        except ValueError as e:
            logger.warning(f"Skipping conversion error: {line}, Error: {e}")
        except Exception as e:
            logger.error(f"Error converting line: {line}, Error: {e}")
    return "\n".join(srt_parts)

# --- Jegyzőkönyv Generátorok ---

async def generate_transcript_docx(text: str, segments: List[Dict[str, Any]], output_path: Path):
    doc = Document()
    try:
        styles = doc.styles['Normal']
    except KeyError:
        logger.warning("Default 'Normal' style not found.")
    if segments:
        for segment in segments:
            text_content = segment.get('text', '').strip()
            speaker = segment.get('speaker', '')
            timestamp = f"[{segment['start']:.2f}-{segment['end']:.2f}]"
            if not text_content: continue
            p = doc.add_paragraph()
            p.add_run(f"{timestamp} ").bold = True
            if speaker: p.add_run(f"Beszélő {speaker}: ").italic = True
            p.add_run(text_content)
    elif text:
        doc.add_paragraph(text)
    else:
        doc.add_paragraph("No content.")
    loop = asyncio.get_running_loop()
    await loop.run_in_executor(None, doc.save, str(output_path))

async def generate_transcript_pdf(text: str, segments: List[Dict[str, Any]], output_path: Path):
    doc = SimpleDocTemplate(str(output_path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    if segments:
        for segment in segments:
            text_content = segment.get('text', '').strip()
            speaker = segment.get('speaker', '')
            timestamp = f"[{segment['start']:.2f}-{segment['end']:.2f}]"
            if not text_content: continue
            speaker_info = f"Beszélő {speaker}: " if speaker else ""
            from reportlab.lib.utils import escapeOnce
            line_text = f"<b>{escapeOnce(timestamp)}</b> <i>{escapeOnce(speaker_info)}</i>{escapeOnce(text_content)}"
            story.append(Paragraph(line_text, styles['Normal']))
            story.append(Spacer(1, 6))
    elif text:
        from reportlab.lib.utils import escapeOnce
        story.append(Paragraph(escapeOnce(text), styles['Normal']))
    else:
        story.append(Paragraph("No content.", styles['Normal']))
    loop = asyncio.get_running_loop()
    await loop.run_in_executor(None, doc.build, story)

async def generate_transcript_txt(segments: List[Dict[str, Any]], output_path: Path, with_timestamps: bool = True, with_speakers: bool = True):
    """Transzkript generálása szabványos SRT formátumban."""
    # Használjuk a create_srt_content függvényt a szabványos SRT formátumhoz
    content = create_srt_content(segments)
    async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
        await f.write(content)
        
async def generate_transcript_legacy_txt(segments: List[Dict[str, Any]], output_path: Path, with_timestamps: bool = True, with_speakers: bool = True):
    """Transzkript generálása legacy TXT formátumban [start-end] formátumú időbélyegekkel."""
    async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
        current_speaker = None
        first_line = True
        for segment in segments:
            text = segment.get('text', '').strip()
            if not text: continue
            speaker = segment.get('speaker', '') if with_speakers else ''
            timestamp_text = f"[{segment['start']:.2f}-{segment['end']:.2f}] " if with_timestamps else ""
            line_prefix = timestamp_text
            speaker_change_line = ""
            if speaker and speaker != current_speaker:
                if not first_line: speaker_change_line += "\n"
                speaker_change_line += f"{timestamp_text}Beszélő {speaker}:\n"
                current_speaker = speaker
                line_prefix = ""
            elif speaker and speaker == current_speaker:
                line_prefix = timestamp_text
            elif not speaker:
                line_prefix = timestamp_text
                if current_speaker is not None and not first_line:
                    speaker_change_line += "\n"
                current_speaker = None
            if speaker_change_line:
                await f.write(speaker_change_line)
            await f.write(f"{line_prefix}{text}\n")
            first_line = False

async def post_process_transcript_with_styles(segments, output_path):
    speaker_colors = {
        "1": "\033[94m",
        "2": "\033[92m",
        "3": "\033[91m",
        "4": "\033[93m",
        "5": "\033[95m",
        "6": "\033[96m"
    }
    reset_color = "\033[0m"
    default_color = "\033[97m"
    async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
        current_speaker = None
        first_line = True
        for segment in segments:
            text_content = segment.get('text', '').strip()
            if not text_content: continue
            speaker = str(segment.get('speaker', ''))
            timestamp_text = f"[{segment['start']:.2f}-{segment['end']:.2f}] "
            color_code = speaker_colors.get(speaker, default_color)
            line_prefix = timestamp_text
            speaker_change_line = ""
            if speaker and speaker != current_speaker:
                if not first_line: speaker_change_line += "\n"
                speaker_change_line += f"{timestamp_text}{color_code}Beszélő {speaker}{reset_color}:\n"
                current_speaker = speaker
                line_prefix = ""
            elif speaker and speaker == current_speaker:
                line_prefix = timestamp_text
            elif not speaker:
                line_prefix = timestamp_text
                if current_speaker is not None and not first_line:
                    speaker_change_line += "\n"
                current_speaker = None
            if speaker_change_line:
                await f.write(speaker_change_line)
            await f.write(f"{line_prefix}{color_code}{text_content}{reset_color}\n")
            first_line = False

# --- Core Transzkripciós Logika ---

async def _transcribe_audio(audio_path: Path, identify_speakers: bool = False, fast_mode: bool = False) -> tuple:
    if client is None:
        logger.error("OpenAI client not available.")
        raise HTTPException(status_code=503, detail="Transcription service unavailable.")
    loop = asyncio.get_running_loop()
    audio_file_to_transcribe = audio_path
    temp_mp3_path = None
    try:
        logger.debug(f"Loading audio file: {audio_path}")
        audio = await loop.run_in_executor(None, AudioSegment.from_file, str(audio_path))
        logger.debug(f"Audio loaded. Duration: {audio.duration_seconds:.2f}s")
        if audio.duration_seconds < 0.1:
            logger.warning(f"Audio file {audio_path.name} too short")
            raise HTTPException(status_code=400, detail="Audio file too short.")
        if audio_path.suffix.lower() not in [".mp3", ".wav", ".m4a", ".ogg", ".flac", ".opus"]:
            temp_mp3_path = audio_path.with_suffix(".mp3")
            logger.info(f"Converting {audio_path.name} to MP3...")
            await loop.run_in_executor(None, audio.export, str(temp_mp3_path), format="mp3", bitrate="64k")
            audio_file_to_transcribe = temp_mp3_path
            logger.info(f"Converted to {temp_mp3_path.name}")
        file_size_mb = audio_file_to_transcribe.stat().st_size / (1024*1024)
        if file_size_mb > 25:
            logger.error(f"Audio file {audio_file_to_transcribe.name} exceeds 25MB limit")
            raise HTTPException(status_code=413, detail="Audio file size exceeds 25MB limit.")
        logger.debug(f"Audio file size: {file_size_mb:.2f}MB")
        logger.info(f"Starting transcription for {audio_file_to_transcribe.name}...")
        with open(audio_file_to_transcribe, "rb") as f:
            transcript = await client.audio.transcriptions.create(
                model="whisper-1",
                file=f,
                response_format="verbose_json",
                timestamp_granularities=["segment"]
            )
        if not transcript or not hasattr(transcript, 'text') or not hasattr(transcript, 'segments'):
            raise ValueError("Invalid Whisper API response.")
        logger.info(f"Transcription finished for {audio_file_to_transcribe.name}.")
        full_transcript_text = transcript.text or ""
        raw_segments = [s.model_dump() for s in transcript.segments] if transcript.segments else []
    except CouldntDecodeError:
        logger.error(f"Could not decode: {audio_path}")
        raise HTTPException(status_code=400, detail=f"Unsupported format: {audio_path.name}")
    except FileNotFoundError:
        logger.error(f"File not found: {audio_path}")
        raise HTTPException(status_code=404, detail=f"File not found: {audio_path.name}")
    except HTTPException as e:
        raise e
    except Exception as e:
        logger.error(f"Transcription error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Transcription failed: {e}")
    finally:
        if temp_mp3_path and temp_mp3_path.exists():
            try:
                logger.debug(f"Removing temp MP3: {temp_mp3_path}")
                temp_mp3_path.unlink()
            except OSError as e:
                logger.warning(f"Could not remove temp MP3 {temp_mp3_path}: {e}")

    effective_identify_speakers = identify_speakers and not fast_mode
    final_segments = []
    if effective_identify_speakers and raw_segments:
        logger.info(f"Starting speaker ID for {audio_path.name}...")
        segments_for_analysis = [s.copy() for s in raw_segments]
        if len(segments_for_analysis) > 1:
            speaker_system_prompt = """Azonosítsd a különböző beszélőket a következő átiratban. Minden szegmenst jelölj meg egy beszélővel (1, 2, 3, stb.). Törekedj a konzisztenciára. A válaszod KIZÁRÓLAG egy JSON objektum legyen, amely tartalmaz egy "segments" kulcsot."""
            user_content_json = json.dumps([{"start": s['start'], "end": s['end'], "text": s['text']} for s in segments_for_analysis], ensure_ascii=False)
            try:
                speaker_response = await client.chat.completions.create(
                    model=os.getenv("MODEL_NAME", "gpt-3.5-turbo"),
                    messages=[
                        {"role": "system", "content": speaker_system_prompt},
                        {"role": "user", "content": f"Elemezd és azonosítsd a beszélőket: {user_content_json}"}
                    ],
                    response_format={"type": "json_object"},
                    temperature=0.3
                )
                response_content = speaker_response.choices[0].message.content
                if not response_content:
                    raise ValueError("Empty response from speaker ID API.")
                logger.debug(f"Speaker ID raw response: {response_content}")
                speaker_analysis = json.loads(response_content)
                processed_segments_map = {}
                if "segments" in speaker_analysis and isinstance(speaker_analysis["segments"], list) and len(speaker_analysis["segments"]) == len(segments_for_analysis):
                    for i, resp_seg in enumerate(speaker_analysis["segments"]):
                        original_seg = segments_for_analysis[i]
                        if isinstance(resp_seg, dict) and "speaker" in resp_seg:
                            original_seg["speaker"] = str(resp_seg.get("speaker", "1")).strip() or "1"
                        else:
                            logger.warning(f"Bad speaker segment {i}")
                            original_seg["speaker"] = "1"
                        processed_segments_map[original_seg['start']] = original_seg
                    for seg in raw_segments:
                        seg['speaker'] = processed_segments_map.get(seg['start'], {}).get('speaker', '1')
                    final_segments = raw_segments
                    logger.info("Speaker ID results merged.")
                else:
                    logger.warning(f"Speaker ID response format/count incorrect. Assigning default speakers.")
                    for segment in raw_segments:
                        segment["speaker"] = "1"  # Default speaker
                    final_segments = raw_segments
            except Exception as e:
                logger.error(f"Error during speaker identification API call or processing: {str(e)}", exc_info=True)
                # Fallback on error: add empty speaker tag
                for segment in raw_segments:
                    segment["speaker"] = ""
                final_segments = raw_segments
        else:
            logger.info("Not enough segments for speaker identification.")
            for segment in raw_segments:
                segment["speaker"] = "1"  # Alapértelmezett beszélő
            final_segments = raw_segments
    else:
        if fast_mode:
            logger.info(f"Fast mode skipping speaker ID for {audio_path.name}.")
        else:
            logger.info(f"Speaker ID disabled for {audio_path.name}.")
        for segment in raw_segments:
            segment["speaker"] = ""
        final_segments = raw_segments

    final_segments = sorted(final_segments, key=lambda x: x.get("start", 0))
    return full_transcript_text, {"type": "audio", "segments": final_segments}

# --- Endpointok ---

@router.post("/")
async def transcribe_audio(
    file: UploadFile = File(...),
    live_translation: bool = Form(False), target_lang: str = Form("en"),
    identify_speakers: bool = Form(True), fast_mode: bool = Form(False),
    connection_id: Optional[str] = Form(None)
):
    """Audio fájl transzkripciója, alapértelmezett kimenetként SRT-t generál."""
    if not connection_id:
        connection_id = str(uuid.uuid4())
    work_dir = TEMP_DIR / connection_id
    work_dir.mkdir(exist_ok=True)
    input_path = None
    try:
        sanitized_name = sanitize_filename(file.filename) if file.filename else Path(f"audio_{connection_id}.tmp")
        input_path = work_dir / sanitized_name
        # Fájl mentése darabokban
        async with aiofiles.open(input_path, "wb") as f:
            while True:
                chunk = await file.read(1024 * 1024)  # Olvasás darabokban
                if not chunk:
                    break
                await f.write(chunk)
        file_size = input_path.stat().st_size
        logger.info(f"File saved to {input_path} ({file_size} bytes)")
        await manager.send_progress(connection_id, 10, "Processing audio file...")
        transcript_text, format_info = await _transcribe_audio(input_path, identify_speakers=identify_speakers, fast_mode=fast_mode)
        segments = format_info.get("segments", [])
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = input_path.stem
        transcript_filename_srt = f"transcript_{base_filename}_{timestamp}.srt"
        transcript_path_srt = SYSTEM_DOWNLOADS / transcript_filename_srt

        # SRT tartalom generálása
        srt_content = create_srt_content(segments)
        # SRT fájl mentése
        async with aiofiles.open(transcript_path_srt, "w", encoding="utf-8") as f:
            await f.write(srt_content)
        logger.info(f"SRT transcript saved to: {transcript_path_srt}")

        # JSON adatok mentése
        transcript_data = {
            "text": transcript_text,
            "segments": segments,
            "live_translation_text": "",
            "original_filename": str(sanitized_name)
        }
        transcript_data_path = work_dir / "transcript_data.json"
        async with aiofiles.open(transcript_data_path, "w", encoding="utf-8") as f:
            await f.write(json.dumps(transcript_data, ensure_ascii=False, indent=2))

        await manager.send_progress(connection_id, 100, "Transcription complete")
        return JSONResponse({
            "transcription_text": transcript_text,
            "download_url": f"/download/{transcript_filename_srt}",
            "api_transcript_url": f"/download_transcript/{connection_id}/srt"
        })
    except HTTPException as e:
        logger.error(f"HTTP Error in transcribe_audio (ID: {connection_id}): {e.detail}")
        raise e
    except Exception as e:
        logger.error(f"Error in transcribe_audio (ID: {connection_id}): {str(e)}", exc_info=True)
        await manager.send_progress(connection_id, 100, f"Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")
    finally:
        logger.debug(f"Work directory {work_dir} kept for transcript download.")

async def transcribe_audio_with_format(file_path: Path, output_format: str, identify_speakers: bool = True, fast_mode: bool = False) -> tuple:
    """Transzkribálja az audiót és visszaadja a kimeneti útvonalat és tartalmat."""
    try:
        if not file_path.exists():
            raise FileNotFoundError(f"Audio file not found: {file_path}")
        transcript_text, transcript_data = await _transcribe_audio(file_path, identify_speakers=identify_speakers, fast_mode=fast_mode)
        segments = transcript_data.get("segments", [])
        output_ext = f".{output_format.lower()}"
        output_filename = file_path.with_suffix(output_ext).name
        output_path = file_path.parent / output_filename
        content = ""
        if output_format.lower() == 'srt':
            content = create_srt_content(segments)
        elif output_format.lower() == 'vtt':
            content = await generate_vtt(segments)
        else:  # TXT
            await generate_transcript_txt(segments, output_path, with_timestamps=True, with_speakers=identify_speakers and not fast_mode)
            async with aiofiles.open(output_path, 'r', encoding='utf-8') as f:
                content = await f.read()
            if output_format.lower() != 'txt':
                content = f"File generated at {output_path}"
        if output_format.lower() in ['srt', 'vtt']:
            async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                await f.write(content)
        return output_path, content
    except Exception as e:
        logger.error(f"Transcription format error for {file_path}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Transcription/Formatting failed: {str(e)}")

@router.post("/cli")
async def transcribe_audio_cli(
    file: UploadFile = File(...),
    output_format: str = Form("srt", description="Output format ('srt', 'txt', or 'vtt')."),
    identify_speakers: bool = Form(True), fast_mode: bool = Form(False)
):
    """CLI audio transzkripció, alapértelmezetten SRT kimenettel."""
    connection_id = str(uuid.uuid4())
    work_dir = TEMP_DIR / connection_id
    work_dir.mkdir(exist_ok=True)
    input_path = None
    supported_formats = ["srt", "txt", "vtt"]
    if output_format.lower() not in supported_formats:
        raise HTTPException(status_code=400, detail=f"Invalid format. Use: {', '.join(supported_formats)}.")
    try:
        sanitized_name = sanitize_filename(file.filename) if file.filename else Path(f"audio_{connection_id}.tmp")
        input_path = work_dir / sanitized_name
        # --- JAVÍTOTT FÁJL MENTÉS ---
        async with aiofiles.open(input_path, "wb") as f:
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                await f.write(chunk)
        # --- JAVÍTOTT FÁJL MENTÉS VÉGE ---
        logger.info(f"CLI request: Transcribing {input_path.name} to {output_format}, speakers={identify_speakers}, fast_mode={fast_mode}")
        output_path, _content = await transcribe_audio_with_format(input_path, output_format, identify_speakers=identify_speakers, fast_mode=fast_mode)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = Path(output_path.name).stem
        download_filename = f"{base_filename}_{timestamp}{output_path.suffix}"
        download_path = SYSTEM_DOWNLOADS / download_filename
        # --- JAVÍTOTT FÁJL MÁSOLÁS (Több sor, behúzás) ---
        async with aiofiles.open(output_path, 'rb') as src, aiofiles.open(download_path, 'wb') as dst:
            while True:
                chunk = await src.read(1024*1024)
                if not chunk:
                    break
                await dst.write(chunk)
        # --- JAVÍTOTT FÁJL MÁSOLÁS VÉGE ---
        logger.info(f"CLI result saved to: {download_path}")
        return JSONResponse({"download_url": f"/download/{download_filename}"})
    except HTTPException as e:
        logger.error(f"HTTP Error in CLI: {e.detail}")
        raise e
    except Exception as e:
        logger.error(f"Error in CLI: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")
    finally:
        if work_dir.exists():
            await async_rmtree(work_dir)
            logger.debug(f"Cleaned up CLI dir: {work_dir}")

@router.get("/download_transcript/{request_id}/{format}")
async def download_transcript(request_id: str, format: str):
    """Legenerálja és letölti a transzkriptumot a megadott formátumban."""
    work_dir = TEMP_DIR / request_id
    transcript_data_path = work_dir / "transcript_data.json"
    logger.info(f"Download request for ID {request_id}, format {format}.")
    if not transcript_data_path.is_file():
        logger.warning(f"Transcript data not found for {request_id}")
        raise HTTPException(status_code=404, detail=f"Transcript data not found for request ID {request_id}.")
    try:
        async with aiofiles.open(transcript_data_path, "r", encoding="utf-8") as f:
            transcript_data = json.loads(await f.read())
        logger.debug(f"Loaded transcript data for {request_id}.")
    except Exception as e:
        logger.error(f"Failed to read transcript data {request_id}: {e}")
        raise HTTPException(status_code=500, detail="Failed to load data.")
    segments = transcript_data.get("segments", [])
    full_text = transcript_data.get("text", "")
    original_filename = transcript_data.get("original_filename", f"transcript_{request_id}")
    base_filename = Path(sanitize_filename(original_filename)).stem
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"{base_filename}_{timestamp}.{format.lower()}"
    output_path = work_dir / output_filename
    logger.info(f"Generating {format} file: {output_path}")
    media_type = "application/octet-stream"
    try:
        if format.lower() == "json":
            media_type = 'application/json'
            output_path = transcript_data_path
        elif format.lower() == "txt":
            await generate_transcript_txt(segments, output_path, True, True)
            media_type = 'text/plain; charset=utf-8'
        elif format.lower() == "srt":
            content = create_srt_content(segments)
            async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
                await f.write(content)
            media_type = 'application/x-subrip'
        elif format.lower() == "vtt":
            content = await generate_vtt(segments)
            async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
                await f.write(content)
            media_type = 'text/vtt'
        elif format.lower() == "docx":
            await generate_transcript_docx(full_text, segments, output_path)
            media_type = 'app/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif format.lower() == "pdf":
            await generate_transcript_pdf(full_text, segments, output_path)
            media_type = 'application/pdf'
        else:
            logger.warning(f"Unsupported format: {format}")
            raise HTTPException(status_code=400, detail=f"Unsupported format: {format}.")
        if not output_path.is_file():
            logger.error(f"Generated file {output_path} not found")
            raise HTTPException(status_code=500, detail=f"Failed to generate {format} file.")
        logger.info(f"Serving file: {output_path} as {output_filename}")
        return FileResponse(path=str(output_path), media_type=media_type, filename=output_filename)
    except Exception as e:
        logger.exception(f"Error generating download ({format}) for {request_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate/serve {format} file.")

# Aliasok (subtitle_module-nak)
create_srt_content = create_srt_content
generate_vtt = generate_vtt
convert_txt_to_srt = convert_txt_to_srt