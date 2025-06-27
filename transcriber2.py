# transcriber.py - Javított verzió (fast_mode + convert_txt_to_srt)

import uuid
import json
import os
import asyncio
import logging
import re # Reguláris kifejezések importálása a convert_txt_to_srt-hez
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional

import aiofiles
from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from fastapi.responses import JSONResponse
from pydub import AudioSegment
# from pydub.utils import make_chunks # Nincs használva
from pydub.exceptions import CouldntDecodeError
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from openai import AsyncOpenAI

from config import (
    TEMP_DIR,
    SYSTEM_DOWNLOADS,
    manager,
    logger,
    client,
    sanitize_filename,
    async_rmtree
)

router = APIRouter()

# --- Segédfüggvény az időbélyeg formázásához ---
def format_timestamp(seconds: float) -> str:
    """Másodpercek átalakítása SRT időbélyeg formátumra (HH:MM:SS,mmm)."""
    if seconds < 0: seconds = 0
    total_milliseconds = int(seconds * 1000)
    hours, remainder = divmod(total_milliseconds, 3600000)
    minutes, remainder = divmod(remainder, 60000)
    secs, milliseconds = divmod(remainder, 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{milliseconds:03d}"

# --- Felirat Generátorok ---

# Ezt a függvényt használják a végpontok és a subtitle_module is
def create_srt_content(segments: List[Dict[str, Any]]) -> str:
    """SRT formátumú tartalom generálása szegmensekből.
       A beszélő jelölése itt kikerül, mert az inkább a TXT-be való.
    """
    srt_parts = []
    for i, segment in enumerate(segments, 1):
        start_time = format_timestamp(segment['start'])
        end_time = format_timestamp(segment['end'])
        text = segment.get('text', '').strip()
        if not text: continue

        srt_parts.append(f"{i}\n{start_time} --> {end_time}\n{text}\n")
    return "\n".join(srt_parts)

# Ezt a függvényt a subtitle_module importálja
async def generate_vtt(segments: List[Dict[str, Any]]) -> str:
    """VTT (WebVTT) formátumú tartalom generálása szegmensekből."""
    vtt_output = "WEBVTT\n\n"

    def format_vtt_timestamp(seconds: float) -> str:
        if seconds < 0: seconds = 0
        total_milliseconds = int(seconds * 1000)
        hours, remainder = divmod(total_milliseconds, 3600000)
        minutes, remainder = divmod(remainder, 60000)
        secs, milliseconds = divmod(remainder, 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{milliseconds:03d}"

    for segment in segments:
        start_time = format_vtt_timestamp(segment['start'])
        end_time = format_vtt_timestamp(segment['end'])
        text = segment.get('text', '').strip()
        if not text: continue

        vtt_output += f"{start_time} --> {end_time}\n{text}\n\n"
    return vtt_output

# --- HIÁNYZÓ FÜGGVÉNY HOZZÁADVA ---
def convert_txt_to_srt(txt_content: str) -> str:
    """
    Speciális TXT formátum ([start-end] text) konvertálása SRT formátumra.
    Ez az eredeti kódban szerepelt, és a subtitle_module használja.
    """
    lines = txt_content.strip().split('\n')
    srt_parts = []
    counter = 1

    for line in lines:
        line = line.strip()
        # Robusztusabb mintaillesztés a timestamp-re
        match = re.match(r'^\[(\d+\.?\d*)-(\d+\.?\d*)\]\s*(.*)', line)
        if not match:
            # logger.debug(f"Skipping line in convert_txt_to_srt (format mismatch): {line}")
            continue

        try:
            start_sec_str, end_sec_str, text = match.groups()
            start_sec = float(start_sec_str)
            end_sec = float(end_sec_str)

            if start_sec >= end_sec:
                logger.warning(f"Skipping line due to invalid timestamps (start >= end): {line}")
                continue

            start_time = format_timestamp(start_sec)
            end_time = format_timestamp(end_sec)

            text_content = text.strip()
            if not text_content: continue

            srt_parts.append(f"{counter}\n{start_time} --> {end_time}\n{text_content}\n")
            counter += 1
        except ValueError as e:
            logger.warning(f"Skipping line due to float conversion error: {line}, Error: {e}")
            continue
        except Exception as e:
            logger.error(f"Error converting line to SRT: {line}, Error: {e}")
            continue

    return "\n".join(srt_parts)
# --- HIÁNYZÓ FÜGGVÉNY VÉGE ---


# --- Jegyzőkönyv Generátorok ---

async def generate_transcript_docx(text: str, segments: List[Dict[str, Any]], output_path: Path):
    """DOCX transzkript generálása"""
    doc = Document()
    try:
        styles = doc.styles['Normal']
        font = styles.font
    except KeyError:
        logger.warning("Default 'Normal' style not found in DOCX template.")

    if segments:
        for segment in segments:
            text_content = segment.get('text', '').strip()
            if not text_content: continue

            speaker = segment.get('speaker', '')
            speaker_text = f"[Beszélő {speaker}]: " if speaker else ""
            timestamp = f"[{segment['start']:.2f}-{segment['end']:.2f}]"
            p = doc.add_paragraph()
            p.add_run(f"{timestamp} ").bold = True
            if speaker:
                p.add_run(f"Beszélő {speaker}: ").italic = True
            p.add_run(text_content)
    elif text:
        doc.add_paragraph(text)
    else:
        doc.add_paragraph("No content available.")

    loop = asyncio.get_running_loop()
    await loop.run_in_executor(None, doc.save, str(output_path))

async def generate_transcript_pdf(text: str, segments: List[Dict[str, Any]], output_path: Path):
    """PDF transzkript generálása"""
    doc = SimpleDocTemplate(str(output_path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    if segments:
        for segment in segments:
            text_content = segment.get('text', '').strip()
            if not text_content: continue

            speaker = segment.get('speaker', '')
            speaker_info = f"Beszélő {speaker}: " if speaker else ""
            timestamp = f"[{segment['start']:.2f}-{segment['end']:.2f}]"
            from reportlab.lib.utils import escapeOnce
            line_text = f"<b>{escapeOnce(timestamp)}</b> <i>{escapeOnce(speaker_info)}</i>{escapeOnce(text_content)}"
            story.append(Paragraph(line_text, styles['Normal']))
            story.append(Spacer(1, 6))
    elif text:
        from reportlab.lib.utils import escapeOnce
        story.append(Paragraph(escapeOnce(text), styles['Normal']))
    else:
        story.append(Paragraph("No content available.", styles['Normal']))

    loop = asyncio.get_running_loop()
    await loop.run_in_executor(None, doc.build, story)

async def generate_transcript_txt(segments: List[Dict[str, Any]], output_path: Path, with_timestamps: bool = True, with_speakers: bool = True):
    """TXT jegyzőkönyv generálása időbélyegekkel és/vagy beszélőkkel."""
    async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
        current_speaker = None
        first_line = True
        for segment in segments:
            text = segment.get('text', '').strip()
            if not text: continue

            speaker = segment.get('speaker', '') if with_speakers else ''
            timestamp_text = f"[{segment['start']:.2f}-{segment['end']:.2f}] " if with_timestamps else ""

            line_prefix = timestamp_text
            speaker_change_line = ""

            if speaker and speaker != current_speaker:
                if not first_line: speaker_change_line += "\n"
                speaker_change_line += f"{timestamp_text}Beszélő {speaker}:\n"
                current_speaker = speaker
                line_prefix = ""
            elif speaker and speaker == current_speaker:
                 line_prefix = timestamp_text
            elif not speaker:
                 line_prefix = timestamp_text
                 if current_speaker is not None and not first_line: speaker_change_line += "\n"
                 current_speaker = None

            if speaker_change_line:
                 await f.write(speaker_change_line)

            await f.write(f"{line_prefix}{text}\n")
            first_line = False

async def post_process_transcript_with_styles(segments, output_path):
    """Továbbfejlesztett jegyzőkönyv generálás ANSI színekkel a konzolhoz."""
    speaker_colors = { "1": "\033[94m", "2": "\033[92m", "3": "\033[91m", "4": "\033[93m", "5": "\033[95m", "6": "\033[96m" }
    reset_color = "\033[0m"
    default_color = "\033[97m"

    async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
        current_speaker = None
        first_line = True
        for segment in segments:
            text_content = segment.get('text', '').strip()
            if not text_content: continue

            speaker = str(segment.get('speaker', ''))
            timestamp_text = f"[{segment['start']:.2f}-{segment['end']:.2f}] "
            color_code = speaker_colors.get(speaker, default_color)

            line_prefix = timestamp_text
            speaker_change_line = ""

            if speaker and speaker != current_speaker:
                if not first_line: speaker_change_line += "\n"
                speaker_change_line += f"{timestamp_text}{color_code}Beszélő {speaker}{reset_color}:\n"
                current_speaker = speaker
                line_prefix = ""
            elif speaker and speaker == current_speaker:
                 line_prefix = timestamp_text
            elif not speaker:
                 line_prefix = timestamp_text
                 if current_speaker is not None and not first_line: speaker_change_line += "\n"
                 current_speaker = None

            if speaker_change_line:
                 await f.write(speaker_change_line)

            await f.write(f"{line_prefix}{color_code}{text_content}{reset_color}\n")
            first_line = False

# --- Core Transzkripciós Logika (fast_mode támogatással) ---

async def _transcribe_audio(audio_path: Path, identify_speakers: bool = False, fast_mode: bool = False) -> tuple:
    """Transzkribál egy audió fájlt (Whisper) és opcionálisan azonosítja a beszélőket (GPT).
       Fast mode esetén kihagyja a beszélőazonosítást.
    """
    if client is None:
        logger.error("OpenAI client not available. Cannot transcribe.")
        raise HTTPException(status_code=503, detail="Transcription service unavailable (OpenAI client not configured).")

    loop = asyncio.get_running_loop()
    audio_file_to_transcribe = audio_path
    temp_mp3_path = None

    try:
        logger.debug(f"Loading audio file: {audio_path}")
        audio = await loop.run_in_executor(None, AudioSegment.from_file, str(audio_path))
        logger.debug(f"Audio loaded. Duration: {audio.duration_seconds:.2f}s")

        if audio.duration_seconds < 0.1:
            logger.warning(f"Audio file {audio_path.name} is too short: {audio.duration_seconds:.3f}s")
            raise HTTPException(status_code=400, detail="Audio file is too short (minimum 0.1 seconds).")

        # Konvertálás MP3-ba, ha szükséges (Whisper jobban szereti)
        if audio_path.suffix.lower() not in [".mp3", ".wav", ".m4a", ".ogg", ".flac", ".opus"]:
            temp_mp3_path = audio_path.with_suffix(".mp3")
            logger.info(f"Converting {audio_path.name} to MP3 for transcription...")
            await loop.run_in_executor(None, audio.export, str(temp_mp3_path), format="mp3", bitrate="64k")
            audio_file_to_transcribe = temp_mp3_path
            logger.info(f"Converted to {temp_mp3_path.name}")

        # Fájlméret ellenőrzése
        file_size_mb = audio_file_to_transcribe.stat().st_size / (1024 * 1024)
        if file_size_mb > 25: # OpenAI limit
             logger.error(f"Audio file {audio_file_to_transcribe.name} exceeds 25MB limit ({file_size_mb:.2f}MB).")
             raise HTTPException(status_code=413, detail="Audio file size exceeds the 25MB limit for transcription.")
        logger.debug(f"Audio file size for transcription: {file_size_mb:.2f}MB")

        # --- Whisper Transzkripció ---
        transcript = None
        logger.info(f"Starting transcription for {audio_file_to_transcribe.name}...")
        with open(audio_file_to_transcribe, "rb") as f:
            transcript = await client.audio.transcriptions.create(
                model="whisper-1",
                file=f,
                response_format="verbose_json",
                timestamp_granularities=["segment"]
            )

        if not transcript or not hasattr(transcript, 'text') or not hasattr(transcript, 'segments'):
            raise ValueError("Invalid response received from Whisper API.")

        logger.info(f"Transcription finished for {audio_file_to_transcribe.name}.")
        full_transcript_text = transcript.text if transcript.text else ""
        raw_segments = [segment.model_dump() for segment in transcript.segments] if transcript.segments else []

    except CouldntDecodeError:
        logger.error(f"Could not decode audio file: {audio_path}")
        raise HTTPException(status_code=400, detail=f"Unsupported audio format or corrupt file: {audio_path.name}")
    except FileNotFoundError:
        logger.error(f"Audio file not found for transcription: {audio_path}")
        raise HTTPException(status_code=404, detail=f"Audio file not found: {audio_path.name}")
    except HTTPException as http_exc:
         raise http_exc # Továbbadjuk a méretkorlát hibát
    except Exception as e:
        logger.error(f"Error during audio loading or transcription: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Transcription failed: {str(e)}")
    finally:
        # Ideiglenes MP3 törlése
        if temp_mp3_path and temp_mp3_path.exists():
            try:
                logger.debug(f"Removing temporary MP3 file: {temp_mp3_path}")
                temp_mp3_path.unlink()
            except OSError as e:
                logger.warning(f"Could not remove temporary MP3 file {temp_mp3_path}: {e}")

    # --- Beszélőazonosítás ---
    effective_identify_speakers = identify_speakers and not fast_mode
    final_segments = []

    if effective_identify_speakers and raw_segments:
        logger.info(f"Starting speaker identification for {audio_path.name}...")
        segments_for_analysis = [s.copy() for s in raw_segments]

        if len(segments_for_analysis) > 1:
            speaker_system_prompt = """
            Azonosítsd a különböző beszélőket a következő átiratban.
            Minden szegmenst jelölj meg egy beszélővel (1, 2, 3, stb.).
            Törekedj a konzisztenciára - ugyanazon beszélő mindig ugyanazt a számot kapja.
            A válaszod KIZÁRÓLAG egy JSON objektum legyen, amely tartalmaz egy "segments" kulcsot.
            Ennek értéke egy lista legyen, ahol minden elem egy eredeti szegmensnek megfelelő objektum,
            kiegészítve egy "speaker": "ID" párral. Például:
            { "segments": [ { "start": 0.5, "end": 2.1, "text": "Hello.", "speaker": "1" }, ... ] }
            Ne adj hozzá semmilyen magyarázatot a JSON-on kívül. Csak a JSON objektumot add vissza.
            """
            user_content_json = json.dumps([{"start": s['start'], "end": s['end'], "text": s['text']} for s in segments_for_analysis], ensure_ascii=False)

            try:
                speaker_response = await client.chat.completions.create(
                    model=os.getenv("MODEL_NAME", "gpt-3.5-turbo"),
                    messages=[
                        {"role": "system", "content": speaker_system_prompt},
                        {"role": "user", "content": f"A következő átirat szegmenseit elemezd és azonosítsd a beszélőket: {user_content_json}"}
                    ],
                    response_format={"type": "json_object"},
                    temperature=0.3
                )
                response_content = speaker_response.choices[0].message.content
                if not response_content: raise ValueError("Empty response from speaker ID API.")
                logger.debug(f"Speaker ID raw response: {response_content}")
                speaker_analysis = json.loads(response_content)

                # Eredmények integrálása
                processed_segments_map = {}
                if "segments" in speaker_analysis and isinstance(speaker_analysis["segments"], list) and len(speaker_analysis["segments"]) == len(segments_for_analysis):
                    for i, resp_seg in enumerate(speaker_analysis["segments"]):
                        original_seg = segments_for_analysis[i]
                        if isinstance(resp_seg, dict) and "speaker" in resp_seg:
                            speaker_id = str(resp_seg.get("speaker", "1")).strip() or "1"
                            original_seg["speaker"] = speaker_id
                        else:
                            logger.warning(f"Unexpected format in speaker segment {i}, assigning default '1'. Response segment: {resp_seg}")
                            original_seg["speaker"] = "1"
                        processed_segments_map[original_seg['start']] = original_seg
                    # Visszaillesztés
                    for seg in raw_segments:
                        seg['speaker'] = processed_segments_map.get(seg['start'], {}).get('speaker', '1')
                    final_segments = raw_segments
                    logger.info("Speaker identification results merged successfully.")
                else:
                    logger.warning(f"Speaker ID response format mismatch or count incorrect. Assigning default speakers.")
                    for segment in raw_segments: segment["speaker"] = "1"
                    final_segments = raw_segments

            except Exception as e:
                logger.error(f"Error during speaker identification API call or processing: {str(e)}", exc_info=True)
                for segment in raw_segments: segment["speaker"] = ""
                final_segments = raw_segments
        else: # Túl kevés szegmens
            logger.info("Not enough segments for speaker identification.")
            for segment in raw_segments: segment["speaker"] = "1"
            final_segments = raw_segments

    else: # Beszélőazonosítás kihagyva
        if fast_mode: logger.info(f"Fast mode enabled, skipping speaker identification for {audio_path.name}.")
        else: logger.info(f"Speaker identification disabled for {audio_path.name}.")
        for segment in raw_segments: segment["speaker"] = ""
        final_segments = raw_segments

    final_segments = sorted(final_segments, key=lambda x: x.get("start", 0))
    return full_transcript_text, {"type": "audio", "segments": final_segments}


# --- Endpointok ---

@router.post("/")
async def transcribe_audio(
    file: UploadFile = File(...),
    live_translation: bool = Form(False),
    target_lang: str = Form("en"),
    identify_speakers: bool = Form(True),
    fast_mode: bool = Form(False),
    connection_id: Optional[str] = Form(None)
):
    """Audio fájl transzkripciója beszélő-felismeréssel és fast_mode opcióval."""
    if not connection_id: connection_id = str(uuid.uuid4())
    work_dir = TEMP_DIR / connection_id
    work_dir.mkdir(exist_ok=True)
    input_path = None

    try:
        sanitized_name = sanitize_filename(file.filename) if file.filename else Path(f"audio_{connection_id}.tmp")
        input_path = work_dir / sanitized_name
        async with aiofiles.open(input_path, "wb") as f:
            content = await file.read()
            await f.write(content)
            file_size = len(content)
        logger.info(f"File saved to {input_path} ({file_size} bytes)")

        await manager.send_progress(connection_id, 10, "Processing audio file...")

        transcript_text, format_info = await _transcribe_audio(
            input_path, identify_speakers=identify_speakers, fast_mode=fast_mode
        )
        segments = format_info.get("segments", [])

        transcript_data = {
            "text": transcript_text, "segments": segments,
            "live_translation_text": "", "original_filename": str(sanitized_name)
        }
        transcript_data_path = work_dir / "transcript_data.json"
        async with aiofiles.open(transcript_data_path, "w", encoding="utf-8") as f:
            await f.write(json.dumps(transcript_data, ensure_ascii=False, indent=2))

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = input_path.stem
        transcript_filename_txt = f"transcript_{base_filename}_{timestamp}.txt"
        transcript_path_txt = SYSTEM_DOWNLOADS / transcript_filename_txt

        await generate_transcript_txt(
             segments, transcript_path_txt, with_timestamps=True,
             with_speakers=identify_speakers and not fast_mode
        )

        await manager.send_progress(connection_id, 100, "Transcription complete")

        return JSONResponse({
            "transcription_text": transcript_text,
            "download_url": f"/download/{transcript_filename_txt}",
            "api_transcript_url": f"/download_transcript/{connection_id}/txt"
        })

    except HTTPException as e:
         logger.error(f"HTTP Error in transcribe_audio (ID: {connection_id}): {e.detail}", exc_info=False) # No need for full traceback for HTTP Exceptions
         raise e
    except Exception as e:
        logger.error(f"Error in transcribe_audio (ID: {connection_id}): {str(e)}", exc_info=True)
        await manager.send_progress(connection_id, 100, f"Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred: {str(e)}")
    finally:
        logger.debug(f"Work directory {work_dir} kept for potential transcript download.")


async def transcribe_audio_with_format(file_path: Path, output_format: str, identify_speakers: bool = True, fast_mode: bool = False) -> tuple:
    """Transzkribálja az audiót és visszaadja a kimeneti útvonalat és tartalmat (fast_mode figyelembevételével)."""
    try:
        if not file_path.exists():
            raise FileNotFoundError(f"Audio file not found: {file_path}")

        transcript_text, transcript_data = await _transcribe_audio(
            file_path, identify_speakers=identify_speakers, fast_mode=fast_mode
        )
        segments = transcript_data.get("segments", [])

        output_ext = f".{output_format.lower()}"
        output_filename = file_path.with_suffix(output_ext).name
        output_path = file_path.parent / output_filename

        content = ""
        if output_format.lower() == 'srt':
            content = create_srt_content(segments)
        elif output_format.lower() == 'vtt':
             content = await generate_vtt(segments)
        else:
            await generate_transcript_txt(
                segments, output_path, with_timestamps=True,
                with_speakers=identify_speakers and not fast_mode
            )
            async with aiofiles.open(output_path, 'r', encoding='utf-8') as f:
                content = await f.read()
            if output_format.lower() != 'txt':
                 content = f"File generated at {output_path}"

        if output_format.lower() in ['srt', 'vtt']:
            async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                await f.write(content)

        return output_path, content

    except Exception as e:
        logger.error(f"Transcription with format error for {file_path}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Transcription/Formatting failed: {str(e)}")


@router.post("/cli")
async def transcribe_audio_cli(
    file: UploadFile = File(...),
    output_format: str = Form("txt"),
    identify_speakers: bool = Form(True),
    fast_mode: bool = Form(False)
):
    """CLI audio transzkripció beszélő-felismeréssel és fast_mode opcióval."""
    connection_id = str(uuid.uuid4())
    work_dir = TEMP_DIR / connection_id
    work_dir.mkdir(exist_ok=True)
    input_path = None

    try:
        sanitized_name = sanitize_filename(file.filename) if file.filename else Path(f"audio_{connection_id}.tmp")
        input_path = work_dir / sanitized_name
        async with aiofiles.open(input_path, "wb") as f:
            await f.write(await file.read())

        logger.info(f"CLI request: Transcribing {input_path.name} to {output_format}, speakers={identify_speakers}, fast_mode={fast_mode}")

        output_path, _content = await transcribe_audio_with_format(
            input_path, output_format, identify_speakers=identify_speakers, fast_mode=fast_mode
        )

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = Path(output_path.name).stem
        download_filename = f"{base_filename}_{timestamp}{output_path.suffix}"
        download_path = SYSTEM_DOWNLOADS / download_filename

        async with aiofiles.open(output_path, 'rb') as src, aiofiles.open(download_path, 'wb') as dst:
            while True:
                chunk = await src.read(1024 * 1024)
                if not chunk: break
                await dst.write(chunk)
        logger.info(f"CLI result saved to: {download_path}")

        return JSONResponse({"download_url": f"/download/{download_filename}"})

    except HTTPException as e:
         logger.error(f"HTTP Error in transcribe_audio_cli (ID: {connection_id}): {e.detail}", exc_info=False)
         raise e
    except Exception as e:
        logger.error(f"Error in transcribe_audio_cli (ID: {connection_id}): {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred: {str(e)}")
    finally:
        if work_dir.exists():
             await async_rmtree(work_dir)
             logger.debug(f"Cleaned up CLI work directory: {work_dir}")


# --- Aliasok a visszamenőleges kompatibilitáshoz ---
# Hogy a subtitle_module is megtalálja őket
# create_srt_content = create_srt_content # Már létezik
# generate_vtt = generate_vtt             # Már létezik
# convert_txt_to_srt = convert_txt_to_srt # Már létezik
# -------------------------------------------------
