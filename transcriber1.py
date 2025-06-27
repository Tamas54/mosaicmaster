import uuid
import json
import os
import asyncio
import logging
from datetime import datetime, timedelta
from pathlib import Path

import aiofiles
from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from fastapi.responses import JSONResponse
from pydub import AudioSegment
from pydub.utils import make_chunks
from pydub.exceptions import CouldntDecodeError
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

from config import (
    TEMP_DIR,
    SYSTEM_DOWNLOADS,
    manager,
    logger,
    client,
    sanitize_filename,
    async_rmtree
)

router = APIRouter()


async def generate_srt(segments: list) -> str:
    """SRT felirat generálása a szegmensekből"""
    srt_output = ""
    for i, segment in enumerate(segments, 1):
        start_time = format_timestamp(segment['start'])
        end_time = format_timestamp(segment['end'])
        srt_output += f"{i}\n{start_time} --> {end_time}\n{segment['text']}\n\n"
    return srt_output

async def generate_vtt(segments: list) -> str:
    """VTT felirat generálása a szegmensekből"""
    vtt_output = "WEBVTT\n\n"
    for segment in segments:
        start_time = "0" + str(timedelta(seconds=segment['start']))
        end_time = "0" + str(timedelta(seconds=segment['end']))
        vtt_output += f"{start_time} --> {end_time}\n{segment['text']}\n\n"
    return vtt_output

async def generate_transcript_docx(text: str, segments: list, output_path: Path):
    """DOCX transzkript generálása"""
    doc = Document()
    if segments:
        for segment in segments:
            speaker = segment.get('speaker', '')
            speaker_text = f"[Beszélő {speaker}]: " if speaker else ""
            doc.add_paragraph(f"[{segment['start']:.2f}-{segment['end']:.2f}] {speaker_text}{segment['text']}")
    else:
        doc.add_paragraph(text)
    doc.save(str(output_path))

async def generate_transcript_pdf(text: str, segments: list, output_path: Path):
    """PDF transzkript generálása"""
    doc = SimpleDocTemplate(str(output_path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    if segments:
        for segment in segments:
            speaker = segment.get('speaker', '')
            speaker_text = f"[Beszélő {speaker}]: " if speaker else ""
            story.append(Paragraph(f"[{segment['start']:.2f}-{segment['end']:.2f}] {speaker_text}{segment['text']}", styles['Normal']))
            story.append(Spacer(1, 12))
    else:
        story.append(Paragraph(text, styles['Normal']))
    doc.build(story)

async def generate_transcript_txt(segments: list, output_path: Path, with_timestamps: bool = True):
    """TXT jegyzőkönyv generálása beszélők megkülönböztetésével"""
    async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
        current_speaker = None
        for segment in segments:
            speaker = segment.get('speaker', '')
            timestamp_text = f"[{segment['start']:.2f}-{segment['end']:.2f}] " if with_timestamps else ""
            
            # Ha új beszélő, akkor jelezzük
            if speaker != current_speaker:
                await f.write(f"\n{timestamp_text}Beszélő {speaker}:\n")
                current_speaker = speaker
            else:
                # Ha ugyanaz a beszélő folytatja, csak új sorban
                await f.write(f"{timestamp_text}")
            
            # A szöveg kiírása
            await f.write(f"{segment['text']}\n")

async def post_process_transcript_with_styles(segments, output_path):
    """Továbbfejlesztett jegyzőkönyv generálás különböző stílusokkal a beszélőkhöz"""
    
    # Beszélő-színek definiálása
    speaker_colors = {
        "1": "\033[94m",  # Kék
        "2": "\033[92m",  # Zöld
        "3": "\033[91m",  # Piros
        "4": "\033[93m",  # Sárga
        "5": "\033[95m",  # Magenta
        "6": "\033[96m",  # Cyan
        "7": "\033[97m",  # Fehér
        "8": "\033[90m",  # Szürke
    }
    reset_color = "\033[0m"
    
    async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
        current_speaker = None
        for segment in segments:
            speaker = segment.get('speaker', '')
            timestamp_text = f"[{segment['start']:.2f}-{segment['end']:.2f}] "
            
            # A megfelelő szín kiválasztása (vagy alapértelmezett, ha nincs)
            color_code = speaker_colors.get(speaker, "")
            
            # Ha új beszélő, akkor jelezzük
            if speaker != current_speaker:
                await f.write(f"\n{timestamp_text}{color_code}Beszélő {speaker}{reset_color}:\n")
                current_speaker = speaker
            else:
                # Ha ugyanaz a beszélő folytatja, csak új sorban
                await f.write(f"{timestamp_text}")
            
            # A szöveg kiírása
            await f.write(f"{color_code}{segment['text']}{reset_color}\n")

# Audio transzkripció segédfüggvénye beszélő-felismeréssel - JAVÍTOTT
async def _transcribe_audio(audio_path: Path, identify_speakers: bool = False, fast_mode: bool = False) -> tuple:
    """Transzkribál egy audió fájlt és opcionálisan megkülönbözteti a beszélőket.
    
    Args:
        audio_path: Az audio fájl elérési útja
        identify_speakers: True, ha beszélőket kell azonosítani
        fast_mode: True, ha gyors feldolgozási módot kell használni (nem élő stream esetén)
    """
    try:
        audio = AudioSegment.from_file(str(audio_path)) # Bármilyen formátum
    except CouldntDecodeError:
        raise HTTPException(status_code=400, detail="Unsupported audio format. Could not decode.")

    # Ellenőrizzük, hogy a hangfájl hossza legalább 0.1 másodperc-e
    if audio.duration_seconds < 0.1:
        raise HTTPException(status_code=400, detail="Audio file is too short. Minimum audio length is 0.1 seconds.")
    
    # Ne használjunk chunkolást, az egész hangfájlt egyszerre küldjük
    full_transcript = ""
    all_segments = []
    speaker_segments = []
    
    # Az egész fájlt használjuk
    audio_file_path = audio_path
    
    # Ha szükséges, konvertáljuk MP3-ba a teljes fájlt
    if audio_path.suffix.lower() != ".mp3":
        audio_file_path = audio_path.parent / f"{audio_path.stem}.mp3"  # MP3 kiterjesztés
        audio.export(str(audio_file_path), format="mp3", bitrate="64k")  # MP3-ba, 64 kbps

    with open(audio_file_path, "rb") as f:  # SZINKRON olvasás
        # Alap transzkripció az OpenAI Whisper API-val
        transcript = await client.audio.transcriptions.create(
            model="whisper-1",
            file=f,
            response_format="verbose_json"
        )

    full_transcript = transcript.text + "\n"
    
    # Ha van szegmens és kértek beszélő-azonosítást
    if hasattr(transcript, 'segments') and identify_speakers:
        # OpenAI API beszélő-felismerés implementációja
        # Kéri a ChatGPT-t, hogy azonosítsa a beszélőket a szegmensekben
        segments_for_analysis = []
        for segment in transcript.segments:
            segment_dict = {
                "start": segment.start,
                "end": segment.end,
                "text": segment.text
            }
            segments_for_analysis.append(segment_dict)
            
        # Csak akkor kérünk beszélő-azonosítást, ha van elég szegmens
        if len(segments_for_analysis) > 1:
            speaker_system_prompt = """
            Azonosítsd a különböző beszélőket a következő átiratban. 
            Minden szegmenst jelölj meg egy beszélővel (1, 2, 3, stb.). 
            Törekedj a konzisztenciára - ugyanazon beszélő mindig ugyanazt a számot kapja.
            Csak a válaszod legyen a JSON formátumú beszélőazonosítás.
            """
            
            try:
                speaker_response = await client.chat.completions.create(
                    model=os.getenv("MODEL_NAME", "gpt-3.5-turbo"),
                    messages=[
                        {"role": "system", "content": speaker_system_prompt},
                        {"role": "user", "content": f"A következő átirat szegmenseit elemezd és azonosítsd a beszélőket: {json.dumps(segments_for_analysis)}"}
                    ],
                    response_format={"type": "json_object"}
                )
                
                speaker_analysis = json.loads(speaker_response.choices[0].message.content)
                
                # Beszélőazonosítás eredményeinek integrálása
                if "segments" in speaker_analysis and isinstance(speaker_analysis["segments"], list):
                    for i, segment in enumerate(segments_for_analysis):
                        if i < len(speaker_analysis["segments"]):
                            # Ellenőrizzük, hogy a speaker_analysis["segments"][i] egy dict-e
                            segment_data = speaker_analysis["segments"][i]
                            if isinstance(segment_data, dict):
                                segment["speaker"] = segment_data.get("speaker", "")
                            elif isinstance(segment_data, list):
                                # Ha lista, akkor az első elemét próbáljuk meg használni
                                segment["speaker"] = "1"
                            else:
                                # Egyéb esetben egyszerűen stringként kezeljük
                                segment["speaker"] = str(segment_data) if segment_data else "1"
                        else:
                            segment["speaker"] = "1"  # Alapértelmezett beszélő
                
                # Ha nem a várt formátumban kaptuk az eredményt, megpróbáljuk értelmezni
                elif isinstance(speaker_analysis, dict):
                    for i, segment in enumerate(segments_for_analysis):
                        segment_id = str(i)
                        if segment_id in speaker_analysis:
                            segment_data = speaker_analysis[segment_id]
                            if isinstance(segment_data, dict):
                                segment["speaker"] = segment_data.get("speaker", "")
                            else:
                                segment["speaker"] = str(segment_data) if segment_data else "1"
                        else:
                            segment["speaker"] = "1"  # Alapértelmezett beszélő
                
            except Exception as e:
                logger.error(f"Error during speaker identification: {str(e)}")
                # Hiba esetén speaker azonosítás nélkül folytatjuk
                for segment in segments_for_analysis:
                    segment["speaker"] = ""
            
            # Hozzáadjuk az eredményeket a teljes listához
            speaker_segments.extend(segments_for_analysis)
        else:
            # Ha túl kevés szegmens van az elemzéshez
            for segment in segments_for_analysis:
                segment["speaker"] = "1"  # Alapértelmezetten 1-es beszélő
            speaker_segments.extend(segments_for_analysis)
            
    elif hasattr(transcript, 'segments'):
        # Ha nincs beszélőazonosítás, egyszerűen hozzáadjuk a szegmenseket
        for segment in transcript.segments:
            segment_dict = {
                "start": segment.start,
                "end": segment.end,
                "text": segment.text
            }
            all_segments.append(segment_dict)

    # Ha van beszélőazonosítás, akkor azokat a szegmenseket használjuk
    final_segments = speaker_segments if speaker_segments else all_segments
    
    # Rendezzük a szegmenseket az idő szerint
    final_segments = sorted(final_segments, key=lambda x: x["start"])

    return full_transcript, {"type": "audio", "segments": final_segments}


def format_timestamp(seconds):
    """Convert seconds to SRT timestamp format (HH:MM:SS,mmm)"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = seconds % 60
    milliseconds = int((secs % 1) * 1000)
    secs = int(secs)

    return f"{hours:02d}:{minutes:02d}:{secs:02d},{milliseconds:03d}"

def create_srt_content(segments):
    """Create SRT formatted content from transcript segments"""
    srt_parts = []
    for i, segment in enumerate(segments, 1):
        start_time = format_timestamp(segment['start'])
        end_time = format_timestamp(segment['end'])
        speaker = segment.get('speaker', '')
        text = segment['text'].strip()
        
        # Ha van beszélő, akkor azt is megjelenítjük
        if speaker:
            text = f"{text}"  # A beszélő jelölését csak a felhasználói felületen jelenítjük meg

        srt_parts.append(
            f"{i}\n"
            f"{start_time} --> {end_time}\n"
            f"{text}\n"
        )

    return "\n".join(srt_parts)

def convert_txt_to_srt(txt_content):
    """Converts timestamp format [start-end] text to SRT format
    
    Example input:
    [6.88-7.88]  in America.
    [7.88-13.44]  Not all of these still exist, but there are at least a dozen Romes still scattered around
    
    Example output:
    1
    00:00:06,880 --> 00:00:07,880
    in America.
    
    2
    00:00:07,880 --> 00:00:13,440
    Not all of these still exist, but there are at least a dozen Romes still scattered around
    """
    lines = txt_content.strip().split('\n')
    srt_parts = []
    counter = 1
    
    for line in lines:
        line = line.strip()
        if not line or not line.startswith('['):
            continue
            
        try:
            # Extract timestamps and text
            timestamp_part, text = line.split(']', 1)
            timestamp_part = timestamp_part[1:]  # Remove the opening bracket
            start_sec, end_sec = map(float, timestamp_part.split('-'))
            
            # Convert to SRT format
            start_time = format_timestamp(start_sec)
            end_time = format_timestamp(end_sec)
            
            srt_parts.append(
                f"{counter}\n"
                f"{start_time} --> {end_time}\n"
                f"{text.strip()}\n"
            )
            counter += 1
        except Exception as e:
            logger.error(f"Error converting line to SRT: {line}, Error: {e}")
            continue
    
    return "\n".join(srt_parts)

async def transcribe_audio_with_format(file_path: Path, output_format: str, identify_speakers: bool = True) -> tuple:
    """Transzkribálja az audiót és visszaadja a kimeneti útvonalat és tartalmat"""
    try:
        if not file_path.exists():
            raise FileNotFoundError(f"Audio file not found: {file_path}")

        transcript_text, transcript_data = await _transcribe_audio(
            file_path, 
            identify_speakers=identify_speakers
        )

        output_ext = '.srt' if output_format == 'srt' else '.txt'
        output_path = file_path.with_suffix(output_ext)

        if output_format == 'srt':
            content = create_srt_content(transcript_data['segments'])
            async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                await f.write(content)
        else:
            # TXT jegyzőkönyv generálása
            await generate_transcript_txt(
                transcript_data['segments'], 
                output_path,
                with_timestamps=True
            )
            async with aiofiles.open(output_path, 'r', encoding='utf-8') as f:
                content = await f.read()

        return output_path, content
    except Exception as e:
        logger.error(f"Transcription error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/")
async def transcribe_audio(
    file: UploadFile = File(...),
    live_translation: bool = Form(False),
    target_lang: str = Form("en"),
    identify_speakers: bool = Form(True),  # Beszélő-felismerés kapcsoló
    connection_id: str = Form(None)
):
    """Audio fájl transzkripciója beszélő-felismeréssel"""
    if not connection_id:
        connection_id = str(uuid.uuid4())
    work_dir = TEMP_DIR / connection_id
    work_dir.mkdir(exist_ok=True)

    try:
        input_path = work_dir / sanitize_filename(file.filename)
        async with aiofiles.open(input_path, "wb") as f:
            await f.write(await file.read())

        await manager.send_progress(connection_id, 10, "Processing audio file...")

        transcript_text, format_info = await _transcribe_audio(
            input_path, 
            identify_speakers=identify_speakers
        )
        segments = format_info.get("segments", [])

        # 1. Transcript data (JSON) mentése a WORK_DIR-be (a /download_transcript végpontnak)
        transcript_data = {
            "text": transcript_text,
            "segments": segments,
            "live_translation_text": "",  # Üres, mert most nincs live translation
            "original_filename": file.filename
        }
        transcript_data_path = work_dir / "transcript_data.json"
        async with aiofiles.open(transcript_data_path, "w", encoding="utf-8") as f:
            await f.write(json.dumps(transcript_data, ensure_ascii=False))

        # 2. TXT jegyzőkönyv generálása a SYSTEM_DOWNLOADS-ba
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        transcript_filename = f"transcript_{sanitize_filename(file.filename).stem}_{timestamp}.txt"
        transcript_path = SYSTEM_DOWNLOADS / transcript_filename
        
        # Jegyzőkönyv generálása TXT formátumban
        await generate_transcript_txt(segments, transcript_path)

        await manager.send_progress(connection_id, 100, "Transcription with speaker identification complete")

        return JSONResponse({
            "transcription_text": transcript_text,
            "download_url": f"/download/{transcript_filename}",
            "api_transcript_url": f"/download_transcript/{connection_id}/txt"
        })

    except Exception as e:
        logger.error(f"Error in transcribe_audio: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        # Ne töröljük a munkakönyvtárat, mert a transcript_data.json-t még használni fogjuk
        pass

@router.post("/cli")
async def transcribe_audio_cli(
    file: UploadFile = File(...),
    output_format: str = Form("txt"),
    identify_speakers: bool = Form(True)  # Beszélő-felismerés kapcsoló
):
    """CLI audio transzkripció beszélő-felismeréssel"""
    connection_id = str(uuid.uuid4())
    work_dir = TEMP_DIR / connection_id
    work_dir.mkdir(exist_ok=True)

    try:
        input_path = work_dir / sanitize_filename(file.filename)
        async with aiofiles.open(input_path, "wb") as f:
            await f.write(await file.read())

        output_path, content = await transcribe_audio_with_format(
            input_path, 
            output_format,
            identify_speakers=identify_speakers
        )
        
        # Másoljuk a generált fájlt a SYSTEM_DOWNLOADS könyvtárba
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        download_filename = f"transcript_{input_path.stem}_{timestamp}{output_path.suffix}"
        download_path = SYSTEM_DOWNLOADS / download_filename
        
        # Másolás
        async with aiofiles.open(output_path, 'rb') as src, aiofiles.open(download_path, 'wb') as dst:
            content = await src.read()
            await dst.write(content)

        return JSONResponse({
            "download_url": f"/download/{download_filename}"
        })
    except Exception as e:
        logger.error(f"Error in transcribe_audio_cli: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        await async_rmtree(work_dir)  # aszinkron törlés
