import uuid
import json
import os
import asyncio
import logging
import aiohttp
import shutil 
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple, Union

import aiofiles
from fastapi import APIRouter, HTTPException, UploadFile, File, Form, BackgroundTasks
from fastapi.responses import JSONResponse
import tiktoken
from bs4 import BeautifulSoup

from config import (
    TEMP_DIR, 
    SYSTEM_DOWNLOADS, 
    manager, 
    logger, 
    process_semaphore,
    client, 
    sanitize_filename,
    chunk_text_by_tokens
)
from document_processor import ConversionProcessor, temp_manager

logger = logging.getLogger(__name__)
router = APIRouter()

# Globális változók a folyamatok követésére
active_processes = {}

async def translate_chunks(chunks: List[str], target_lang: str, connection_id: str) -> List[str]:
    """Fordítja a szöveg darabokat párhuzamosan, rate-limit és haladás jelzéssel"""
    tasks = []
    
    async def translate_with_rate_limit(chunk: str, index: int) -> Tuple[int, str]:
        """Egy chunk fordítása rate-limit kezeléssel"""
        async with process_semaphore:
            try:
                response = await asyncio.wait_for(
                    client.chat.completions.create(
                        model=os.getenv("MODEL_NAME", "gpt-3.5-turbo"),
                        messages=[
                            {
                                "role": "system",
                                "content": (
                                    "You are a professional translation system. "
                                    "Accurately translate the provided text while preserving its original meaning, "
                                    "formatting, and special characters. Return only the translated text."
                                )
                            },
                            {
                                "role": "user",
                                "content": f"Please translate this text to {target_lang}:\n\n{chunk}"
                            }
                        ],
                        temperature=0.3
                    ),
                    timeout=45
                )
                result = response.choices[0].message.content.strip()
                progress = min(90, ((index + 1) / len(chunks)) * 90)
                await manager.send_progress(connection_id, progress, f"Translating: {index + 1}/{len(chunks)} chunks...")
                await asyncio.sleep(0.2)
                return index, result
            except asyncio.TimeoutError:
                logger.error(f"Timeout during translation of chunk {index}.")
                raise HTTPException(status_code=408, detail="Translation request timed out.")
            except Exception as e:
                logger.error(f"Error translating chunk {index}: {str(e)}")
                raise
    
    for i, chunk in enumerate(chunks):
        tasks.append(asyncio.create_task(translate_with_rate_limit(chunk, i)))
    
    results = await asyncio.gather(*tasks)
    sorted_results = sorted(results, key=lambda x: x[0])  # Eredmények rendezése
    return [result[1] for result in sorted_results]

async def summarize_translated_text(text: str, connection_id: str) -> str:
    """AI összefoglaló készítése a lefordított szövegről"""
    try:
        await manager.send_progress(connection_id, 92, "Generating AI summary...")
        
        response = await client.chat.completions.create(
            model=os.getenv("MODEL_NAME", "gpt-3.5-turbo"),
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Create a comprehensive executive summary of the following document. "
                        "Include key points, findings, and conclusions. Be concise but thorough."
                    )
                },
                {
                    "role": "user",
                    "content": f"Please summarize this text:\n\n{text[:4000]}..." # Limit to first 4000 chars
                }
            ],
            temperature=0.3,
            max_tokens=1000
        )
        
        summary = response.choices[0].message.content.strip()
        await manager.send_progress(connection_id, 95, "AI summary generated!")
        return summary
    except Exception as e:
        logger.error(f"Summary generation error: {str(e)}")
        return ""  # Return empty if failed, this is non-critical

async def fetch_url_content(url: str, connection_id: str) -> str:
    """Szöveg letöltése URL-ből"""
    try:
        await manager.send_progress(connection_id, 10, f"Fetching content from URL: {url}")
        
        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                if response.status != 200:
                    raise HTTPException(
                        status_code=400, 
                        detail=f"Failed to fetch URL: HTTP {response.status}"
                    )
                
                content_type = response.headers.get('Content-Type', '').lower()
                
                if 'text/html' in content_type:
                    html = await response.text()
                    soup = BeautifulSoup(html, 'html.parser')
                    
                    # Remove scripts, styles, and other non-content elements
                    for element in soup(['script', 'style', 'header', 'footer', 'nav']):
                        element.decompose()
                    
                    # Extract text from main content areas
                    main_content = soup.find('main') or soup.find('article') or soup.find('body')
                    if main_content:
                        text = main_content.get_text(separator='\n', strip=True)
                    else:
                        text = soup.get_text(separator='\n', strip=True)
                    
                    # Clean up excessive whitespace
                    import re
                    text = re.sub(r'\n+', '\n', text)
                    text = re.sub(r'\s+', ' ', text)
                    
                    return text
                elif 'application/pdf' in content_type:
                    # For PDFs, we download and process with our document processor
                    pdf_data = await response.read()
                    work_dir = TEMP_DIR / connection_id
                    work_dir.mkdir(exist_ok=True)
                    pdf_path = work_dir / "downloaded.pdf"
                    
                    async with aiofiles.open(pdf_path, "wb") as f:
                        await f.write(pdf_data)
                    
                    text, _ = await ConversionProcessor.process_pdf(pdf_path)
                    return text
                else:
                    # Just plain text or unknown
                    return await response.text()
    except aiohttp.ClientError as e:
        logger.error(f"URL fetch error: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Failed to fetch URL: {str(e)}")

@router.post("/")
async def translate_content(
    file: Optional[UploadFile] = File(None),
    text: Optional[str] = Form(None),
    url: Optional[str] = Form(None),
    target_lang: str = Form("en"),
    target_format: str = Form("txt"),
    preserve_format: bool = Form(False),
    generate_summary: bool = Form(False),
    connection_id: Optional[str] = Form(None),
    background_tasks: BackgroundTasks = None
):
    """
    Fordítja a megadott fájlt, szöveget vagy URL-t a célnyelvre
    
    Paraméterek:
    - file: Opcionális feltöltött fájl
    - text: Opcionális szöveg fordításra
    - url: Opcionális webcím, ahonnan a szöveget kinyeri
    - target_lang: Célnyelv kódja (pl. "en", "hu", "de")
    - target_format: Kimeneti fájlformátum
    - preserve_format: Megtartsa-e az eredeti formázást
    - generate_summary: Készítsen-e AI összefoglalót
    - connection_id: Kapcsolat azonosító a folyamat követéséhez
    """
    source_count = sum(1 for x in [file, text, url] if x)
    if source_count == 0:
        raise HTTPException(status_code=400, detail="Either file, text, or URL must be provided")
    elif source_count > 1:
        raise HTTPException(status_code=400, detail="Only one source (file, text, or URL) should be provided")

    if not connection_id:
        connection_id = str(uuid.uuid4())
    
    active_processes[connection_id] = True
    
    # Ideiglenes munkakönyvtár létrehozása
    async with temp_manager.temp_dir(connection_id) as work_dir:
        original_filename = "text"
        format_info = None
        content_to_translate = ""
        source_type = "text"

        try:
            # 1. Forrás feldolgozása és szöveg kinyerése
            if file:
                source_type = "file"
                original_filename = sanitize_filename(file.filename)
                file_path = work_dir / original_filename
                async with aiofiles.open(file_path, "wb") as f:
                    await f.write(await file.read())
                
                await manager.send_progress(connection_id, 10, "Processing file...")
                
                # MIME típus normalizálása
                ext = Path(file.filename).suffix.lower()[1:]
                
                # JAVÍTÁS: TXT fájl felismerése még akkor is, ha a content_type üres
                if ext == "txt" or (file.content_type and "text/plain" in file.content_type):
                    file_type = "txt"
                # Képfájlok azonosítása kiterjesztés alapján
                elif ext in ["jpg", "jpeg", "png", "gif", "bmp", "tiff", "mpo"]:
                    mapping = {
                        "jpg": "image/jpeg", "jpeg": "image/jpeg", "mpo": "image/jpeg",
                        "png": "image/png", "gif": "image/gif",
                        "bmp": "image/bmp", "tiff": "image/tiff"
                    }
                    file_type = mapping.get(ext, f"image/{ext}")
                elif file.content_type:
                    # Használjuk a kapott content_type-ot
                    file_type = file.content_type
                else:
                    # Fallback: használjuk a kiterjesztést
                    file_type = ext
                
                logger.info(f"File type determined as: {file_type} for file {file.filename}")
                
                # Képfájlok kezelése OCR-rel
                if file_type.startswith("image/"):
                    try:
                        from document_processor import process_image
                        content_to_translate, format_info = await process_image(file_path)
                    except ImportError:
                        raise HTTPException(
                            status_code=400, 
                            detail="OCR processing is not available. Cannot translate image files directly."
                        )
                # TXT fájlok közvetlen kezelése
                elif file_type == "txt":
                    logger.info(f"Processing TXT file directly: {file_path}")
                    try:
                        # TXT fájl közvetlen olvasása UTF-8 kódolással
                        async with aiofiles.open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                            content_to_translate = await f.read()
                        format_info = {"type": "txt"}
                    except Exception as e:
                        logger.error(f"Error reading TXT file: {str(e)}")
                        raise HTTPException(status_code=500, detail=f"Error reading text file: {str(e)}")
                else:
                    # Egyéb fájlok feldolgozása
                    try:
                        content_to_translate, format_info = await ConversionProcessor.process_file(
                            file_path, file_type, preserve_format
                        )
                    except HTTPException as e:
                        if e.status_code == 400 and "Unsupported file type" in e.detail:
                            # Próbáljuk meg átalakítani egy kezelhető formátumba
                            logger.info(f"Attempting to convert unsupported file type: {file_type}")
                            temp_output = work_dir / f"converted_temp.txt"
                            
                            try:
                                await ConversionProcessor.convert_document(
                                    file_path, temp_output, ext, "txt"
                                )
                                
                                async with aiofiles.open(temp_output, "r", encoding="utf-8", errors="ignore") as f:
                                    content_to_translate = await f.read()
                                    
                                format_info = {"type": "txt"}
                            except Exception as conv_e:
                                logger.error(f"Conversion attempt failed: {str(conv_e)}")
                                raise e  # Raise the original error if conversion fails
                        else:
                            raise
            elif text:
                source_type = "text"
                await manager.send_progress(connection_id, 10, "Processing input text...")
                content_to_translate = text
                format_info = {"type": "txt"}
            elif url:
                source_type = "url"
                content_to_translate = await fetch_url_content(url, connection_id)
                original_filename = sanitize_filename(Path(url).name or "webpage")
                format_info = {"type": "txt"}

            if not content_to_translate:
                raise HTTPException(status_code=400, detail="No content found to translate")

            # 2. Szöveg darabolása és fordítása
            chunks = chunk_text_by_tokens(content_to_translate)
            await manager.send_progress(connection_id, 30, f"Translating content to {target_lang}...")
            translated_chunks = await translate_chunks(chunks, target_lang, connection_id)
            translated_text = "\n".join(translated_chunks)
            
            # 3. Opcionális AI összefoglaló
            summary = ""
            if generate_summary:
                summary = await summarize_translated_text(translated_text, connection_id)

            # 4. Eredmények mentése
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            name_without_ext = Path(original_filename).stem

            # 4.1 Transcript adatok mentése
            transcript_data = {
                "text": translated_text,
                "summary": summary,
                "segments": format_info.get("segments", []) if format_info else [],
                "original_filename": str(original_filename),
                "source_type": source_type
            }
            
            # Munkamappában mentjük a transcript_data.json fájlt (a későbbi /download_transcript végpont számára)
            transcript_data_path = work_dir / "transcript_data.json"
            async with aiofiles.open(transcript_data_path, "w", encoding="utf-8") as f:
                await f.write(json.dumps(transcript_data, ensure_ascii=False))
            
            # 4.2 Letölthető txt formátumú transzkript fájl a letöltési mappában
            transcript_filename = f"transcript_{name_without_ext}_{timestamp}.txt"
            transcript_path = SYSTEM_DOWNLOADS / transcript_filename
            async with aiofiles.open(transcript_path, "w", encoding="utf-8") as f:
                await f.write(translated_text)
                if summary:
                    await f.write("\n\n--- AI SUMMARY ---\n\n")
                    await f.write(summary)

            # 5. Eredmény fájl generálása a kért formátumban
            download_url = None
            if target_format.lower() != "txt":
                output_filename = f"translated_{name_without_ext}_{timestamp}.{target_format.lower()}"
                output_path = SYSTEM_DOWNLOADS / output_filename
                
                await manager.send_progress(connection_id, 92, f"Converting to {target_format}...")
                source_format = format_info.get("type", "txt") if format_info else "txt"
                
                # JAVÍTÁS: Ha DocX/Doc/ODT a cél, használjunk direkt konverziós megoldást
                if target_format.lower() in ["docx", "doc", "odt"]:
                    try:
                        # Készítsünk egy ideiglenes TXT fájlt a fordításból
                        temp_txt_path = work_dir / f"translated_text_{timestamp}.txt"
                        async with aiofiles.open(temp_txt_path, "w", encoding="utf-8") as f:
                            await f.write(translated_text)
                            
                        if target_format.lower() == "docx":
                            from docx import Document
                            
                            def create_docx():
                                doc = Document()
                                # Soronként adjuk hozzá a szöveget, hogy az ékezetek megmaradjanak
                                paragraphs = translated_text.split('\n')
                                for para in paragraphs:
                                    if para.strip():
                                        doc.add_paragraph(para)
                                    else:
                                        doc.add_paragraph()  # üres bekezdés
                                doc.save(str(output_path))
                                return True
                            
                            loop = asyncio.get_event_loop()
                            await loop.run_in_executor(None, create_docx)
                        else:
                            # Más formátumok esetén használjuk a ConversionProcessor-t
                            await ConversionProcessor.convert_document(
                                temp_txt_path, output_path, "txt", target_format.lower()
                            )
                    except Exception as e:
                        logger.error(f"Direct format conversion error: {str(e)}")
                        # Fallback a normál konverzióhoz hiba esetén
                        await ConversionProcessor.convert_document(
                            temp_txt_path, output_path, "txt", target_format.lower()
                        )
                else:
                    # Nem DocX/Doc/ODT cél - normál konverzió
                    # Készítsünk egy ideiglenes TXT fájlt a fordításból
                    temp_txt_path = work_dir / f"translated_text_{timestamp}.txt"
                    async with aiofiles.open(temp_txt_path, "w", encoding="utf-8") as f:
                        await f.write(translated_text)
                        
                    await ConversionProcessor.convert_document(
                        temp_txt_path, output_path, "txt", target_format.lower()
                    )
                
                download_url = f"/download/{output_filename}"
            else:
                # Alapértelmezett TXT kimenet
                output_filename = f"translated_{name_without_ext}_{timestamp}.txt"
                output_path = SYSTEM_DOWNLOADS / output_filename
                
                # Ha már van transcript, csak másoljuk át
                if Path(transcript_path).exists():
                    shutil.copy(transcript_path, output_path)
                else:
                    async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
                        await f.write(translated_text)
                        if summary:
                            await f.write("\n\n--- AI SUMMARY ---\n\n")
                            await f.write(summary)
                
                download_url = f"/download/{output_filename}"

            await manager.send_progress(connection_id, 100, "Translation complete")
            
            # 6. Válasz előkészítése
            response_data = {
                "download_url": download_url,
                "transcript_download_url": f"/download/{transcript_filename}",
                "api_transcript_url": f"/download_transcript/{connection_id}/txt"
            }
            
            if summary:
                response_data["summary"] = summary
                
            return JSONResponse(response_data)
        
        except Exception as e:
            logger.error(f"Translation error: {str(e)}")
            await manager.send_progress(connection_id, 100, f"Error: {str(e)}")
            raise HTTPException(status_code=500, detail=str(e))
        
        finally:
            active_processes.pop(connection_id, None)
            # Ideiglenes könyvtár törlése a háttérben, ha van BackgroundTasks
            if background_tasks:
                background_tasks.add_task(temp_manager.cleanup, connection_id)
